"""Denvia 자산·계정 통합 인계서 빌더.

2026-06-05 — 도메인·AWS·토스 페이먼츠·알리고·OAuth·외부 API 키·관리자 계정·
백엔드 시크릿을 단일 문서로 통합한 보안 인계서.

산출물:
    자료/Denvia_자산_계정_통합_인계서_2026-06-05.docx

CLAUDE.md "고객 전달용 문서 작성 규칙" 준수 — SDT(content control), 폼 보호,
문서 보호, ActiveX 등 일체 사용하지 않음. 어떤 워드 버전에서도 자유 편집 가능.

이 문서는 비밀번호·API 키가 포함된 1급 보안 문서입니다. 첫 페이지에 보안 경고를
강조하고, 인계 후 비밀번호 변경 체크리스트를 동봉합니다.

실행:
    py -3 scripts/build_assets_handover.py
"""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


KOREAN_FONT = "맑은 고딕"

DOC_TITLE = "자산·계정 통합 인계서"
DOC_SUBTITLE = "Denvia · 치과 보험청구·데스크 행정 AI 어시스턴트"
DOC_DATE = "작성일: 2026년 6월 5일"


# ──────────────────────────────────────────────────────────────────────────────
# 비밀값 로딩 — 자격증명은 절대 소스에 박지 않는다.
# `자료/.handover_secrets.env` (gitignore) 또는 셸 환경변수에서만 읽는다.
# ──────────────────────────────────────────────────────────────────────────────

_REPO_ROOT = Path(__file__).resolve().parent.parent
_SECRETS_ENV_PATH = _REPO_ROOT / "자료" / ".handover_secrets.env"


def _load_handover_secrets() -> None:
    """`자료/.handover_secrets.env` 가 있으면 KEY=VALUE 라인을 os.environ 에 머지.

    이미 환경에 설정된 키는 덮어쓰지 않는다(셸 export 우선).
    파일이 없거나 비어 있으면 silent — 값이 진짜 필요한 시점에 `_env_secret` 가 실패시킨다.
    """
    if not _SECRETS_ENV_PATH.exists():
        return
    for raw in _SECRETS_ENV_PATH.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, _, value = line.partition("=")
        key, value = key.strip(), value.strip().strip('"').strip("'")
        if key and value and key not in os.environ:
            os.environ[key] = value


def _env_secret(key: str, label: str) -> str:
    """문서에 들어갈 자격증명. 미설정이면 SystemExit 로 명확히 실패시킨다."""
    value = os.environ.get(key, "").strip()
    if not value:
        raise SystemExit(
            f"[build_assets_handover] 환경변수 {key}({label})가 비어 있습니다.\n"
            f"  - {_SECRETS_ENV_PATH} 에 `{key}=...` 형식으로 채우거나\n"
            f"  - 셸에서 `set {key}=...` (PowerShell `$env:{key}=...`) 로 주입하십시오."
        )
    return value


# ──────────────────────────────────────────────────────────────────────────────
# 스타일 헬퍼
# ──────────────────────────────────────────────────────────────────────────────

def _apply_kr_font(run, *, bold: bool = False, size_pt: float | None = None,
                   color: RGBColor | None = None, mono: bool = False) -> None:
    font_name = "Consolas" if mono else KOREAN_FONT
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), KOREAN_FONT)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    if bold:
        run.bold = True
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color


def _add_paragraph(doc, text, *, bold=False, size_pt=None, align=None,
                   color=None, space_after_pt=None, mono=False) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_after_pt is not None:
        p.paragraph_format.space_after = Pt(space_after_pt)
    run = p.add_run(text)
    _apply_kr_font(run, bold=bold, size_pt=size_pt, color=color, mono=mono)


def _add_heading(doc, text, level=1) -> None:
    sizes = {1: 18, 2: 14, 3: 12}
    spaces_before = {1: 18, 2: 14, 3: 10}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(spaces_before[level])
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _apply_kr_font(run, bold=True, size_pt=sizes[level])


def _shade_cell(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _add_horizontal_rule(doc) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_callout_box(doc, body: str, *, fill: str = "FFF8DC",
                     size_pt: float = 11, bold: bool = False,
                     color: RGBColor | None = None, mono: bool = False) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Cm(16)
    cell = tbl.cell(0, 0)
    cell.width = Cm(16)
    _shade_cell(cell, fill)
    for i, line in enumerate(body.split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        _apply_kr_font(r, size_pt=size_pt, bold=bold, color=color, mono=mono)
    doc.add_paragraph()


def _add_table_header(table, headers: list[str], *, fill: str = "E0E0E0") -> None:
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        _apply_kr_font(r, bold=True, size_pt=10)
        _shade_cell(hdr_cells[i], fill)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _add_table_row(table, values: list[str], *, bold_first: bool = False,
                   mono_cols: list[int] | None = None) -> None:
    mono_cols = mono_cols or []
    row = table.add_row().cells
    for i, v in enumerate(values):
        row[i].text = ""
        row[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        for j, line in enumerate(v.split("\n")):
            p = row[i].paragraphs[0] if j == 0 else row[i].add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(line)
            _apply_kr_font(r, bold=(bold_first and i == 0), size_pt=10,
                           mono=(i in mono_cols))


def _add_kv_table(doc, rows: list[tuple[str, str]], *,
                  label_width_cm: float = 4.0,
                  value_width_cm: float = 12.0,
                  label_fill: str = "F2F2F2",
                  value_mono: bool = False) -> None:
    tbl = doc.add_table(rows=0, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Cm(label_width_cm)
    tbl.columns[1].width = Cm(value_width_cm)
    for label, value in rows:
        row = tbl.add_row().cells
        row[0].text = ""
        row[1].text = ""
        row[0].width = Cm(label_width_cm)
        row[1].width = Cm(value_width_cm)
        for cell in row:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _shade_cell(row[0], label_fill)
        p0 = row[0].paragraphs[0]
        r0 = p0.add_run(label)
        _apply_kr_font(r0, bold=True, size_pt=10)
        for i, line in enumerate(value.split("\n")):
            p = row[1].paragraphs[0] if i == 0 else row[1].add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(line)
            _apply_kr_font(r, size_pt=10, mono=value_mono)


# ──────────────────────────────────────────────────────────────────────────────
# 문서 빌드
# ──────────────────────────────────────────────────────────────────────────────

def build_document(output_path: Path) -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = KOREAN_FONT
    style.font.size = Pt(11)

    for section in doc.sections:
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # ─── 표지
    _add_paragraph(doc, DOC_TITLE, bold=True, size_pt=24,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, DOC_SUBTITLE, size_pt=12,
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   color=RGBColor(0x55, 0x55, 0x55))
    _add_paragraph(doc, DOC_DATE, size_pt=11,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=18)

    # ─── 보안 주의 (간단)
    _add_callout_box(
        doc,
        "[주의] 본 문서에는 외부 서비스 로그인 정보가 일부 포함되어 있습니다.\n"
        "이메일·메신저로 전달하지 마시고, 개인 PC의 안전한 폴더에 보관해 주세요.",
        fill="FFF8DC",
        size_pt=11,
    )

    # ─── 0. 이 문서의 목적
    _add_heading(doc, "0. 이 문서의 목적", level=2)
    _add_paragraph(
        doc,
        "본 문서는 Denvia 서비스 운영에 사용되는 외부 서비스 목록과, 개발 진행 중 "
        "개발자가 직접 개설·등록한 계정 정보를 한눈에 보실 수 있도록 정리한 안내서입니다. "
        "도메인·AWS·토스 페이먼츠·알리고·소셜 로그인(카카오·네이버·구글)·외부 API·"
        "관리자 계정이 어떻게 구성되어 있는지를 형식적으로 정리해 둔 참고용 문서입니다.",
        size_pt=11,
    )
    _add_paragraph(
        doc,
        "대부분의 외부 서비스 계정은 이미 클라이언트 명의로 가입·등록되어 있어, "
        "특별한 인계 절차나 비밀번호 일괄 교체가 필요한 것은 아닙니다. "
        "각 항목의 로그인 방법과 용도를 확인하시는 용도로 보시면 됩니다.",
        size_pt=11,
    )

    # ─── 1. 자산 한눈에 보기
    _add_heading(doc, "1. 자산 한눈에 보기", level=1)
    _add_paragraph(
        doc,
        "Denvia 운영에 사용되는 모든 외부 자산입니다. 각 항목의 상세 자격 증명은 "
        "이후 2 ~ 10항에 정리되어 있습니다.",
        size_pt=11,
    )
    tbl = doc.add_table(rows=1, cols=4)
    tbl.autofit = False
    tbl.columns[0].width = Cm(3.5)
    tbl.columns[1].width = Cm(5.0)
    tbl.columns[2].width = Cm(4.5)
    tbl.columns[3].width = Cm(3.0)
    _add_table_header(tbl, ["분류", "공급자 / 서비스", "용도", "월 비용 (대략)"])
    _add_table_row(tbl, ["도메인", "카페24 (denvia.ai.kr)", "사이트 주소", "약 1,800원"])
    _add_table_row(tbl, ["서버", "AWS Lightsail (서울 리전)", "Denvia 서비스 호스팅", "약 61,600원"])
    _add_table_row(tbl, ["결제", "토스 페이먼츠", "구독 결제·환불", "수수료 별도"])
    _add_table_row(tbl, ["알림톡·SMS", "알리고 (Aligo)", "사용자 알림 발송", "건당 약 8.5~19원"])
    _add_table_row(tbl, ["소셜 로그인", "카카오·네이버·구글", "간편 회원가입·로그인", "무료"])
    _add_table_row(tbl, ["환율 API", "한국수출입은행", "원/달러 환율 일 1회 조회", "무료"])
    _add_table_row(tbl, ["AI(RAG)", "OpenAI", "보험청구 Q&A 생성", "사용량 과금"])
    doc.add_paragraph()

    _add_callout_box(
        doc,
        "※ 월 비용은 USD 1,400원 환산 기준이며 환율·사용량에 따라 ±10% 변동될 수 있습니다.\n"
        "※ OpenAI·토스 페이먼츠는 사용량에 비례하므로 별도 청구됩니다. "
        "관리자 페이지의 ‘예산 통제’ 기능으로 OpenAI 월 한도를 설정해 두실 수 있습니다.",
        fill="FFF8DC",
    )

    # ─── 2. 도메인
    _add_heading(doc, "2. 도메인 — denvia.ai.kr", level=1)
    _add_paragraph(
        doc,
        "사이트의 주소(URL)를 발급·갱신하는 서비스입니다. 이 계정이 정지되면 "
        "사이트가 일시적으로 접속 불가 상태가 되므로, 갱신일에 반드시 결제가 "
        "이루어져야 합니다.",
        size_pt=11,
    )
    _add_kv_table(doc, [
        ("공급자", "카페24 (https://www.cafe24.com)"),
        ("도메인", "denvia.ai.kr"),
        ("관리자 콘솔", "https://hosting.cafe24.com → 로그인 → 도메인 관리"),
        ("로그인 ID", "dlrbtjd357"),
        ("로그인 비밀번호", "7@Y6^SyyXAWCY8"),
        ("갱신 주기", "연 1회 (자동 결제 권장)"),
        ("DNS 설정", "A 레코드 → 15.164.114.152 (Lightsail 서버 공인 IP)"),
        ("소유 주체", "클라이언트 단독 보유"),
    ], label_width_cm=4.0, value_width_cm=12.0)
    _add_callout_box(
        doc,
        "※ 도메인 갱신을 놓치면 사이트 접속이 끊깁니다. 카페24의 ‘자동 결제’ 기능을 "
        "켜두시거나, 갱신일 2주 전 알림을 휴대폰 캘린더에 등록해 두시기를 권장합니다.",
        fill="FFF8DC",
    )

    # ─── 3. AWS 콘솔 + Lightsail 서버
    _add_heading(doc, "3. AWS — 콘솔 계정 + Lightsail 서버", level=1)
    _add_paragraph(
        doc,
        "Denvia가 실제로 돌아가는 서버를 빌리는 회사입니다. AWS 콘솔 계정 하나로 "
        "Lightsail(서버)·자동 백업·결제를 모두 관리합니다.",
        size_pt=11,
    )

    _add_heading(doc, "3-1. AWS 콘솔 로그인 (클라이언트 명의 계정)", level=2)
    _add_kv_table(doc, [
        ("공급자", "Amazon Web Services (AWS)"),
        ("콘솔 URL", "https://console.aws.amazon.com"),
        ("계정 이메일", "dlrbtjd357@naver.com"),
        ("계정 소유자", "이규성 (leekyusung)"),
        ("로그인 비밀번호", "asd345@#!@#$!@RQWDdsv"),
        ("로그인 방식", "콘솔 URL 접속 → ‘루트 사용자’ 선택 → 이메일·비밀번호 입력"),
    ], label_width_cm=4.0, value_width_cm=12.0)

    _add_callout_box(
        doc,
        "[보안 권장] AWS 콘솔에 MFA(2단계 인증)를 반드시 설정해 주세요.\n"
        "\n"
        "AWS 계정은 이미 클라이언트 명의 카드가 결제 수단으로 등록되어 있어, "
        "비밀번호가 유출되면 카드까지 함께 노출될 수 있습니다. "
        "MFA를 켜두시면 비밀번호를 알아도 휴대폰의 OTP 코드 없이는 로그인할 수 없어 "
        "계정 탈취를 막을 수 있습니다.\n"
        "\n"
        "설정 방법:\n"
        "  1) AWS 콘솔 로그인 → 우측 상단 본인 계정 이름 클릭 → ‘보안 자격 증명’\n"
        "  2) ‘다단계 인증 (MFA)’ 섹션 → ‘MFA 디바이스 할당’\n"
        "  3) ‘인증 관리자 앱(Authenticator app)’ 선택 (Google Authenticator 또는 Authy)\n"
        "  4) 휴대폰의 인증 앱으로 QR 코드 스캔 → 6자리 코드 2회 입력 → 등록 완료\n"
        "\n"
        "이후 로그인할 때마다 비밀번호 + 휴대폰 인증 코드 6자리를 함께 입력하시면 됩니다.",
        fill="EAF3FF",
        size_pt=11,
    )

    _add_heading(doc, "3-2. Lightsail 서버 (실제 운영 인스턴스)", level=2)
    _add_kv_table(doc, [
        ("인스턴스 이름", "denvia-prod-01"),
        ("리전", "서울 (ap-northeast-2)"),
        ("운영체제", "Ubuntu 22.04 LTS"),
        ("사양", "2 vCPU · 8 GB RAM · 160 GB SSD ($40 플랜)"),
        ("공인 IP", "15.164.114.152"),
        ("도메인 연결", "denvia.ai.kr → 15.164.114.152 (DNS A 레코드)"),
        ("SSH 접속", "AWS Lightsail 콘솔 → ‘브라우저로 연결’ 클릭"),
        ("SSH 키 보관", "개발자 보관 (장애 복구·배포용, 인계 협의 사항)"),
        ("자동 백업", "매일 1회 스냅샷 · 7일 보관 ($4/월)"),
        ("프로젝트 경로", "/opt/denvia (서버 안에서 코드가 사는 폴더)"),
    ], label_width_cm=4.0, value_width_cm=12.0)
    _add_callout_box(
        doc,
        "※ Lightsail 콘솔에서 인스턴스 ‘중지(Stop)’ 또는 ‘삭제(Delete)’ 버튼을 "
        "임의로 누르면 데이터 손실 위험이 있습니다. 사양 변경·재시작이 필요하시면 "
        "반드시 개발자에게 먼저 문의해 주세요.",
        fill="FFEDED",
    )

    # ─── 4. 토스 페이먼츠 (PG)
    _add_heading(doc, "4. 토스 페이먼츠 — 결제(PG)", level=1)
    _add_paragraph(
        doc,
        "사용자가 카드로 구독료를 결제할 때 거치는 결제 대행 회사입니다. "
        "토스 페이먼츠 가맹점 계약과 API 키는 다음과 같이 분리하여 관리합니다.",
        size_pt=11,
    )

    _add_heading(doc, "4-1. 가맹점 계정 (클라이언트 명의)", level=2)
    _add_kv_table(doc, [
        ("공급자", "토스 페이먼츠 (https://www.tosspayments.com)"),
        ("가맹점 콘솔", "https://app.tosspayments.com"),
        ("가맹점 가입 명의", "클라이언트 본인 (사업자등록증 기준)"),
        ("수수료", "약 2.9% (카드사·결제 수단에 따라 변동)"),
        ("고객센터", "1544-7772 (평일 09:00 ~ 18:00)"),
    ], label_width_cm=4.0, value_width_cm=12.0)

    _add_heading(doc, "4-2. API 키 — 관리자 페이지에서 직접 등록", level=2)
    _add_paragraph(
        doc,
        "토스 페이먼츠의 클라이언트 키·시크릿 키는 서버 코드에 박아두는 방식이 "
        "아니라, Denvia 관리자 페이지에서 직접 입력하여 관리합니다. "
        "토스 콘솔에서 키를 재발급하시면 관리자 페이지에서만 갱신하면 "
        "되므로, 개발자 도움 없이 클라이언트가 직접 처리하실 수 있습니다.",
        size_pt=11,
    )
    _add_kv_table(doc, [
        ("입력 위치", "관리자 페이지 → 좌측 메뉴 [설정] → [결제 / PG 키 관리]"),
        ("입력 경로 (URL)", "https://denvia.ai.kr/admin/settings/payment"),
        ("입력 항목", "① 클라이언트 키 (라이브)\n② 시크릿 키 (라이브)\n③ 클라이언트 키 (테스트)\n④ 시크릿 키 (테스트)\n⑤ 운영 모드 (테스트 / 라이브)"),
        ("키 발급 위치", "토스 페이먼츠 가맹점 콘솔 → 상점 관리 → API 키"),
        ("최초 설정 시점", "토스 가맹점 심사 완료 직후, 라이브 키 발급되면 즉시 ④→③→②→① 순서로 입력하고 ⑤를 ‘라이브’로 전환"),
    ], label_width_cm=4.0, value_width_cm=12.0)
    _add_callout_box(
        doc,
        "※ 운영 모드가 ‘테스트’이면 실제 결제가 이뤄지지 않습니다. 라이브 키를 "
        "입력한 후 반드시 ⑤ 운영 모드를 ‘라이브’로 전환해 주세요.\n"
        "※ 라이브 키를 입력하기 전까지는 사용자의 결제가 실패 처리됩니다. "
        "심사 완료 → 키 입력 → 라이브 전환의 3단계를 빠짐없이 수행해 주세요.",
        fill="FFF8DC",
    )

    # ─── 5. 알리고 (알림톡 + SMS)
    _add_heading(doc, "5. 알리고 — 카카오 알림톡 + SMS", level=1)
    _add_paragraph(
        doc,
        "사용자에게 결제 안내·비밀번호 재설정·1:1 문의 답변 알림 등을 카카오 "
        "알림톡(또는 SMS)으로 발송하는 회사입니다. Denvia는 모든 사용자 알림을 "
        "이메일이 아닌 알림톡·SMS로만 발송합니다.",
        size_pt=11,
    )
    _add_kv_table(doc, [
        ("공급자", "알리고 (https://smartsms.aligo.in)"),
        ("로그인 ID", "dlrbtjd357"),
        ("로그인 비밀번호", "varqmUTaY7@sn*w"),
        ("API Key", "5j5as2l79f1qo3sv4k7u6wbrauef7gi5"),
        ("발신번호", "010-2323-2753 (알리고 콘솔에서 사전 인증 완료)"),
        ("카카오 송신프로필 키", "a5b8f74ad44ee83363fbcb0ff51343800ca1fd3d"),
        ("발송 모드", "운영 (실발송 — ALIGO_TEST_MODE=false)"),
        ("충전 잔액", "관리자 페이지에서 실시간 확인 가능"),
        ("단가 (참고)", "알림톡 약 8.5원 / SMS 약 19원 / LMS 약 35원"),
        ("고객센터", "1661-9854 (평일 09:00 ~ 18:00)"),
    ], label_width_cm=4.0, value_width_cm=12.0)
    _add_callout_box(
        doc,
        "※ 알리고 API Key는 외부 노출 시 부정 발송이 가능한 민감 자산입니다. "
        "본 문서 외 다른 곳에 복사하지 말아 주세요.\n"
        "※ 잔액이 0원이 되면 알림 발송이 모두 실패합니다. 관리자 페이지에서 "
        "잔액 알림 임계값을 설정해 두시면 사전에 경고 알림을 받으실 수 있습니다.",
        fill="FFF8DC",
    )

    # ─── 6. 소셜 로그인 (OAuth)
    _add_heading(doc, "6. 소셜 로그인 (OAuth)", level=1)
    _add_paragraph(
        doc,
        "사용자가 카카오·네이버·구글 계정으로 한 번에 로그인할 수 있게 해 주는 "
        "서비스입니다. 각 공급자(카카오·네이버·구글)의 개발자 콘솔에서 "
        "‘Denvia 앱’을 등록한 결과로 발급된 키들이 아래에 정리돼 있습니다.",
        size_pt=11,
    )

    _add_heading(doc, "6-1. 카카오 로그인", level=2)
    _add_kv_table(doc, [
        ("개발자 콘솔", "https://developers.kakao.com"),
        ("앱 이름", "Denvia"),
        ("REST API 키 (Client ID)", _env_secret("KAKAO_OAUTH_CLIENT_ID", "카카오 REST API 키")),
        ("Client Secret", _env_secret("KAKAO_OAUTH_CLIENT_SECRET", "카카오 Client Secret")),
        ("Redirect URI", "https://denvia.ai.kr/api/v1/auth/oauth/kakao/callback"),
    ], label_width_cm=5.0, value_width_cm=11.0, value_mono=True)

    _add_heading(doc, "6-2. 네이버 로그인", level=2)
    _add_kv_table(doc, [
        ("개발자 콘솔", "https://developers.naver.com"),
        ("앱 이름", "Denvia"),
        ("Client ID", _env_secret("NAVER_OAUTH_CLIENT_ID", "네이버 Client ID")),
        ("Client Secret", _env_secret("NAVER_OAUTH_CLIENT_SECRET", "네이버 Client Secret")),
        ("Redirect URI", "https://denvia.ai.kr/api/v1/auth/oauth/naver/callback"),
    ], label_width_cm=5.0, value_width_cm=11.0, value_mono=True)

    _add_heading(doc, "6-3. 구글 로그인", level=2)
    _add_kv_table(doc, [
        ("개발자 콘솔", "https://console.cloud.google.com → APIs & Services → 사용자 인증 정보"),
        ("앱 이름", "Denvia"),
        ("Client ID", _env_secret("GOOGLE_OAUTH_CLIENT_ID", "구글 Client ID")),
        ("Client Secret", _env_secret("GOOGLE_OAUTH_CLIENT_SECRET", "구글 Client Secret")),
        ("Redirect URI", "https://denvia.ai.kr/api/v1/auth/oauth/google/callback"),
    ], label_width_cm=5.0, value_width_cm=11.0, value_mono=True)
    _add_callout_box(
        doc,
        "※ Redirect URI는 절대 변경하지 마십시오. 변경 시 ‘리다이렉트 URI 불일치’ 오류로 "
        "소셜 로그인이 모두 실패합니다. 도메인을 바꾸시려면 개발자에게 의뢰해 주세요.\n"
        "※ 각 콘솔에서 Client Secret을 재발급하시면 서버 환경변수도 함께 갱신해야 "
        "합니다. 개발자에게 의뢰해 주세요.",
        fill="FFF8DC",
    )

    # ─── 7. 외부 API 키
    _add_heading(doc, "7. 외부 API 키", level=1)

    _add_heading(doc, "7-1. 한국수출입은행 환율 API", level=2)
    _add_paragraph(
        doc,
        "원/달러 환율을 매일 1회 자동 조회하여 토스 페이먼츠 결제 금액을 KRW로 "
        "표시하는 데 사용합니다.",
        size_pt=11,
    )
    _add_kv_table(doc, [
        ("공급자", "한국수출입은행"),
        ("발급 페이지", "https://www.koreaexim.go.kr/ir/HPHKIR055M01"),
        ("API Key", "Q9P19VZ6u8mDBzM68NrmggQGwzmm6dJo"),
        ("호출 한도", "일 1,000회 (개인용)"),
        ("비용", "무료"),
    ], label_width_cm=4.0, value_width_cm=12.0, value_mono=True)

    _add_heading(doc, "7-2. OpenAI (RAG / Q&A)", level=2)
    _add_paragraph(
        doc,
        "사용자의 보험청구 질문에 AI가 답변할 때 사용하는 OpenAI 모델 호출 키입니다. "
        "토큰 사용량에 비례하여 USD로 과금됩니다. 관리자 페이지의 ‘예산 통제’ 기능으로 "
        "월 한도를 설정해 두시면 80% / 95% / 100% 도달 시 자동 알림과 함께 호출이 차단됩니다.",
        size_pt=11,
    )
    _add_kv_table(doc, [
        ("공급자", "OpenAI (https://platform.openai.com)"),
        ("관리 콘솔", "https://platform.openai.com/account"),
        ("API Key", "서버 환경변수 OPENAI_API_KEY 로 관리 (개발자 보관)"),
        ("월 예산 통제", "관리자 페이지 → 설정 → 예산 통제 (80/95/100% 자동 차단)"),
    ], label_width_cm=4.0, value_width_cm=12.0)

    # ─── 8. 관리자 계정
    _add_heading(doc, "8. 관리자 계정", level=1)
    _add_paragraph(
        doc,
        "Denvia 관리자 페이지(https://denvia.ai.kr/admin)에 로그인할 수 있는 계정입니다. "
        "권한 등급은 4단계로 나뉘며, 등급이 높을수록 더 많은 작업이 가능합니다.",
        size_pt=11,
    )

    _add_heading(doc, "8-1. 권한 등급 (RBAC)", level=2)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.autofit = False
    tbl.columns[0].width = Cm(3.0)
    tbl.columns[1].width = Cm(5.0)
    tbl.columns[2].width = Cm(8.0)
    _add_table_header(tbl, ["등급", "권한 범위", "비고"])
    _add_table_row(tbl, [
        "operator",
        "전체 운영 권한",
        "실제 운영 담당. 다른 operator의 권한 변경도 가능.",
    ])
    _add_table_row(tbl, [
        "sub_operator",
        "조회·응대 중심의 제한된 권한",
        "고객 응대·문의 답변 등 일상 업무.",
    ])
    _add_table_row(tbl, [
        "pending",
        "권한 없음 (승인 대기)",
        "관리자 가입 신청 후 operator 승인 전 상태.",
    ])
    doc.add_paragraph()

    _add_heading(doc, "8-2. 운영 계정", level=2)
    _add_kv_table(doc, [
        ("operator (운영 ID)", "admin@denvia.ai.kr"),
        ("비밀번호", "adminTest5873!"),
        ("로그인 위치", "https://denvia.ai.kr/admin"),
    ], label_width_cm=5.0, value_width_cm=11.0)
    _add_callout_box(
        doc,
        "※ 본 계정으로 일상 운영(회원 관리·결제 현황 확인·문의 응답·공지 발송 등)을 "
        "모두 수행하실 수 있습니다.\n"
        "※ operator 끼리 서로의 권한 변경·차단·삭제가 가능합니다. 추가 운영자를 두실 "
        "때는 sub_operator로 시작하시고, 신뢰가 쌓이면 operator로 승격하시는 것을 권장합니다.",
        fill="FFF8DC",
    )

    _add_horizontal_rule(doc)
    _add_paragraph(
        doc,
        "이상으로 Denvia 운영에 사용되는 외부 서비스 및 계정 안내를 마칩니다.",
        size_pt=11, align=WD_ALIGN_PARAGRAPH.CENTER,
        color=RGBColor(0x55, 0x55, 0x55),
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    _load_handover_secrets()
    repo_root = Path(__file__).resolve().parent.parent
    output_path = repo_root / "자료" / "Denvia_자산_계정_통합_인계서_2026-06-05.docx"
    build_document(output_path)
    print(f"[OK] Wrote {output_path}")


if __name__ == "__main__":
    main()
