"""고객 인계용 로컬 개발환경 셋업 안내서 빌더.

2026-06-10 — 인계받은 git 저장소 압축본을 자신의 PC에서 띄우고 직접 코드를
수정할 수 있게 만들기 위한 안내서. 비개발자(고객사 담당자)도 따라 할 수 있게
모든 명령어·환경변수에 한국어 괄호 설명을 붙인다.

산출물:
    자료/Denvia_로컬_개발환경_셋업_안내서.docx
    자료/Denvia_로컬_개발환경_셋업_안내서.pdf  (docx2pdf 사용, Word 필요)

CLAUDE.md "고객 전달용 문서 작성 규칙" 준수 — SDT(content control), 폼 보호,
문서 보호, ActiveX 등 일체 사용하지 않음. 어떤 워드 버전에서도 자유 편집 가능.

실행:
    py -3 scripts/build_local_setup_guide.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


KOREAN_FONT = "맑은 고딕"
MONO_FONT = "Consolas"

DOC_TITLE = "Denvia 로컬 개발환경 셋업 안내서"
DOC_SUBTITLE = "인계받은 코드 압축본을 내 PC에서 띄우고 수정하는 절차"
DOC_DATE = "작성일: 2026년 6월 10일"


# ──────────────────────────────────────────────────────────────────────────────
# 스타일 헬퍼
# ──────────────────────────────────────────────────────────────────────────────

def _apply_kr_font(run, *, bold: bool = False, size_pt: float | None = None,
                   color: RGBColor | None = None, mono: bool = False) -> None:
    font = MONO_FONT if mono else KOREAN_FONT
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), KOREAN_FONT)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:ascii"), font)
    if bold:
        run.bold = True
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color


def _add_paragraph(doc, text, *, bold=False, size_pt=None, align=None,
                   color=None, space_after_pt=None, mono=False) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_after_pt is not None:
        p.paragraph_format.space_after = Pt(space_after_pt)
    run = p.add_run(text)
    _apply_kr_font(run, bold=bold, size_pt=size_pt, color=color, mono=mono)


def _add_heading(doc, text, level=1) -> None:
    sizes = {1: 18, 2: 14, 3: 12}
    spaces_before = {1: 18, 2: 14, 3: 10}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(spaces_before[level])
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _apply_kr_font(run, bold=True, size_pt=sizes[level])


def _shade_cell(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _add_horizontal_rule(doc) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_callout_box(doc, body: str, *, fill: str = "FFF8DC",
                     size_pt: float = 10.5) -> None:
    """안내 박스(연한 배경의 단일 셀 표)."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Cm(16)
    cell = tbl.cell(0, 0)
    cell.width = Cm(16)
    _shade_cell(cell, fill)
    for i, line in enumerate(body.split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        _apply_kr_font(r, size_pt=size_pt)
    doc.add_paragraph()


def _add_code_box(doc, body: str, *, fill: str = "F4F4F4",
                  size_pt: float = 10) -> None:
    """명령어/코드 박스 (등폭 폰트, 회색 배경)."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Cm(16)
    cell = tbl.cell(0, 0)
    cell.width = Cm(16)
    _shade_cell(cell, fill)
    for i, line in enumerate(body.split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        _apply_kr_font(r, size_pt=size_pt, mono=True)
    doc.add_paragraph()


def _add_table_header(table, headers: list[str], *, fill: str = "E0E0E0") -> None:
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        _apply_kr_font(r, bold=True, size_pt=10)
        _shade_cell(hdr_cells[i], fill)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _add_table_row(table, values: list[str], *, bold_first: bool = False) -> None:
    row = table.add_row().cells
    for i, v in enumerate(values):
        row[i].text = ""
        row[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        for j, line in enumerate(v.split("\n")):
            p = row[i].paragraphs[0] if j == 0 else row[i].add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(line)
            _apply_kr_font(r, bold=(bold_first and i == 0), size_pt=10)


def _add_kv_table(doc, rows: list[tuple[str, str]], *,
                  label_width_cm: float = 4.5,
                  value_width_cm: float = 11.5,
                  label_fill: str = "F2F2F2") -> None:
    tbl = doc.add_table(rows=0, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Cm(label_width_cm)
    tbl.columns[1].width = Cm(value_width_cm)
    for label, value in rows:
        row = tbl.add_row().cells
        row[0].text = ""
        row[1].text = ""
        row[0].width = Cm(label_width_cm)
        row[1].width = Cm(value_width_cm)
        for cell in row:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _shade_cell(row[0], label_fill)
        p0 = row[0].paragraphs[0]
        r0 = p0.add_run(label)
        _apply_kr_font(r0, bold=True, size_pt=10)
        for i, line in enumerate(value.split("\n")):
            p = row[1].paragraphs[0] if i == 0 else row[1].add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(line)
            _apply_kr_font(r, size_pt=10)


def _add_bullets(doc, items: list[str], *, size_pt: float = 11) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        _apply_kr_font(r, size_pt=size_pt)


# ──────────────────────────────────────────────────────────────────────────────
# 문서 빌드
# ──────────────────────────────────────────────────────────────────────────────

def build_document(output_path: Path) -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = KOREAN_FONT
    style.font.size = Pt(11)

    for section in doc.sections:
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # ─── 표지 ────────────────────────────────────────────────────────────────
    _add_paragraph(doc, DOC_TITLE, bold=True, size_pt=22,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, DOC_SUBTITLE, size_pt=12,
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   color=RGBColor(0x55, 0x55, 0x55))
    _add_paragraph(doc, DOC_DATE, size_pt=11,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=24)

    # ─── 0. 이 문서를 읽는 분께 ──────────────────────────────────────────────
    _add_heading(doc, "이 문서를 읽는 분께", level=2)
    _add_paragraph(
        doc,
        "본 문서는 Denvia 프로젝트 코드(인계 시점의 git 저장소 전체)를 압축본으로 "
        "받으신 분이, 자신의 PC에서 코드를 띄우고 직접 수정·테스트하실 수 있게 하기 "
        "위한 셋업 안내서입니다. 어려운 기술 용어는 가능한 한 풀어서 설명했고, 명령어·"
        "환경변수마다 옆 괄호에 한국어 설명을 붙였습니다.",
        size_pt=11,
    )
    _add_paragraph(
        doc,
        "본 문서는 다음 두 가지를 다룹니다.",
        size_pt=11,
    )
    _add_bullets(doc, [
        "내 PC에 어떤 프로그램을 설치해야 코드가 돌아가는지",
        "압축본을 풀고 첫 실행 명령을 어떻게 내려서 브라우저로 확인할 수 있는지",
    ])
    _add_paragraph(
        doc,
        "다음 항목은 본 문서가 다루지 않습니다(다른 안내서 참고).",
        size_pt=11,
    )
    _add_bullets(doc, [
        "운영 서버(AWS Lightsail) 직접 배포 — 「Denvia_서버_안내.docx」 참조",
        "운영 환경 비밀번호·API 키 발급 절차 — 「클라이언트_준비사항_안내서.docx」 참조",
        "코드 수정 후 git push 및 운영 반영 절차 — 운영팀과 별도 협의 필요",
    ])

    # ─── 1. PC 권장 사양 ────────────────────────────────────────────────────
    _add_heading(doc, "1. 컴퓨터(PC) 권장 사양", level=1)
    _add_paragraph(
        doc,
        "이 시스템은 한 대의 PC 안에서 6개의 프로그램(웹 서버, API 서버, 데이터베이스, "
        "Redis(레디스 — 임시 데이터 저장소), 작업 큐, 작업 스케줄러)을 동시에 띄웁니다. "
        "사양이 낮아도 실행 자체는 되지만, 응답이 느려지거나 첫 부팅이 매우 오래 걸릴 수 "
        "있습니다.",
        size_pt=11,
    )
    tbl = doc.add_table(rows=1, cols=3)
    tbl.autofit = False
    tbl.columns[0].width = Cm(4.5)
    tbl.columns[1].width = Cm(5.5)
    tbl.columns[2].width = Cm(6.0)
    _add_table_header(tbl, ["항목", "최소", "권장"])
    _add_table_row(tbl, ["운영체제",
                         "Windows 10 64bit / macOS 12 / Ubuntu 22.04",
                         "Windows 11 / macOS 14"])
    _add_table_row(tbl, ["CPU(중앙처리장치)", "4 코어", "8 코어 이상"])
    _add_table_row(tbl, ["메모리(RAM)", "8 GB", "16 GB 이상 (Docker가 4~6 GB 사용)"])
    _add_table_row(tbl, ["저장공간", "20 GB 여유",
                         "40 GB 여유 (코드 + Docker 이미지 + DB 데이터)"])
    _add_table_row(tbl, ["인터넷", "안정적 연결",
                         "첫 실행 시 약 2 GB 다운로드"])
    doc.add_paragraph()
    _add_callout_box(
        doc,
        "※ Windows 사용자 주의 — Windows Home 에디션도 동작하지만, Pro 에디션 + "
        "WSL2(윈도우용 리눅스 하위 시스템) 조합이 가장 안정적입니다. Home에서는 "
        "Docker Desktop이 설치 시 WSL2를 자동으로 함께 설치합니다.",
    )

    # ─── 2. 압축 푸는 위치 ──────────────────────────────────────────────────
    _add_heading(doc, "2. 압축 파일을 어디에 풀까", level=1)
    _add_paragraph(
        doc,
        "받은 압축 파일(예: Dental-Chatbot.zip)을 다음 규칙대로 풀어주세요.",
        size_pt=11,
    )
    _add_heading(doc, "2.1 좋은 위치 (권장)", level=2)
    _add_kv_table(doc, [
        ("Windows", "D:\\projects\\Dental Chatbot\\ 또는 C:\\work\\Dental Chatbot\\"),
        ("macOS", "~/projects/dental-chatbot/"),
        ("Linux", "~/projects/dental-chatbot/"),
    ], label_width_cm=3.0, value_width_cm=13.0)

    _add_heading(doc, "2.2 피해야 할 위치 (실행이 깨질 수 있음)", level=2)
    _add_bullets(doc, [
        "한글이 들어간 경로 (예: C:\\내 문서\\치과챗봇\\) — 일부 도구가 한글 경로를 못 읽음",
        "공백이 너무 많은 경로 (예: C:\\Program Files\\... 아래)",
        "OneDrive·Dropbox·Google Drive 동기화 폴더 — 동기화 충돌로 빌드가 깨짐",
        "네트워크 드라이브(\\\\서버\\공유\\... 형태) — Docker 마운트가 동작하지 않음",
    ])

    _add_heading(doc, "2.3 압축을 풀면 보여야 하는 폴더 구조", level=2)
    _add_paragraph(
        doc,
        "압축을 풀고 들어가 보면 아래 폴더들이 보여야 정상입니다.",
        size_pt=11,
    )
    _add_code_box(doc, (
        "Dental Chatbot/\n"
        "├─ api/              ← 백엔드 코드 (Python / FastAPI)\n"
        "├─ web/              ← 프론트엔드 코드 (Next.js)\n"
        "├─ vendor/           ← 외부 자산 (디자인 시스템, RAG 코드)\n"
        "├─ infra/            ← Docker 설정 파일\n"
        "├─ docs/             ← 개발자용 기술 문서\n"
        "├─ scripts/          ← 보조 스크립트\n"
        "├─ 자료/             ← 고객 전달용 문서(본 안내서 포함)\n"
        "├─ .env.example      ← 환경변수 견본 (복사해서 .env 를 만듭니다)\n"
        "└─ CLAUDE.md         ← 프로젝트 운영 규칙"
    ))
    _add_paragraph(
        doc,
        "node_modules/, outputs/, __pycache__/ 같은 폴더는 압축본에 없어도 정상이며 "
        "첫 실행 시 자동으로 생성됩니다.",
        size_pt=10.5, color=RGBColor(0x55, 0x55, 0x55),
    )

    # ─── 3. 설치 프로그램 ───────────────────────────────────────────────────
    _add_heading(doc, "3. 설치해야 하는 프로그램", level=1)
    _add_paragraph(
        doc,
        "필수로 깔아야 하는 건 사실상 1개(Docker Desktop)뿐입니다. 나머지는 편의용·"
        "선택입니다.",
        size_pt=11,
    )

    _add_heading(doc, "3.1 Docker Desktop — 필수", level=2)
    _add_paragraph(
        doc,
        "이 프로젝트의 모든 서버(웹, API, DB 등)는 Docker 컨테이너(작은 가상 박스) 안에서 "
        "돌아갑니다. 이 박스를 띄워주는 프로그램이 Docker Desktop(도커 데스크톱 — "
        "컨테이너 관리 GUI)입니다.",
        size_pt=11,
    )
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Cm(3.0)
    tbl.columns[1].width = Cm(13.0)
    _add_table_header(tbl, ["OS", "다운로드 위치"])
    _add_table_row(tbl, ["Windows",
                         "https://www.docker.com/products/docker-desktop/  → Download for Windows"])
    _add_table_row(tbl, ["macOS",
                         "같은 페이지 → Download for Mac (Apple Silicon / Intel 구분)"])
    _add_table_row(tbl, ["Linux",
                         "https://docs.docker.com/engine/install/ubuntu/"])
    doc.add_paragraph()
    _add_paragraph(
        doc,
        "설치 후 한 번 실행해서 시스템 트레이(윈도우 우측 하단) 또는 메뉴바(macOS 상단)에 "
        "고래 아이콘이 나타나는지 확인하세요. 그다음 명령어 창에서 아래를 실행해 정상 설치를 "
        "확인합니다.",
        size_pt=11,
    )
    _add_code_box(doc, (
        "docker --version\n"
        "docker compose version"
    ))
    _add_paragraph(
        doc,
        "두 명령 모두 버전 번호가 출력되면 정상입니다.",
        size_pt=11,
    )
    _add_callout_box(
        doc,
        "※ Windows 사용자가 처음 Docker Desktop을 깔면 \"WSL2 업데이트 필요\" 메시지가 "
        "뜰 수 있습니다. 안내에 따라 마이크로소프트 사이트에서 업데이트 파일을 받아 "
        "설치한 뒤 PC를 재부팅하세요.",
    )

    _add_heading(doc, "3.2 코드 편집기 — 강력 권장", level=2)
    _add_paragraph(
        doc,
        "코드를 보고 수정하려면 편집기가 필요합니다. Visual Studio Code(VS Code, 브이에스코드 "
        "— 무료 코드 편집기)를 권장합니다.",
        size_pt=11,
    )
    _add_bullets(doc, [
        "다운로드: https://code.visualstudio.com/",
        "설치 후 압축 푼 폴더(예: D:\\projects\\Dental Chatbot)를 통째로 엽니다 — File → Open Folder...",
        "기본 사용법은 「Denvia_VSCode_SSH_접속_안내서.pdf」도 참고 가능 (운영 서버 접속용이지만 사용법은 동일)",
    ])

    _add_heading(doc, "3.3 Git — 권장", level=2)
    _add_paragraph(
        doc,
        "압축이 아니라 git 저장소에서 직접 코드를 받고 싶을 때, 또는 코드 수정 이력을 "
        "관리하고 싶을 때 필요합니다. 압축본만 쓰실 거면 굳이 설치하지 않아도 됩니다.",
        size_pt=11,
    )
    _add_bullets(doc, [
        "Windows: https://git-scm.com/download/win",
        "macOS / Linux: 보통 기본 설치되어 있음 (git --version 으로 확인)",
    ])

    _add_heading(doc, "3.4 (선택) Node.js · Python 등 — 안 깔아도 됨", level=2)
    _add_paragraph(
        doc,
        "Docker 안에서 모든 게 돌아가므로 호스트 PC에는 Node·Python을 굳이 깔지 않아도 됩니다. "
        "다만 Docker 없이 편집기에서 타입 검사·자동완성·테스트를 빠르게 돌리고 싶다면 다음을 "
        "추가로 설치하면 편합니다.",
        size_pt=11,
    )
    tbl = doc.add_table(rows=1, cols=3)
    tbl.autofit = False
    tbl.columns[0].width = Cm(4.0)
    tbl.columns[1].width = Cm(7.5)
    tbl.columns[2].width = Cm(4.5)
    _add_table_header(tbl, ["도구", "용도", "버전"])
    _add_table_row(tbl, ["Node.js(노드 — 자바스크립트 실행기)",
                         "프론트(web) 직접 실행", "20 LTS"])
    _add_table_row(tbl, ["pnpm(피엔피엠 — 노드 패키지 관리자)",
                         "프론트 의존성 설치", "9.x"])
    _add_table_row(tbl, ["Python(파이썬)", "백엔드(api) 직접 실행", "3.12"])
    _add_table_row(tbl, ["uv(유브이 — 파이썬 패키지 관리자)",
                         "백엔드 의존성 설치", "최신"])
    doc.add_paragraph()
    _add_paragraph(
        doc,
        "본 안내서는 Docker만 쓰는 경로를 기본으로 설명합니다. 직접 실행 경로는 "
        "docs/DEVELOPMENT.md 를 참고하세요.",
        size_pt=10.5, color=RGBColor(0x55, 0x55, 0x55),
    )

    # ─── 4. .env 만들기 ─────────────────────────────────────────────────────
    _add_heading(doc, "4. 환경설정 파일(.env) 만들기", level=1)
    _add_paragraph(
        doc,
        "이 시스템은 비밀번호·API 키 같은 비밀값을 코드에 직접 적지 않고 .env(닷이엔브이 — "
        "환경변수 파일, 점으로 시작해서 숨김 파일 취급) 라는 한 파일에 모아둡니다.",
        size_pt=11,
    )

    _add_heading(doc, "4.1 견본 파일 복사", level=2)
    _add_paragraph(
        doc,
        "압축 푼 폴더 최상위에 .env.example(닷이엔브이 닷 예제 — 견본 파일) 이 있습니다. "
        "이걸 복사해서 이름을 .env(점 + env) 로 바꿉니다.",
        size_pt=11,
    )
    _add_paragraph(doc, "Windows PowerShell:", bold=True, size_pt=11)
    _add_code_box(doc, (
        "cd \"D:\\projects\\Dental Chatbot\"\n"
        "Copy-Item .env.example .env"
    ))
    _add_paragraph(doc, "macOS / Linux:", bold=True, size_pt=11)
    _add_code_box(doc, (
        "cd ~/projects/dental-chatbot\n"
        "cp .env.example .env"
    ))

    _add_heading(doc, "4.2 채워야 하는 값 / 비워둬도 되는 값", level=2)
    _add_paragraph(
        doc,
        ".env 를 편집기로 열면 항목이 많아 보이지만, 로컬 테스트만 할 거라면 채워야 하는 건 "
        "사실상 1~2개입니다. 나머지는 비어 있어도 서버는 뜹니다(다만 해당 기능만 동작하지 않음).",
        size_pt=11,
    )
    tbl = doc.add_table(rows=1, cols=3)
    tbl.autofit = False
    tbl.columns[0].width = Cm(5.5)
    tbl.columns[1].width = Cm(2.5)
    tbl.columns[2].width = Cm(8.0)
    _add_table_header(tbl, ["항목", "채워야?", "설명"])
    rows = [
        ("DATABASE_URL\n(DB 접속 주소)", "그대로", "Docker 안 Postgres에 맞춰 이미 채워져 있음"),
        ("REDIS_URL\n(레디스 접속 주소)", "그대로", "Docker 안 Redis에 맞춰 이미 채워져 있음"),
        ("DENVIA_JWT_SECRET\n(로그인 토큰 서명용)", "권장", "견본값으로도 동작. 안전을 위해 32자 임의 문자열로 교체"),
        ("DENVIA_ADMIN_EMAIL\n(관리자 이메일)", "그대로", "기본값 admin@denvia.ai.kr 사용"),
        ("DENVIA_ADMIN_INITIAL_PASSWORD\n(관리자 초기 비번)", "그대로", "기본값 change_me_in_production 사용"),
        ("OPENAI_API_KEY\n(오픈AI 키)", "필요",
         "RAG 챗봇 테스트하려면 발급 필수. 비어 있으면 챗봇 답변 안 옴"),
        ("KAKAO_* / GOOGLE_* / NAVER_*\n(소셜 로그인 키)", "비워둠",
         "이메일 로그인은 가능. 소셜 로그인 버튼만 동작 안 함"),
        ("TOSS_SECRET_KEY / NEXT_PUBLIC_TOSS_CLIENT_KEY\n(토스 결제 키)", "그대로",
         "견본에 토스 공식 샌드박스 키가 들어 있음(실 결제 X)"),
        ("MESSAGING_PROVIDER\n(메시징 공급자)", "stub 유지",
         "stub(스텁 — 가짜 발송, 로그만 남김)이면 실제 카톡/SMS 안 나감"),
        ("ALIGO_*\n(알리고 알림톡 키)", "비워둠",
         "MESSAGING_PROVIDER=stub 이면 어차피 호출 안 됨"),
        ("KOREAEXIM_API_KEY\n(한국수출입은행 환율)", "비워둠",
         "비면 환율 갱신만 안 함, 기본 1,400원 사용"),
        ("BILLING_KEY_ENC_KEY\n(결제 빌링키 암호화 키)", "결제 모듈 띄우려면 필요",
         "비면 결제 API가 시작 시 에러. 다음 절(4.3) 명령으로 생성"),
    ]
    for label, must, desc in rows:
        _add_table_row(tbl, [label, must, desc])
    doc.add_paragraph()
    _add_callout_box(
        doc,
        "※ OpenAI 키 발급 — https://platform.openai.com/api-keys 에서 계정을 만들고 키를 "
        "발급(sk- 로 시작)합니다. 신규 계정은 약간의 무료 크레딧이 있고, 이후 사용량만큼 "
        "결제카드에서 차감됩니다. 본 안내서 수준의 테스트라면 1~2달러 미만 비용입니다.",
    )

    _add_heading(doc, "4.3 BILLING_KEY_ENC_KEY 생성하기", level=2)
    _add_paragraph(
        doc,
        "결제 모듈을 띄우려면 이 키를 채워야 합니다. 다음 중 하나의 명령을 실행하면 한 줄 "
        "문자열이 출력됩니다. 그걸 복사해 .env 의 BILLING_KEY_ENC_KEY= 뒤에 붙여넣으세요.",
        size_pt=11,
    )
    _add_paragraph(doc, "PC에 Python이 설치되어 있을 때:", bold=True, size_pt=11)
    _add_code_box(doc, (
        "python -c \"from cryptography.fernet import Fernet; print(Fernet.generate_key().decode())\""
    ))
    _add_paragraph(doc, "Docker만 있을 때 (먼저 Docker가 떠 있어야 함):", bold=True, size_pt=11)
    _add_code_box(doc, (
        "docker compose -f infra/docker-compose.yml run --rm api \\\n"
        "  uv --project /workspace/api run python -c \\\n"
        "  \"from cryptography.fernet import Fernet; print(Fernet.generate_key().decode())\""
    ))

    # ─── 5. 처음 실행하기 ──────────────────────────────────────────────────
    _add_heading(doc, "5. 처음 실행하기", level=1)
    _add_paragraph(
        doc,
        "여기서부터는 명령어 한 줄로 끝납니다.",
        size_pt=11,
    )

    _add_heading(doc, "5.1 명령어 창 열기", level=2)
    _add_bullets(doc, [
        "Windows: 시작 메뉴 → \"PowerShell\" 검색 → 실행",
        "macOS: Launchpad → \"Terminal\" 실행",
        "Linux: 기본 터미널 실행",
    ])

    _add_heading(doc, "5.2 압축 푼 폴더로 이동", level=2)
    _add_code_box(doc, "cd \"D:\\projects\\Dental Chatbot\"")
    _add_paragraph(
        doc,
        "(macOS·Linux는 본인이 푼 경로에 맞게 바꿔서 입력)",
        size_pt=10.5, color=RGBColor(0x55, 0x55, 0x55),
    )

    _add_heading(doc, "5.3 첫 실행 명령", level=2)
    _add_code_box(doc,
                  "docker compose -f infra/docker-compose.yml up --build")
    _add_paragraph(doc, "이 명령이 하는 일:", bold=True, size_pt=11)
    _add_bullets(doc, [
        "infra/docker-compose.yml(컴포즈 설정 파일)을 읽음",
        "필요한 Docker 이미지를 인터넷에서 받아옴 (Postgres·Redis·Node·Python 등)",
        "각 컨테이너 안에 코드 의존성을 자동 설치 (pnpm install, uv sync)",
        "DB 스키마(테이블)를 자동 생성 (alembic upgrade head)",
        "관리자 계정을 자동 생성 (seed_admin.py)",
        "7개 서비스(web, api, worker, beat, postgres, redis, nginx)를 모두 기동",
    ])
    _add_callout_box(
        doc,
        "※ 첫 실행 소요 시간 — 인터넷 속도에 따라 10~20분. 처음에만 오래 걸리고, "
        "두 번째 실행부터는 1분 내 부팅됩니다.",
    )

    _add_heading(doc, "5.4 정상 실행 확인", level=2)
    _add_paragraph(
        doc,
        "명령어 창에 새 로그가 멈추고 다음과 비슷한 메시지가 보이면 성공입니다.",
        size_pt=11,
    )
    _add_code_box(doc, (
        "api-1     | INFO:     Uvicorn running on http://0.0.0.0:8000\n"
        "web-1     | ✓ Ready in 3.4s\n"
        "web-1     | - Local:   http://localhost:3000"
    ))
    _add_paragraph(
        doc,
        "이제 브라우저를 열고 다음 두 주소를 확인하세요.",
        size_pt=11,
    )
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Cm(6.0)
    tbl.columns[1].width = Cm(10.0)
    _add_table_header(tbl, ["주소", "보여야 하는 화면"])
    _add_table_row(tbl, ["http://localhost:3000", "Denvia 메인 페이지"])
    _add_table_row(tbl, ["http://localhost:8000/healthz",
                         "{\"status\":\"ok\"} 같은 JSON 한 줄"])
    doc.add_paragraph()

    _add_heading(doc, "5.5 관리자 로그인 테스트", level=2)
    _add_paragraph(
        doc,
        "http://localhost:3000/login 에 접속해서 다음 정보로 로그인합니다.",
        size_pt=11,
    )
    _add_kv_table(doc, [
        ("이메일", "admin@denvia.ai.kr"),
        ("비밀번호",
         "change_me_in_production  (.env 의 DENVIA_ADMIN_INITIAL_PASSWORD 값)"),
    ], label_width_cm=3.5, value_width_cm=12.5)
    _add_paragraph(
        doc,
        "로그인 후 우측 상단 메뉴에서 /admin (관리자 페이지)으로 진입되면 모든 설정이 정상입니다.",
        size_pt=11,
    )

    # ─── 6. 자주 쓰는 명령 ─────────────────────────────────────────────────
    _add_heading(doc, "6. 자주 쓰는 명령어 모음", level=1)
    _add_paragraph(
        doc,
        "압축 푼 폴더 최상위에서 실행합니다.",
        size_pt=11,
    )
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Cm(5.0)
    tbl.columns[1].width = Cm(11.0)
    _add_table_header(tbl, ["목적", "명령"])
    cmds = [
        ("띄우기 (로그도 같이 봄)", "docker compose -f infra/docker-compose.yml up"),
        ("띄우기 (백그라운드)", "docker compose -f infra/docker-compose.yml up -d"),
        ("로그만 보기 (web + api)",
         "docker compose -f infra/docker-compose.yml logs -f web api"),
        ("멈추기", "docker compose -f infra/docker-compose.yml down"),
        ("데이터까지 완전 초기화", "docker compose -f infra/docker-compose.yml down -v"),
        ("컨테이너 상태 확인", "docker compose -f infra/docker-compose.yml ps"),
        ("API 컨테이너에 직접 들어가기",
         "docker compose -f infra/docker-compose.yml exec api bash"),
    ]
    for purpose, cmd in cmds:
        _add_table_row(tbl, [purpose, cmd])
    doc.add_paragraph()
    _add_callout_box(
        doc,
        "※ down -v 주의 — -v 옵션은 Postgres DB 내용까지 함께 지웁니다. 회원·결제 데이터를 "
        "초기화하고 싶을 때만 쓰세요. 평소엔 down 만 씁니다.",
        fill="FFE8E8",
    )

    # ─── 7. 코드 수정 반영 ─────────────────────────────────────────────────
    _add_heading(doc, "7. 코드를 수정하면 어떻게 반영되나", level=1)
    _add_paragraph(
        doc,
        "Docker가 떠 있는 상태에서 편집기로 코드 파일을 저장하면 다시 띄울 필요 없이 "
        "자동 반영됩니다(이를 \"Hot Reload\"(핫 리로드 — 자동 새로고침) 라고 부릅니다).",
        size_pt=11,
    )
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Cm(6.0)
    tbl.columns[1].width = Cm(10.0)
    _add_table_header(tbl, ["수정 위치", "자동 반영 방식"])
    rows = [
        ("web/ 안의 .tsx / .css 파일", "브라우저 새로고침 없이 화면 갱신"),
        ("api/src/ 안의 .py 파일", "API 서버가 1~2초 안에 자동 재시작"),
        ("api/alembic/versions/ 안의 새 마이그레이션",
         "다음 시작 시 자동 적용. 즉시 적용은 exec 명령으로 alembic upgrade head"),
        (".env 파일 변경",
         "자동 반영 안 됨. down 후 다시 up 해야 함"),
        ("pyproject.toml / package.json 의존성 추가",
         "컨테이너 재시작 필요"),
        ("Celery worker / beat (작업 큐) 코드",
         "재시작 필요: docker compose ... restart worker beat"),
    ]
    for left, right in rows:
        _add_table_row(tbl, [left, right])
    doc.add_paragraph()
    _add_paragraph(
        doc,
        "호스트 PC의 폴더(D:\\projects\\Dental Chatbot\\web\\src\\...)와 Docker 안의 폴더"
        "(/workspace/web/src/...)는 같은 파일을 공유합니다. 호스트에서 저장한 변경이 그대로 "
        "컨테이너 안에 보입니다.",
        size_pt=11,
    )

    # ─── 8. 트러블슈팅 ────────────────────────────────────────────────────
    _add_heading(doc, "8. 자주 만나는 문제와 해결", level=1)

    _add_heading(doc, "8.1 \"port is already allocated\" — 포트 충돌", level=2)
    _add_paragraph(
        doc,
        "다른 프로그램이 이미 3000/8000/5432/6379/80 번 포트를 쓰고 있어서 못 띄움. "
        "대표적으로 다른 Postgres·Redis가 호스트에 깔려 있을 때 발생.",
        size_pt=11,
    )
    _add_code_box(doc, (
        "docker ps   # 어떤 컨테이너가 무슨 포트를 쓰는지 확인\n"
        "# 호스트의 Postgres·Redis 서비스를 잠시 끄거나\n"
        "# infra/docker-compose.yml 에서 ports 매핑을 바꾸기 (예: 5432:5432 → 5433:5432)"
    ))

    _add_heading(doc, "8.2 첫 실행이 너무 오래 걸린다", level=2)
    _add_bullets(doc, [
        "첫 실행은 정상적으로 10~20분 걸립니다. 인터넷 속도와 PC 성능에 따라 더 오래 걸릴 수도 있습니다.",
        "pnpm install(프론트 의존성 설치) 단계에서 멈춘 듯 보여도 실제로는 다운로드 중입니다. 작업관리자/Activity Monitor에서 네트워크 사용량을 확인해보세요.",
        "30분이 지나도 진행이 없으면 Ctrl+C로 중단하고 다시 up 명령을 실행하세요.",
    ])

    _add_heading(doc, "8.3 관리자 로그인이 안 됨 (401 에러)", level=2)
    _add_bullets(doc, [
        ".env 의 DENVIA_ADMIN_INITIAL_PASSWORD 값을 정확히 입력했는지 확인 (.env 에 따옴표 없이 적혀 있는 그대로)",
        "다른 비밀번호를 시도해본 적이 있다면 잠금 상태일 수 있음. 5분 기다린 뒤 다시 시도",
        "DB가 초기화되어 관리자 시드가 안 됐을 수도 있음. 아래 명령으로 다시 시드:",
    ])
    _add_code_box(doc, (
        "docker compose -f infra/docker-compose.yml exec api \\\n"
        "  uv --project /workspace/api run python /workspace/api/scripts/seed_admin.py"
    ))

    _add_heading(doc, "8.4 화면은 뜨는데 챗봇 답변이 안 옴", level=2)
    _add_paragraph(
        doc,
        ".env 의 OPENAI_API_KEY 값이 비어 있거나 잘못된 경우입니다. 키를 채운 뒤 반드시 "
        "컨테이너를 재시작하세요.",
        size_pt=11,
    )
    _add_code_box(doc, (
        "docker compose -f infra/docker-compose.yml down\n"
        "docker compose -f infra/docker-compose.yml up -d"
    ))

    _add_heading(doc, "8.5 알림톡(카카오톡)이 안 옴", level=2)
    _add_bullets(doc, [
        "로컬 환경에서는 알림톡·SMS가 실제 발송되지 않습니다 (의도된 동작, MESSAGING_PROVIDER=stub).",
        "발송 시도 기록만 API 로그에서 확인 가능:",
    ])
    _add_code_box(doc,
                  "docker compose -f infra/docker-compose.yml logs api | grep \"alimtalk\"")
    _add_paragraph(
        doc,
        "실제 발송 테스트는 운영 서버(「Denvia_서버_안내.docx」 참조)에서만 가능합니다.",
        size_pt=11,
    )

    _add_heading(doc, "8.6 디스크 용량 부족 경고", level=2)
    _add_paragraph(
        doc,
        "Docker 이미지·볼륨이 점점 쌓입니다. 안 쓰는 컨테이너를 정리하세요. "
        "현재 떠 있는 것은 지워지지 않습니다.",
        size_pt=11,
    )
    _add_code_box(doc, "docker system prune -a   # y 입력하면 안 쓰는 이미지 모두 삭제")

    _add_heading(doc, "8.7 윈도우에서 \"Filesharing is not enabled\" 에러", level=2)
    _add_paragraph(
        doc,
        "Docker Desktop 설정 → Resources(리소스) → File sharing(파일 공유) 탭에서 D:\\ 드라이브가 "
        "허용 목록에 있는지 확인하세요. 없으면 추가하고 Apply & Restart 클릭.",
        size_pt=11,
    )

    _add_heading(doc, "8.8 그 외 — 일단 한 번 다시 띄워본다", level=2)
    _add_paragraph(
        doc,
        "대부분의 일시적 문제는 다음 절차로 해결됩니다.",
        size_pt=11,
    )
    _add_code_box(doc, (
        "docker compose -f infra/docker-compose.yml down\n"
        "docker compose -f infra/docker-compose.yml up -d\n"
        "docker compose -f infra/docker-compose.yml logs -f web api"
    ))
    _add_paragraph(
        doc,
        "그래도 안 되면 로그 전체를 운영팀에 전달해 주세요.",
        size_pt=11,
    )

    # ─── 9. 종료 ──────────────────────────────────────────────────────────
    _add_heading(doc, "9. 종료할 때", level=1)
    _add_paragraph(
        doc,
        "작업이 끝나면 컨테이너를 멈춰서 PC 자원을 돌려받으세요.",
        size_pt=11,
    )
    _add_code_box(doc, "docker compose -f infra/docker-compose.yml down")
    _add_paragraph(
        doc,
        "PC를 끄거나 Docker Desktop을 종료해도 자동으로 멈춥니다. 다음에 다시 작업할 때는 "
        "up -d 명령만 다시 실행하면 됩니다. DB 데이터(회원·결제·게시판 등)는 PC 안의 Docker "
        "볼륨에 그대로 남아 있어서, 다음 실행 시 같은 상태에서 이어집니다.",
        size_pt=11,
    )

    # ─── 10. 큰 그림 ─────────────────────────────────────────────────────
    _add_heading(doc, "10. 코드 수정 → 운영 서버 반영까지의 큰 그림", level=1)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.autofit = False
    tbl.columns[0].width = Cm(3.5)
    tbl.columns[1].width = Cm(4.5)
    tbl.columns[2].width = Cm(8.0)
    _add_table_header(tbl, ["단계", "어디서", "어떻게"])
    _add_table_row(tbl, ["1. 코드 수정", "내 PC, 편집기에서",
                         "web/, api/ 안의 파일 편집"])
    _add_table_row(tbl, ["2. 로컬에서 확인", "내 PC, 브라우저에서",
                         "http://localhost:3000 에서 동작 확인"])
    _add_table_row(tbl, ["3. git 커밋", "내 PC, 명령어 창",
                         "git add . → git commit -m \"...\" → git push"])
    _add_table_row(tbl, ["4. 운영 서버 배포", "운영 서버 SSH",
                         "「Denvia_서버_안내.docx」의 배포 절차"])
    doc.add_paragraph()
    _add_callout_box(
        doc,
        "※ 본 안내서는 1~2단계까지 다룹니다. 3~4단계는 운영팀과 협의 후 진행하세요. "
        "운영 서버 직접 배포는 실수 시 서비스 중단으로 이어지므로 신중히 다루셔야 합니다.",
        fill="FFE8E8",
    )

    # ─── 11. 더 자세한 문서 ───────────────────────────────────────────────
    _add_heading(doc, "11. 더 자세한 기술 문서", level=1)
    _add_paragraph(
        doc,
        "본 안내서는 비전공자도 따라 할 수 있게 핵심만 추렸습니다. 코드를 깊이 들여다보거나 "
        "새 기능을 개발할 때는 다음 문서를 참고하세요 (모두 docs/ 폴더 안에 있음).",
        size_pt=11,
    )
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Cm(6.0)
    tbl.columns[1].width = Cm(10.0)
    _add_table_header(tbl, ["문서", "용도"])
    rows = [
        ("docs/DEVELOPMENT.md", "개발자용 Docker·Hot Reload 상세"),
        ("docs/ARCHITECTURE_OVERVIEW.md", "시스템 전체 구조도"),
        ("docs/CODING_CONVENTIONS.md", "코드 작성 규칙"),
        ("docs/ONBOARDING.md", "인수자 첫 운영 작업 4종"),
        ("docs/OPERATIONS.md", "일상 운영·모니터링"),
        ("docs/SECURITY.md", "보안 운영"),
        ("docs/RUNBOOK_DEPLOY.md", "배포·CI/CD"),
        ("docs/RUNBOOK_INCIDENT.md", "장애 대응"),
        ("docs/RUNBOOK_RAG.md", "RAG(챗봇 지식베이스) 운영"),
        ("docs/ALIMTALK_TEMPLATES.md", "알림톡 템플릿 카탈로그"),
    ]
    for left, right in rows:
        _add_table_row(tbl, [left, right])
    doc.add_paragraph()

    # ─── 문의 ────────────────────────────────────────────────────────────
    _add_horizontal_rule(doc)
    _add_heading(doc, "문의", level=2)
    _add_paragraph(
        doc,
        "본 안내서대로 진행했는데도 막히는 부분이 있으면, ① 막힌 단계 번호, "
        "② 명령어 창에 나온 마지막 에러 메시지, ③ 스크린샷, 이 세 가지를 첨부해 "
        "운영팀에 전달해 주세요. 그 세 가지가 있으면 대부분 1시간 안에 원인을 짚어드릴 "
        "수 있습니다.",
        size_pt=11,
    )
    _add_paragraph(doc, "감사합니다.", bold=True, size_pt=12,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=12)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


# ──────────────────────────────────────────────────────────────────────────────
# PDF 변환 (docx2pdf — Windows + Word 필요)
# ──────────────────────────────────────────────────────────────────────────────

def convert_to_pdf(docx_path: Path) -> Path | None:
    pdf_path = docx_path.with_suffix(".pdf")
    try:
        from docx2pdf import convert
    except ImportError:
        print("[WARN] docx2pdf 미설치 — PDF 생성 스킵.")
        print("       설치: py -3 -m pip install docx2pdf")
        return None
    try:
        convert(str(docx_path), str(pdf_path))
    except Exception as exc:
        print(f"[WARN] PDF 변환 실패: {exc}")
        print("       Word(또는 LibreOffice)가 설치된 환경에서 다시 실행해 주세요.")
        return None
    return pdf_path


def main() -> None:
    repo_root = Path(__file__).resolve().parent.parent
    docx_path = repo_root / "자료" / "Denvia_로컬_개발환경_셋업_안내서.docx"
    build_document(docx_path)
    print(f"[OK] Wrote {docx_path}")

    pdf_path = convert_to_pdf(docx_path)
    if pdf_path is not None:
        print(f"[OK] Wrote {pdf_path}")


if __name__ == "__main__":
    main()
