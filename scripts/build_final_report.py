"""Denvia 개발 완료 보고서 (최종) 빌더.

2026-06-05 — 전체 개발 완료 시점 보고서.

산출물:
    자료/Denvia_개발완료보고서_2026-06-05.docx

CLAUDE.md "고객 전달용 문서 작성 규칙" 준수 — SDT(content control), 폼 보호,
문서 보호, ActiveX 등 일체 사용하지 않음. 어떤 워드 버전에서도 자유 편집 가능.

톤:
- 중간보고서(2026-04-29)와 동일한 비개발자 친화 한글 문체.
- Epic/Story 번호·SSOT 편차·DB 마이그레이션 등 기술 디테일은 "별첨"으로만 노출.
- 계약 후 추가/변경된 항목은 "변경·추가 사항" 별도 섹션에서 명확히 정리.

실행:
    py -3 scripts/build_final_report.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ──────────────────────────────────────────────────────────────────────────────
# 콘텐츠 정의
# ──────────────────────────────────────────────────────────────────────────────

DOC_TITLE = "Denvia 개발 완료 보고서"
DOC_SUBTITLE = "전 개발 범위 완료 · 인계 준비 단계"
DOC_DATE = "2026년 6월 5일"


# 전체 진행 현황 — 큰 영역별 요약
PROGRESS_SUMMARY: list[dict] = [
    {
        "area": "사용자(고객) 화면",
        "scope": "회원가입 / 로그인 / Q&A 채팅 / 구독 결제 / 마이페이지 / 쪽지·문의 / 안내 팝업",
        "status": "완료",
        "note": "이메일·소셜(카카오/네이버/구글) 가입, 무료/유료 분기, 자동 결제, 환불 신청까지 모두 동작.",
    },
    {
        "area": "관리자 화면",
        "scope": (
            "대시보드 / 회원 관리 / 통계 / 매출 / RAG 데이터 / 알림톡 / "
            "이상탐지 / 예산 통제 / CS 관리 / 설정 관리 / 관리자 정보 수정 / 관리자 관리"
        ),
        "status": "완료",
        "note": "관리자 4등급 권한 체계 포함. 마스터/운영자/부운영자/대기 분리.",
    },
    {
        "area": "AI 응답 엔진",
        "scope": "RAG(치과 행정·보험청구 지식 검색) + GPT 응답 + 동의어 사전",
        "status": "완료",
        "note": "고객 인수자료(RAG 코드) 그대로 사용. 무중단 재빌드 + 동의어 사전 관리 기능 추가.",
    },
    {
        "area": "결제 시스템",
        "scope": "토스페이먼츠 빌링키 / 자동 갱신 / 결제 실패 재시도 / 환불 3종 분기",
        "status": "완료",
        "note": "청약철회 7일 자동 / 관리자 수동환불(부분환불 포함) / 구독 해지 및 철회.",
    },
    {
        "area": "알림톡 메시지",
        "scope": "결제·구독·문의·이상탐지 등 사용자 13종 + 관리자 7종",
        "status": "완료",
        "note": (
            "메시지 본문 수정·검수, 카카오 비즈 검수, 알리고 콘솔 템플릿 등록, "
            "실제 발송 테스트까지 전 단계 완료. 추가 검수 없음."
        ),
    },
    {
        "area": "운영 서버",
        "scope": "AWS Lightsail 운영 서버 구축 / 도메인 연결 / SSL 인증서 / 자동 백업",
        "status": "완료",
        "note": (
            "denvia.ai.kr 운영 중. AWS Lightsail 서울 리전(ap-northeast-2), "
            "Ubuntu 22.04 LTS, 2 vCPU / 7.6 GB RAM / 155 GB SSD. "
            "7개 서비스(웹·API·DB·Redis·워커·스케줄러·nginx) 정상 동작."
        ),
    },
    {
        "area": "법적 문서",
        "scope": "이용약관 / 개인정보처리방침",
        "status": "완료",
        "note": "사이트에 게재 완료. 개인정보 보호책임자 = denvia@naver.com 확정.",
    },
]


# 사용자 화면별 결과
USER_SCREENS: list[dict] = [
    {
        "title": "(1) 랜딩 페이지 — 첫 진입 화면",
        "what": "비회원이 처음 사이트에 들어왔을 때 보이는 메인 화면입니다.",
        "result": (
            "• 서비스 소개와 ‘무료로 시작하기’ 버튼이 제공됩니다.\n"
            "• 로그인하지 않아도 채팅창에서 바로 질문을 시도할 수 있고, 시도하면 로그인 팝업이 자동으로 뜹니다.\n"
            "• 모바일·태블릿·PC 모두 자동 정렬되어 보입니다."
        ),
        "shots": ["user_01_landing.png"],
    },
    {
        "title": "(2) 로그인 / 회원가입 팝업 — 4가지 인증 수단",
        "what": "이메일과 카카오·네이버·구글 소셜 로그인 4가지를 한 팝업에서 선택할 수 있습니다.",
        "result": (
            "• 이메일 가입: 이름·생년월일·성별·이메일·휴대폰번호·비밀번호를 직접 입력합니다. "
            "휴대폰 본인인증(문자)과 가입유형(치과의사/치과위생사/기타) 선택까지 한 흐름으로 진행됩니다.\n"
            "• 카카오 가입: 카카오에서 받을 수 있는 기본 정보(이름·이메일 등)에 더해, "
            "사용자가 ‘선택 동의’ 한 경우 생년월일·성별·휴대폰 등 기타 정보도 자동으로 수집됩니다.\n"
            "• 네이버 가입: 네이버에서 받을 수 있는 기본 정보에 더해, "
            "‘선택 동의’ 항목을 통해 생년월일·성별·휴대폰 등 기타 정보도 자동으로 수집됩니다.\n"
            "• 구글 가입: 구글은 정책상 이메일과 비밀번호(인증) 외 다른 개인정보를 제공하지 않습니다. "
            "기타 정보가 필요한 경우 마이페이지에서 직접 입력해야 합니다.\n"
            "• 같은 이메일로 소셜 가입 후 이메일 로그인을 시도하면 ‘XX로 가입되어 있습니다’ 안내가 뜹니다."
        ),
        "shots": ["user_02_login_popup.png"],
    },
    {
        "title": "(3) 이메일/비밀번호 찾기",
        "what": "비밀번호를 잊거나 가입한 이메일이 기억나지 않을 때 사용하는 화면입니다.",
        "result": (
            "• 휴대폰 본인인증으로 비밀번호 재설정 가능.\n"
            "• 이메일(아이디) 찾기도 휴대폰 본인인증으로 진행.\n"
            "• 무차별 시도 차단(연속 실패 시 일정 시간 잠금)도 적용되어 있습니다."
        ),
        "shots": ["user_03_password_reset.png"],
    },
    {
        "title": "(4) Q&A 채팅 — 핵심 가치 화면",
        "what": "사용자가 치과 보험청구·데스크 행정 질문을 입력하면 AI가 답변하는 메인 기능입니다.",
        "result": (
            "• 무료 사용자: 하루 정해진 횟수까지 사용. 답변은 의도적으로 잠시 지연 표시.\n"
            "• 유료(Pro) 사용자: 무제한 + 답변이 한 글자씩 실시간 흘러나오는 형태로 즉시 표시.\n"
            "• 답변 아래 ‘도움됨/도움 안 됨’ 피드백 버튼.\n"
            "• 대화 초기화·질문 다시 묻기·후속 질문 자동 인식 기능 포함."
        ),
        "shots": ["user_04_qa_chat_guest.png"],
    },
    {
        "title": "(5) 구독 결제 (Pro 가입)",
        "what": "Pro 구독을 시작하기 위해 카드 정보를 등록하고 결제하는 화면입니다.",
        "result": (
            "• 토스페이먼츠 카드 등록 → 즉시 결제 → Pro 권한 활성화.\n"
            "• 매달 같은 날 자동 결제가 진행됩니다.\n"
            "• 결제가 실패하면 1일·3일 간격으로 자동 재시도하며, 알림톡이 발송됩니다."
        ),
        "shots": ["user_07_subscribe.png"],
    },
    {
        "title": "(6) 마이페이지",
        "what": "사용자가 본인의 구독 상태와 결제 내역을 확인하고 관리하는 영역입니다.",
        "result": (
            "• 현재 구독 상태(요금제·다음 결제일) 확인.\n"
            "• 결제 내역 표 (최근 결제·금액·카드 마지막 4자리).\n"
            "• 구독 해지(다음 결제 중단) / 해지 철회.\n"
            "• 청약철회 7일 환불 신청 (조건 충족 시 즉시 환불 처리).\n"
            "• 회원 탈퇴 기능 (휴대폰 본인인증 후 즉시 처리)."
        ),
        "shots": ["user_06_mypage.png"],
    },
    {
        "title": "(7) 받은 쪽지함 / 1:1 문의 게시판",
        "what": "운영 공지·관리자 답변을 받는 곳과, 사용자가 문의글을 작성·조회하는 게시판입니다.",
        "result": (
            "• 받은 쪽지함: 공지사항·환불·이상탐지 등 안내 메시지가 도착.\n"
            "• 1:1 문의: 게시판 형식. 6가지 문의 유형(결제·계정·기능·오류·환불·기타)에서 선택, 이미지 최대 5장 첨부 가능.\n"
            "• 관리자가 답변하면 알림톡으로 통지됩니다."
        ),
        "shots": ["user_08_inbox.png", "user_09_inquiries.png"],
    },
    {
        "title": "(8) 자동 안내 팝업 (광고·공지)",
        "what": "사이트 진입 시 운영자가 띄우는 공지 팝업입니다.",
        "result": (
            "• ‘오늘 하루 보지 않기’ 옵션 제공.\n"
            "• 야간(21시~익일 8시)에는 광고성 알림이 자동으로 보류됩니다(법적 요구사항 준수)."
        ),
        "shots": ["user_05_qa_chat_authed.png"],
    },
]


# 관리자 화면별 결과
ADMIN_SCREENS: list[dict] = [
    {
        "title": "(1) 관리자 로그인 / 가입 / 권한 체계",
        "what": "운영자 전용 로그인 화면과 권한 체계입니다.",
        "result": (
            "• 일반 사용자와 완전히 분리된 별도 로그인 경로(/admin/login).\n"
            "• 관리자 가입 시 ‘대기’ 상태로 들어오며, 마스터가 등급을 부여해야 활성화됩니다.\n"
            "• 4등급 권한: 마스터(개발자 1인) / 운영자 / 부운영자 / 대기.\n"
            "• 운영자끼리 서로 등급을 조정·차단·삭제할 수 있습니다."
        ),
        "shots": ["admin_01_login.png", "admin_16_admins.png"],
    },
    {
        "title": "(2) 관리자 대시보드 (메인 화면)",
        "what": "운영자가 로그인하면 가장 먼저 보는 종합 현황 화면입니다.",
        "result": (
            "• 가입자 수·구독자 수·매출·API 비용 등 위젯 한눈에.\n"
            "• 월 예산 게이지(80%/95%/100% 단계별 색상).\n"
            "• 새로운 이상탐지·문의·환불 요청이 있으면 상단에 표시."
        ),
        "shots": ["admin_02_dashboard.png"],
    },
    {
        "title": "(3) 회원 통합 검색 / 회원 상세",
        "what": "이름·이메일·휴대폰으로 회원을 검색하고 상세 정보·이용 기록을 볼 수 있는 화면입니다.",
        "result": (
            "• 통합 검색(이름·이메일·휴대폰 부분 일치).\n"
            "• 회원 상세 패널: 가입 정보·로그인 이력·결제 이력·질문 이력.\n"
            "• 권한 편집: Pro 권한 부여/해제·1일 한도 변경·답변 속도 조정·차단/해제.\n"
            "• 모든 권한 변경은 ‘수정 이력’ 페이지에 기록(누가 언제 무엇을 바꿨는지)."
        ),
        "shots": ["admin_03_users_list.png", "admin_04_users_edits.png"],
    },
    {
        "title": "(4) 통계 — 가입자·피드백·매출·가입유형별",
        "what": "운영 의사결정을 돕는 통계 화면들입니다.",
        "result": (
            "• 가입자/구독자 추이(기간별 그래프 + 엑셀 내보내기).\n"
            "• 답변 피드백 통계(좋아요/싫어요 비율, 질문 분석).\n"
            "• 매출 대시보드(월별·환율 KRW 동시 표기·엑셀).\n"
            "• 가입유형별 통계(치과의사/치과위생사/기타 분포·연차 분포)."
        ),
        "shots": ["admin_05_analytics_signups.png", "admin_06_analytics_segments.png"],
    },
    {
        "title": "(5) RAG 데이터 관리",
        "what": "AI가 답변할 때 참고하는 ‘지식 자료’를 등록·수정·삭제하는 화면입니다.",
        "result": (
            "• 텍스트(.txt) 형식 자료 업로드 + 미리보기.\n"
            "• 기존 자료 수정·삭제.\n"
            "• 자료 변경 후 ‘재빌드’ 버튼 한 번으로 무중단 반영.\n"
            "• 동의어 사전 관리(예: ‘청구’ = ‘청구하기’ = ‘보험청구’)."
        ),
        "shots": ["admin_08_rag_data.png"],
    },
    {
        "title": "(6) AI 프롬프트 / 모델 파라미터",
        "what": "AI가 어떤 톤·길이로 답할지 조절하는 화면입니다.",
        "result": (
            "• 시스템 프롬프트 편집 + 차단 키워드 설정.\n"
            "• 모델 파라미터(창의성·답변 길이 등) 슬라이더 조정.\n"
            "• 변경 사항은 모두 ‘설정 이력’으로 기록되어 되돌릴 수 있습니다."
        ),
        "shots": ["admin_09_rag_prompts.png"],
    },
    {
        "title": "(7) 이상탐지 / 수동 차단",
        "what": "비정상적인 사용 패턴을 자동으로 감지하고 운영자가 대응하는 화면입니다.",
        "result": (
            "• 자동 감지: 비밀번호 5회 오류·동일 질문 반복·짧은 시간 동시 IP 로그인·동일 IP 반복.\n"
            "• 감지 즉시 운영자에게 알림톡 발송.\n"
            "• 운영자는 ‘차단’ 또는 ‘해제’ 직접 처리.\n"
            "• 차단된 사용자에게는 차단 사유가 인앱 쪽지로 전달됩니다."
        ),
        "shots": ["admin_10_anomaly.png"],
    },
    {
        "title": "(8) 결제·환불 관리",
        "what": "전체 결제 내역과 환불 요청을 관리하는 화면입니다.",
        "result": (
            "• 결제 이벤트 타임라인(누가 언제 결제했고 어떤 결과였는지).\n"
            "• 환불 검토 탭에서 부분환불 가능(일할 권장값 자동 계산).\n"
            "• 5년치 결제 데이터 보관."
        ),
        "shots": ["admin_11_finance_payments.png", "admin_07_finance_revenue.png"],
    },
    {
        "title": "(9) 예산·비상 정지 (Kill Switch)",
        "what": "API 비용 폭증을 막기 위한 안전장치 화면입니다.",
        "result": (
            "• 월 예산 설정 + 자동 알림(80%/95%/100%).\n"
            "• 100% 도달 시 무료 사용자 질의가 자동 차단됩니다.\n"
            "• 수동 ‘전체 정지’ 버튼(사유 입력 + 2단계 확인 필수).\n"
            "• 정지 기간만큼 유료 구독자의 다음 결제일도 자동 연기됩니다."
        ),
        "shots": ["admin_12_killswitch.png"],
    },
    {
        "title": "(10) CS 문의 관리 + 답변",
        "what": "사용자 1:1 문의에 답변하는 화면입니다.",
        "result": (
            "• 신규 문의가 오면 운영자에게 알림톡 발송.\n"
            "• 답변 본문 서식(글머리·굵게·링크) 지원.\n"
            "• 답변 등록 시 사용자에게 알림톡 자동 발송."
        ),
        "shots": ["admin_13_cs.png"],
    },
    {
        "title": "(11) 팝업·공지 관리 / 전역 설정",
        "what": "사이트 전체에 표시되는 공지 팝업과 전역 토글을 관리하는 화면입니다.",
        "result": (
            "• 팝업 등록·수정·삭제 + 노출 기간 설정.\n"
            "• 글로벌 토글: 구독 버튼 ON/OFF · 1일 한도 · 답변 지연시간 · 업그레이드 권유 팝업."
        ),
        "shots": ["admin_14_popups.png", "admin_15_settings.png"],
    },
    {
        "title": "(12) 관리자 활동 로그",
        "what": "모든 관리자 행동(권한 변경·차단·환불·설정 변경 등)이 기록되는 감사 로그 화면입니다.",
        "result": (
            "• 누가 / 언제 / 무엇을 했는지 시간순 정렬.\n"
            "• 변경 전/후 값 비교 보기.\n"
            "• 액션 종류·대상 회원·기간으로 필터링."
        ),
        "shots": ["admin_17_admins_logs.png"],
    },
]


# 인계 시 함께 전달될 문서 목록 (간소화)
DELIVERABLES: list[dict] = [
    {
        "category": "계약·정산",
        "items": [
            "최종 납품 확인서 (검수 체크리스트)",
        ],
    },
    {
        "category": "운영 안내",
        "items": [
            "Denvia 서버 안내서",
            "관리자 사이트 사용 매뉴얼",
            "결제·환불 운영 가이드",
            "장애 대응 안내서 (긴급 연락처 포함)",
        ],
    },
    {
        "category": "자산·계정 인계",
        "items": [
            "자산·계정 통합 인계서 (도메인 / AWS / 토스 / 알리고 / OAuth / API 키 / 관리자 계정)",
        ],
    },
    {
        "category": "법적 문서",
        "items": [
            "이용약관 (최종본 PDF)",
            "개인정보처리방침 (최종본 PDF)",
        ],
    },
    {
        "category": "유지보수",
        "items": [
            "유지보수 계약서 (별도)",
            "하자보수 범위·기간 안내",
        ],
    },
]


# ──────────────────────────────────────────────────────────────────────────────
# 스타일 헬퍼 (build_alimtalk_guide_v4.py와 동일한 호환 안전 패턴)
# ──────────────────────────────────────────────────────────────────────────────

KOREAN_FONT = "맑은 고딕"


def _apply_kr_font(run, *, bold: bool = False, size_pt: float | None = None,
                   color: RGBColor | None = None) -> None:
    run.font.name = KOREAN_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement

        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), KOREAN_FONT)
    rfonts.set(qn("w:hAnsi"), KOREAN_FONT)
    rfonts.set(qn("w:ascii"), KOREAN_FONT)
    if bold:
        run.bold = True
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color


def _add_paragraph(doc, text: str, *, bold: bool = False, size_pt: float | None = None,
                   align: int | None = None, color: RGBColor | None = None,
                   space_after_pt: float | None = None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_after_pt is not None:
        p.paragraph_format.space_after = Pt(space_after_pt)
    run = p.add_run(text)
    _apply_kr_font(run, bold=bold, size_pt=size_pt, color=color)


def _add_heading(doc, text: str, level: int = 1) -> None:
    sizes = {1: 18, 2: 14, 3: 12}
    spaces = {1: 16, 2: 12, 3: 8}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(spaces[level])
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _apply_kr_font(run, bold=True, size_pt=sizes[level])


def _shade_cell(cell, color_hex: str) -> None:
    from docx.oxml import OxmlElement

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, *, bold: bool = False, size_pt: float = 10,
                   color: RGBColor | None = None) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    for i, line in enumerate(text.split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        _apply_kr_font(r, bold=bold, size_pt=size_pt, color=color)


def _add_horizontal_rule(doc) -> None:
    from docx.oxml import OxmlElement

    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


SHOT_DIR = Path(__file__).resolve().parent.parent / "outputs" / "screenshots"


def _add_screenshot(doc, filename: str) -> None:
    """캡쳐 PNG 이미지를 본문 폭에 맞춰 삽입."""
    path = SHOT_DIR / filename
    if not path.exists():
        _add_paragraph(doc, f"[이미지 없음: {filename}]", size_pt=9,
                       color=RGBColor(0xAA, 0x33, 0x33))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(16.5))


def _add_screen_section(doc, screen: dict) -> None:
    _add_heading(doc, screen["title"], level=3)
    _add_paragraph(doc, f"무엇을 하는 화면인가요?", bold=True, size_pt=11)
    _add_paragraph(doc, screen["what"], size_pt=11)
    _add_paragraph(doc, "구현된 내용", bold=True, size_pt=11)
    _add_paragraph(doc, screen["result"], size_pt=11, space_after_pt=6)
    for shot in screen.get("shots", []):
        _add_screenshot(doc, shot)


# ──────────────────────────────────────────────────────────────────────────────
# 문서 빌드
# ──────────────────────────────────────────────────────────────────────────────

def build_document(output_path: Path) -> None:
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # ── 표지 ──
    _add_paragraph(doc, "Denvia", bold=True, size_pt=24,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "치과 보험청구·데스크 행정 AI 어시스턴트",
                   size_pt=11, align=WD_ALIGN_PARAGRAPH.CENTER,
                   color=RGBColor(0x55, 0x55, 0x55))
    _add_paragraph(doc, "", size_pt=8)
    _add_paragraph(doc, DOC_TITLE, bold=True, size_pt=20,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, f"— {DOC_SUBTITLE} —", size_pt=12,
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   color=RGBColor(0x55, 0x55, 0x55))
    _add_paragraph(doc, f"작성일: {DOC_DATE}", size_pt=11,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=24)

    # ── 1. 보고 개요 ──
    _add_heading(doc, "1. 보고 개요", level=1)
    _add_paragraph(
        doc,
        "본 문서는 치과 보험청구·데스크 행정 AI 어시스턴트(이하 “Denvia”) 프로젝트의 "
        "전체 개발이 완료된 시점의 최종 보고서입니다. 계약 범위 내 모든 핵심 기능 개발이 "
        "완료되어 운영 서버에 배포된 상태이며, 본 보고서 회신과 함께 인계 절차를 시작합니다.",
        size_pt=11,
    )
    _add_paragraph(
        doc,
        "중간 보고(2026년 4월 29일) 이후 추가 작업으로는 ① 결제·환불 시스템 전체(토스페이먼츠 연동) "
        "② 마이페이지 및 사용자 결제 흐름 ③ 관리자 통계·매출·이상탐지·예산 통제 ④ 관리자 4등급 권한 체계 "
        "⑤ 1:1 문의 게시판 + 운영자 답변 ⑥ AWS Lightsail 운영 서버 구축이 진행되어 모두 완료되었습니다.",
        size_pt=11,
    )
    _add_paragraph(
        doc,
        "본 보고서에는 (가) 전체 진행 현황 요약 (나) 화면별 개발 결과 "
        "(다) 인계 시 함께 전달되는 문서 목록 (라) 하자보수·유지보수 안내가 포함됩니다.",
        size_pt=11,
    )

    # ── 2. 전체 진행 현황 요약 ──
    _add_heading(doc, "2. 전체 진행 현황 요약", level=1)
    _add_paragraph(
        doc,
        "계약 범위 내 모든 큰 영역의 개발이 완료되었습니다. 영역별 요약은 다음과 같습니다.",
        size_pt=11,
    )

    tbl = doc.add_table(rows=1, cols=4)
    tbl.autofit = False
    tbl.columns[0].width = Cm(3.5)
    tbl.columns[1].width = Cm(6.5)
    tbl.columns[2].width = Cm(2.5)
    tbl.columns[3].width = Cm(4.0)
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["영역", "포함 범위", "상태", "비고"]):
        _set_cell_text(hdr[i], h, bold=True, size_pt=10)
        _shade_cell(hdr[i], "E0E0E0")
    for item in PROGRESS_SUMMARY:
        cells = tbl.add_row().cells
        _set_cell_text(cells[0], item["area"], bold=True, size_pt=10)
        _set_cell_text(cells[1], item["scope"], size_pt=10)
        _set_cell_text(cells[2], item["status"], size_pt=10,
                       color=RGBColor(0x10, 0x80, 0x30))
        _set_cell_text(cells[3], item["note"], size_pt=9,
                       color=RGBColor(0x55, 0x55, 0x55))
    _add_paragraph(doc, "", size_pt=6, space_after_pt=12)

    # ── 3. 화면별 개발 결과 ──
    _add_heading(doc, "3. 화면별 개발 결과", level=1)
    _add_paragraph(
        doc,
        "아래는 사용자가 실제로 만나는 화면들을 사용자/관리자로 나눠 정리한 내용입니다. "
        "각 화면은 운영 서버(denvia.ai.kr)에서 직접 사용해 보실 수 있습니다.",
        size_pt=11,
    )

    _add_heading(doc, "3-1. 사용자(고객) 화면", level=2)
    for screen in USER_SCREENS:
        _add_screen_section(doc, screen)
        _add_horizontal_rule(doc)

    _add_heading(doc, "3-2. 관리자(Admin) 화면", level=2)
    for screen in ADMIN_SCREENS:
        _add_screen_section(doc, screen)
        _add_horizontal_rule(doc)

    # ── 4. 인계 시 함께 전달되는 문서 ──
    _add_heading(doc, "4. 인계 시 함께 전달되는 문서", level=1)
    _add_paragraph(
        doc,
        "본 보고서와 함께 다음 문서들이 ‘자료’ 폴더에 일괄 정리되어 전달됩니다. "
        "별도 문서가 필요한 항목은 추가 제작 후 일괄 전달합니다.",
        size_pt=11,
    )

    deliv_tbl = doc.add_table(rows=1, cols=2)
    deliv_tbl.autofit = False
    deliv_tbl.columns[0].width = Cm(4.0)
    deliv_tbl.columns[1].width = Cm(12.5)
    hdr = deliv_tbl.rows[0].cells
    for i, h in enumerate(["분류", "문서 목록"]):
        _set_cell_text(hdr[i], h, bold=True, size_pt=10)
        _shade_cell(hdr[i], "E0E0E0")
    for cat in DELIVERABLES:
        cells = deliv_tbl.add_row().cells
        _set_cell_text(cells[0], cat["category"], bold=True, size_pt=10)
        _set_cell_text(cells[1], "\n".join(f"• {it}" for it in cat["items"]), size_pt=10)
    _add_paragraph(doc, "", size_pt=6, space_after_pt=12)

    # ── 5. 하자보수·유지보수 안내 ──
    _add_heading(doc, "5. 하자보수·유지보수 안내", level=1)
    _add_paragraph(
        doc,
        "본 개발 완료 시점부터 계약서(제11조)에 정한 무상 하자보수가 진행되며, "
        "이후에는 별도 유지보수 계약을 통해 운영을 지원해 드립니다. "
        "상세 조건은 별도 ‘유지보수 계약서’에 명시됩니다.",
        size_pt=11,
    )

    maint_tbl = doc.add_table(rows=0, cols=2)
    maint_tbl.autofit = False
    maint_tbl.columns[0].width = Cm(4.5)
    maint_tbl.columns[1].width = Cm(12.0)
    for label, value in [
        ("무상 하자보수 범위",
         "합의된 기능명세 내에서 발견된 오작동·오류·버그의 수정 (계약서 제11조 제2항)"),
        ("무상 하자보수 기간",
         "최종 검수 완료일로부터 30일 (계약서 제11조 제1항)"),
        ("무상 범위 제외",
         "기능 추가·디자인 변경·외부 API 정책 변경 대응 등은 별도 유료 작업"),
        ("긴급 장애 대응",
         "운영 서버 다운·결제 시스템 마비 등 서비스 중단 장애는 즉시 대응(24시간 이내 1차 회신)"),
        ("일반 문의 대응",
         "평일 업무시간 기준 1~3 영업일 내 회신"),
    ]:
        row = maint_tbl.add_row().cells
        _set_cell_text(row[0], label, bold=True, size_pt=10)
        _shade_cell(row[0], "F2F2F2")
        _set_cell_text(row[1], value, size_pt=10)
    _add_paragraph(doc, "", size_pt=6, space_after_pt=12)

    # ── 끝맺음 ──
    _add_horizontal_rule(doc)
    _add_paragraph(
        doc,
        "본 보고서에 대한 의견·보완 요청 사항은 회신 부탁드리며, 회신 후 "
        "납품·인계 절차를 진행하겠습니다. 그동안 협력해 주셔서 감사드립니다.",
        size_pt=11, align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _add_paragraph(doc, "Denvia 개발팀", size_pt=11,
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   color=RGBColor(0x55, 0x55, 0x55))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    project_root = Path(__file__).resolve().parent.parent
    out = project_root / "자료" / "Denvia_개발완료보고서_2026-06-05.docx"
    build_document(out)
    print(f"[OK] {out}")


if __name__ == "__main__":
    main()
