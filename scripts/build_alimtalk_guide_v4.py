"""고객 전달용 알림톡 안내서 v4 (확정본) 빌더.

2026-05-18 — Google Docs 회신본(v3 검수)을 반영한 확정본을 생성합니다.

산출물:
    자료/Denvia_카카오_알림톡_메시지_안내서_v4.docx

고객 전달용 문서 작성 규칙 준수 — SDT(content control), 폼 보호,
문서 보호, ActiveX 등 일체 사용하지 않음. 어떤 워드 버전에서도 자유 편집 가능.

실행:
    py -3 scripts/build_alimtalk_guide_v4.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ──────────────────────────────────────────────────────────────────────────────
# 콘텐츠 정의
# ──────────────────────────────────────────────────────────────────────────────

DOC_TITLE = "Denvia 카카오 알림톡 메시지 안내서 v4 (확정본)"
DOC_SUBTITLE = "— 2026-05-18 v3 검수 회신본 반영 완료 —"

# 사용자(Part 1) 알림톡 — v4 확정. 폐기 항목은 별도 §에 표기.
USER_MESSAGES: list[dict] = [
    {
        "section": "① 결제 안내",
        "no": "1-1",
        "title": "첫 구독 결제 완료",
        "when": "처음으로 Pro 구독을 시작하면서 결제가 정상 처리된 직후",
        "to": "결제하신 고객님",
        "body": (
            "안녕하세요, Denvia입니다.\n"
            "결제가 정상적으로 완료되어 Pro 구독이 시작되었습니다.\n"
            "문의사항은 앱 내 “문의작성”을 통해 언제든 문의해주세요.\n"
            "감사합니다."
        ),
        "note": "v3 검수 → 수정 반영. 결제 금액·다음 결제일은 본문에서 제거(마이페이지에서 확인).",
    },
    {
        "section": "① 결제 안내",
        "no": "1-4",
        "title": "결제 실패 안내 (1차)",
        "when": "자동결제 첫 실패 직후 (1일 뒤 자동 재시도 예정)",
        "to": "결제 실패한 고객님",
        "body": (
            "안녕하세요 Denvia 입니다.\n"
            "pro 구독 정기결제에 실패했습니다.\n"
            "1일 후 자동 재시도됩니다.\n"
            "카드 정보를 확인해주세요."
        ),
        "note": "v3 검수 → 수정 반영 (인사말 추가).",
    },
    {
        "section": "① 결제 안내",
        "no": "1-5",
        "title": "결제 실패 안내 (2차)",
        "when": "1차 재시도까지 실패했을 때 (3일 뒤 마지막 재시도 예정)",
        "to": "결제 실패한 고객님",
        "body": (
            "Denvia 구독 결제가 다시 실패했습니다.\n"
            "3일 후 마지막으로 재시도됩니다.\n"
            "3일 후 결제 실패가 일어나는 경우 pro 구독은 자동 해지될 수 있습니다.\n"
            "문의사항은 앱 내 “문의작성”을 통해 언제든 문의해주세요.\n"
            "감사합니다."
        ),
        "note": "v3 검수 → 수정 반영 (자동 해지 안내·문의 안내·감사 인사 추가).",
    },
    {
        "section": "① 결제 안내",
        "no": "1-6",
        "title": "결제 최종 실패",
        "when": "3회(최종) 재시도까지 모두 실패 → 구독 해지 예정",
        "to": "결제 실패한 고객님",
        "body": (
            "pro 구독 결제가 최종 실패했습니다.\n"
            "pro구독이 해지될 예정입니다.\n"
            "고객센터 문의: https://denvia.kr/support"
        ),
        "note": "v3 검수 → 수정 반영. 메시지 하단에 [고객센터 바로가기] 버튼 함께 표시.",
    },
    {
        "section": "① 결제 안내",
        "no": "1-7",
        "title": "환불 처리 완료",
        "when": "청약철회(즉시 해지 + 전액 환불), 운영자 수동 전액 환불, 운영자 수동 부분 환불 — 세 경우에 동일한 템플릿으로 발송되며 '처리 유형'이 자동으로 바뀜",
        "to": "환불받으신 고객님",
        "body": (
            "환불이 완료되었습니다.\n"
            "처리 유형: 즉시 해지 및 전액 환불\n"
            "결제 원금: 30,000원\n"
            "환불 금액: 30,000원\n"
            "처리일: 2026년 5월 12일\n"
            "\n"
            "영업일 3~4일 내 결제 카드로 입금됩니다.\n"
            "문의는 Denvia 1:1 문의 게시판을 이용해주세요."
        ),
        "note": "v3 검수 → 확인. ‘처리 유형’은 ‘즉시 해지 및 전액 환불’ / ‘전액 환불’ / ‘부분 환불’ 중 하나로 표시됨.",
    },
    {
        "section": "② 구독 안내",
        "no": "2-1",
        "title": "구독 해지 예약 완료",
        "when": "마이페이지에서 해지를 신청한 직후 (즉시 해지가 아니라 다음 결제일까지 사용 가능)",
        "to": "해지 신청한 고객님",
        "body": (
            "Pro 구독 해지가 예약되었습니다.\n"
            "2026-06-15까지 pro구독 서비스를 이용하실 수 있으며 이후 무료버전으로 전환됩니다.\n"
            "감사합니다."
        ),
        "note": "v3 검수 → 수정 반영 (무료버전 전환 안내 추가).",
    },
    {
        "section": "② 구독 안내",
        "no": "2-2",
        "title": "구독 해지 완료",
        "when": "해지 예정일에 도달해 실제로 구독이 종료된 시점",
        "to": "해지된 고객님",
        "body": (
            "Denvia Pro 구독이 해지되었습니다.\n"
            "홍길동 님은 2026-02-13 일 부터 무료버전으로 전환 됩니다.\n"
            "감사합니다."
        ),
        "note": "v3 검수 → 수정 반영. 사용자 표시 이름과 전환일은 자동으로 채워집니다.",
    },
    {
        "section": "② 구독 안내",
        "no": "2-3",
        "title": "구독 해지 철회 완료",
        "when": "해지 예약 상태에서 ‘해지 취소’를 누른 직후",
        "to": "해지 철회한 고객님",
        "body": "Pro 구독 해지가 철회되었습니다.",
        "note": "v3 검수 → 수정 반영 (한 줄로 간결화).",
    },
    {
        "section": "③ 고객문의",
        "no": "3-1",
        "title": "고객문의 답변 도착",
        "when": "운영자가 사용자의 1:1 문의에 답변을 등록한 직후",
        "to": "문의를 작성한 고객님",
        "body": (
            "안녕하세요. Denvia 입니다.\n"
            "고객문의에 답변이 등록되었습니다.\n"
            "문의 제목: 결제가 이중으로 빠진 것 같아요\n"
            "쪽지함에서 자세한 내용을 확인해주세요."
        ),
        "note": "v3 검수 → 수정 반영 (인사말 추가). 메시지 하단에 [쪽지함 바로가기] 버튼 함께 표시.",
    },
    {
        "section": "④ 공지사항",
        "no": "4-1",
        "title": "Denvia 공지사항",
        "when": "운영자가 공지 푸시 발송 버튼을 눌렀을 때 (야간 21~08시는 다음날 08시에 묶어서 발송)",
        "to": "공지 대상 고객님 (전체/등급별/세그먼트별)",
        "body": (
            "점검 안내\n"
            "\n"
            "5월 18일 02:00~03:00 사이에 서버 점검이 진행됩니다.\n"
            "점검 시간 동안 서비스 이용이 일시 중단됩니다.\n"
            "이용에 불편을 드려 죄송합니다."
        ),
        "note": "v3 검수 → 확인. 제목과 본문은 공지 내용에 따라 달라집니다.",
    },
]


DEPRECATED_USER_MESSAGES: list[dict] = [
    {"no": "1-2", "title": "구독 자동 갱신 완료", "reason": "고객 검수에서 ‘삭제 요청’으로 회신 — 자동 갱신 시 알림 미발송."},
    {"no": "1-3", "title": "결제 재시도 성공", "reason": "고객 검수에서 ‘삭제 요청’으로 회신 — 재시도 성공 시 알림 미발송."},
    {"no": "5-1", "title": "계정 잠금 안내 (비밀번호 3회 오류)", "reason": "이상탐지 사용자 알림 일괄 폐기. 차단 시 1:1 쪽지·팝업으로 별도 전달."},
    {"no": "5-2", "title": "답변 속도 제한 안내 (비정상 빠른 질문)", "reason": "이상탐지 사용자 알림 일괄 폐기."},
    {"no": "5-3", "title": "동시 로그인 감지 안내", "reason": "이상탐지 사용자 알림 일괄 폐기."},
    {"no": "5-4", "title": "반복 질문 감지 안내", "reason": "이상탐지 사용자 알림 일괄 폐기."},
    {"no": "5-5", "title": "이용 제한 해제 안내", "reason": "이상탐지 사용자 알림 일괄 폐기."},
    {"no": "5-6", "title": "계정 차단 안내", "reason": "이상탐지 사용자 알림 일괄 폐기. 차단 시 인앱 1:1 쪽지·팝업으로 사유 전달."},
]


_ADMIN_RECIPIENT = "Denvia 운영자"


# 관리자(Part 2) 알림톡 — v4 확정
ADMIN_MESSAGES: list[dict] = [
    {
        "no": "6-1",
        "title": "RAG(지식베이스) 재빌드 완료",
        "when": "관리자 페이지에서 시작한 RAG 재빌드가 정상 완료된 직후",
        "to": _ADMIN_RECIPIENT,
        "body": "Denvia RAG 재빌드가 완료되었습니다.\n활성 청크 수: 1,247",
        "note": "v3 검수 → 확인.",
    },
    {
        "no": "6-2",
        "title": "RAG 재빌드 실패",
        "when": "RAG 재빌드 도중 오류가 발생해 실패했을 때",
        "to": _ADMIN_RECIPIENT,
        "body": "Denvia RAG 재빌드가 실패했습니다.\n오류: OpenAI API 시간 초과",
        "note": "v3 검수 → 확인.",
    },
    {
        "no": "6-3",
        "title": "월 예산 80% 도달",
        "when": "이번 달 OpenAI API 사용액이 월 예산의 80%를 처음 넘긴 시점 (한 달에 한 번만)",
        "to": _ADMIN_RECIPIENT,
        "body": (
            "[Denvia] 월 예산 80% 도달\n"
            "\n"
            "이번 달 OpenAI API 비용이 월 예산의 82%에 도달했습니다.\n"
            "현재 사용액: $82.30\n"
            "월 한도: $100.00\n"
            "관리자 대시보드에서 사용 패턴을 점검해주세요."
        ),
        "note": "v3 검수 → 확인.",
    },
    {
        "no": "6-4",
        "title": "월 예산 95% 도달 (경고)",
        "when": "이번 달 사용액이 월 예산의 95%를 처음 넘긴 시점 (한 달에 한 번만)",
        "to": _ADMIN_RECIPIENT,
        "body": (
            "[Denvia] 월 예산 95% 도달 — 경고\n"
            "\n"
            "이번 달 OpenAI API 비용이 월 예산의 96%에 도달했습니다.\n"
            "100% 도달 시 무료 질의가 자동 차단됩니다.\n"
            "현재 사용액: $96.10 / $100.00"
        ),
        "note": "v3 검수 → 확인.",
    },
    {
        "no": "6-5",
        "title": "월 예산 100% 소진 (무료 질의 자동 차단)",
        "when": "이번 달 사용액이 월 예산을 모두 소진해 무료 사용자 질의가 자동 차단된 시점",
        "to": _ADMIN_RECIPIENT,
        "body": (
            "[Denvia] 월 예산 소진 — 무료 질의 자동 차단\n"
            "\n"
            "월 예산 $100.00이 소진되어 무료 사용자 질의가 일시 차단되었습니다.\n"
            "유료 사용자는 영향 없습니다.\n"
            "다음 달 1일 자동 해제 또는 예산 상향 시 즉시 재개됩니다."
        ),
        "note": "v3 검수 → 확인.",
    },
    {
        "no": "6-6",
        "title": "신규 1:1 문의 접수 알림",
        "when": "사용자가 1:1 문의를 등록한 직후",
        "to": _ADMIN_RECIPIENT,
        "body": (
            "[Denvia] 새 1:1 문의 접수\n"
            "\n"
            "새 1:1 문의가 접수되었습니다.\n"
            "작성자: 홍길동\n"
            "문의 제목: 결제가 이중으로 빠진 것 같아요\n"
            "관리자 페이지에서 답변을 처리해주세요."
        ),
        "note": "v3 검수 → 확인. 메시지 하단에 [관리자 페이지 바로가기] 버튼 함께 표시.",
    },
    {
        "no": "6-7",
        "title": "이상탐지 발생 알림 (신규 — v4 추가)",
        "when": "사용자 계정에서 이상탐지가 감지된 직후. 비밀번호 3회 오류 / 동시 로그인 / 동일 질문 반복 / 답변 출력 후 3초 이내 연속 질문 / 중복 IP 등 모든 이상탐지 패턴에 동일하게 발송. 무료/유료 사용자 동일.",
        "to": _ADMIN_RECIPIENT,
        "body": (
            "[Denvia] 이상탐지 발생\n"
            "\n"
            "이상탐지내용: 비밀번호 3회 오류\n"
            "이상탐지 계정: user@example.com\n"
            "관리자 페이지에서 확인 필요합니다."
        ),
        "note": (
            "v4 신규 — 고객 추가요청 반영. 이상탐지는 사용자에게 카톡 미발송, 관리자에게만 통지. "
            "‘이상탐지내용’은 한국어 사유 라벨로 자동 채워집니다(예: ‘비밀번호 3회 오류’, ‘동시 로그인’, ‘동일 질문 반복’, ‘IP 중복’). "
            "운영자는 관리자 페이지에서 기록 확인 후 ‘해제’ 또는 ‘차단’을 직접 처리하며, 차단 시 사용자에게는 인앱 1:1 쪽지·차단 페이지 팝업으로 사유를 안내합니다."
        ),
    },
]


# ──────────────────────────────────────────────────────────────────────────────
# 스타일 헬퍼 (저레벨 OXML — 어떤 워드 버전에서도 호환되도록 단순 표현만 사용)
# ──────────────────────────────────────────────────────────────────────────────

KOREAN_FONT = "맑은 고딕"


def _apply_kr_font(run, *, bold: bool = False, size_pt: float | None = None, color: RGBColor | None = None) -> None:
    """모든 run에 한글 폰트 + 옵션을 일괄 적용."""
    run.font.name = KOREAN_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement

        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), KOREAN_FONT)
    rfonts.set(qn("w:hAnsi"), KOREAN_FONT)
    rfonts.set(qn("w:ascii"), KOREAN_FONT)
    if bold:
        run.bold = True
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color


def _add_paragraph(doc: Document, text: str, *, bold: bool = False, size_pt: float | None = None,
                   align: int | None = None, color: RGBColor | None = None,
                   space_after_pt: float | None = None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_after_pt is not None:
        p.paragraph_format.space_after = Pt(space_after_pt)
    run = p.add_run(text)
    _apply_kr_font(run, bold=bold, size_pt=size_pt, color=color)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    sizes = {1: 18, 2: 14, 3: 12}
    spaces = {1: 16, 2: 12, 3: 8}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(spaces[level])
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _apply_kr_font(run, bold=True, size_pt=sizes[level])


def _shade_cell(cell, color_hex: str) -> None:
    """셀 배경색 지정 (워드 호환 — 단순 fill 속성만 사용)."""
    from docx.oxml import OxmlElement

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _add_table_row(table, label: str, value: str, *, label_bold: bool = True,
                   label_shade: str | None = "F2F2F2") -> None:
    row = table.add_row().cells
    row[0].text = ""
    row[1].text = ""
    for cell in row:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    # 라벨 셀
    p0 = row[0].paragraphs[0]
    r0 = p0.add_run(label)
    _apply_kr_font(r0, bold=label_bold, size_pt=10)
    if label_shade:
        _shade_cell(row[0], label_shade)
    # 값 셀
    for i, line in enumerate(value.split("\n")):
        p = row[1].paragraphs[0] if i == 0 else row[1].add_paragraph()
        r = p.add_run(line)
        _apply_kr_font(r, size_pt=10)


def _add_message_body_box(doc: Document, body: str) -> None:
    """알림톡 본문을 노란 배경 한 칸 표로 렌더."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Cm(15)
    cell = tbl.cell(0, 0)
    cell.width = Cm(15)
    _shade_cell(cell, "FFF8DC")  # 연한 노랑 (cornsilk-ish)
    # 첫 단락에 본문 입력
    for i, line in enumerate(body.split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        _apply_kr_font(r, size_pt=11)
    doc.add_paragraph()  # 박스 다음 간격


def _add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    from docx.oxml import OxmlElement

    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_message_section(doc: Document, msg: dict) -> None:
    _add_heading(doc, f"{msg['no']}. {msg['title']}", level=3)

    # 메타 정보 표
    meta_tbl = doc.add_table(rows=0, cols=2)
    meta_tbl.autofit = False
    meta_tbl.columns[0].width = Cm(3.5)
    meta_tbl.columns[1].width = Cm(13)
    _add_table_row(meta_tbl, "언제 발송되나요?", msg["when"])
    _add_table_row(meta_tbl, "받는 사람", msg["to"])

    doc.add_paragraph()
    _add_paragraph(doc, "실제로 받게 되는 메시지", bold=True, size_pt=11)
    _add_message_body_box(doc, msg["body"])

    if msg.get("note"):
        _add_paragraph(
            doc,
            f"※ {msg['note']}",
            size_pt=10,
            color=RGBColor(0x55, 0x55, 0x55),
            space_after_pt=14,
        )


# ──────────────────────────────────────────────────────────────────────────────
# 문서 생성
# ──────────────────────────────────────────────────────────────────────────────

def build_document(output_path: Path) -> None:
    doc = Document()

    # 페이지 여백
    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # 표지
    _add_paragraph(doc, "Denvia", bold=True, size_pt=24, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "카카오 알림톡 발송 메시지 안내서 (확정본 v4)",
                   bold=True, size_pt=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, DOC_SUBTITLE, size_pt=11,
                   align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x55, 0x55, 0x55))
    _add_paragraph(doc, "작성일: 2026년 5월 18일", size_pt=11,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=24)

    # 머리말
    _add_heading(doc, "이 문서의 목적", level=2)
    _add_paragraph(
        doc,
        "본 문서는 고객님께서 2026-05-18 v3 검수에서 회신해 주신 내용을 모두 반영한 "
        "최종 확정본입니다. 본 문서에 수록된 메시지가 곧 카카오 비즈니스 검수에 등록되며, "
        "검수 통과 후 알리고 콘솔에 그대로 등록됩니다. 추가 수정이 필요한 부분은 "
        "본 문서 회신 후 별도로 알려주시면, 카카오 사전 승인을 다시 요청해 반영합니다.",
        size_pt=11,
    )
    _add_paragraph(
        doc,
        "v3 대비 주요 변경: ① 자동 갱신/재시도 성공 알림 폐기 ② 결제·구독·고객문의 본문 7건 수정 "
        "③ 이상탐지 사용자 알림 6건 전면 폐기 ④ 이상탐지 관리자 알림 신규 추가.",
        size_pt=11,
    )

    # 분류 한눈에 보기
    _add_heading(doc, "발송되는 메시지 한눈에 보기", level=2)
    summary_tbl = doc.add_table(rows=1, cols=4)
    summary_tbl.autofit = False
    hdr = summary_tbl.rows[0].cells
    for i, h in enumerate(["분류", "수신자", "어떤 상황", "메시지 개수"]):
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        _apply_kr_font(r, bold=True, size_pt=10)
        _shade_cell(hdr[i], "E0E0E0")

    summary_rows = [
        ("결제 안내", "사용자", "첫 결제·결제 실패(1/2차)·최종 실패·환불 완료", "5개"),
        ("구독 안내", "사용자", "해지 신청·해지 완료·해지 철회", "3개"),
        ("고객문의", "사용자", "1:1 문의 답변 등록 시", "1개"),
        ("공지사항", "사용자", "전체/그룹 공지 푸시 (야간 보류)", "1개"),
        ("관리자 알림", "관리자", "RAG·예산 경고·1:1 문의 접수·이상탐지", "7개 (이상탐지 1건 신규)"),
    ]
    for s in summary_rows:
        cells = summary_tbl.add_row().cells
        for i, v in enumerate(s):
            p = cells[i].paragraphs[0]
            r = p.add_run(v)
            _apply_kr_font(r, size_pt=10)
    _add_paragraph(doc, "", size_pt=6, space_after_pt=12)

    # Part 1
    _add_heading(doc, "Part 1. 사용자에게 발송되는 메시지", level=1)
    current_section = None
    for msg in USER_MESSAGES:
        if msg["section"] != current_section:
            current_section = msg["section"]
            _add_heading(doc, current_section, level=2)
        _add_message_section(doc, msg)
        _add_horizontal_rule(doc)

    # 폐기 안내
    _add_heading(doc, "[참고] v3 → v4에서 폐기된 메시지", level=2)
    _add_paragraph(
        doc,
        "아래 항목은 고객 검수 결과 또는 추가 요청에 따라 v4에서 발송 대상에서 제외되었습니다. "
        "카카오 비즈니스/알리고 콘솔에도 등록하지 않습니다.",
        size_pt=11,
    )
    dep_tbl = doc.add_table(rows=1, cols=3)
    dep_tbl.autofit = False
    hdr = dep_tbl.rows[0].cells
    for i, h in enumerate(["번호", "메시지 이름", "폐기 사유"]):
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        _apply_kr_font(r, bold=True, size_pt=10)
        _shade_cell(hdr[i], "E0E0E0")
    for dep in DEPRECATED_USER_MESSAGES:
        cells = dep_tbl.add_row().cells
        cells[0].text = ""
        cells[1].text = ""
        cells[2].text = ""
        for i, v in enumerate([dep["no"], dep["title"], dep["reason"]]):
            p = cells[i].paragraphs[0]
            r = p.add_run(v)
            _apply_kr_font(r, size_pt=10)
    _add_paragraph(doc, "", size_pt=6, space_after_pt=12)

    # Part 2
    _add_heading(doc, "Part 2. 운영자(관리자)에게 발송되는 메시지", level=1)
    _add_paragraph(
        doc,
        "아래 메시지들은 Denvia 서비스 운영자 본인의 휴대폰으로 발송되며, "
        "일반 사용자에게는 발송되지 않습니다. 운영 모니터링과 빠른 대응을 위한 알림입니다.",
        size_pt=11,
    )
    for msg in ADMIN_MESSAGES:
        _add_message_section(doc, msg)
        _add_horizontal_rule(doc)

    # 부록 A. SMS
    _add_heading(doc, "부록 A. 일반 SMS로 발송되는 항목 (참고)", level=1)
    _add_paragraph(
        doc,
        "다음 두 가지는 카카오 알림톡이 아닌 일반 SMS 문자로 발송됩니다. "
        "보안상 자유 텍스트로 발송되어야 하는 항목이며, 별도 템플릿 등록이 필요 없습니다.",
        size_pt=11,
    )

    _add_heading(doc, "A-1. SMS 인증번호 발송", level=3)
    _add_paragraph(doc,
                   "트리거: 회원가입 / 아이디찾기 / 비밀번호찾기 화면에서 ‘인증번호 받기’ 클릭 시. "
                   "재발송 제한: 60초 쿨다운, 시간당 최대 3회.",
                   size_pt=11)
    _add_message_body_box(doc, "[Denvia] 인증번호: 482917\n3분 안에 입력해주세요.")

    _add_heading(doc, "A-2. 임시 비밀번호 발송", level=3)
    _add_paragraph(doc,
                   "트리거: 비밀번호찾기에서 이메일·휴대폰 번호 일치 확인 직후. "
                   "임시 비밀번호는 한 번만 표시되며 로그인 후 즉시 변경하셔야 합니다.",
                   size_pt=11)
    _add_message_body_box(doc, "[Denvia] 임시 비밀번호: A7b2K9pX\n로그인 후 즉시 변경해주세요.")

    # 부록 B. 변경 이력
    _add_heading(doc, "부록 B. v1 → v4 변경 이력", level=1)
    history_tbl = doc.add_table(rows=1, cols=2)
    history_tbl.autofit = False
    history_tbl.columns[0].width = Cm(3.0)
    history_tbl.columns[1].width = Cm(13.5)
    hdr = history_tbl.rows[0].cells
    for i, h in enumerate(["일자", "변경 내용"]):
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        _apply_kr_font(r, bold=True, size_pt=10)
        _shade_cell(hdr[i], "E0E0E0")
    history = [
        ("2026-05-07", "v1 최초 작성. 알리고 공급자 확정. 18개 템플릿 카탈로그화."),
        ("2026-05-12", "환불 정책 재편 — `billing.refund_success` 변수 확장(처리 유형·결제 원금·환불 금액 구분 표시), `billing.refund_denied` 폐기."),
        ("2026-05-15", "v3 검수 안내서 발행 — 23개 메시지에 ‘수정/삭제/확인’ 회신 요청. 이상탐지 6건 신규 제안 포함."),
        ("2026-05-18",
         "v4 확정. v3 회신 반영 — ① 폐기 2건(자동 갱신/재시도 성공) ② 본문 수정 7건(첫 결제·결제 실패 1/2/최종·해지 예약/완료/철회·고객문의 답변) "
         "③ 이상탐지 사용자 알림 6건 일괄 폐기 ④ 이상탐지 관리자 알림 신규(6-7)."),
    ]
    for d, c in history:
        cells = history_tbl.add_row().cells
        cells[0].text = ""
        cells[1].text = ""
        for i, v in enumerate([d, c]):
            p = cells[i].paragraphs[0]
            r = p.add_run(v)
            _apply_kr_font(r, size_pt=10)

    _add_paragraph(doc, "", size_pt=6, space_after_pt=24)
    _add_paragraph(doc, "감사합니다.", bold=True, size_pt=12,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    repo_root = Path(__file__).resolve().parent.parent
    output_path = repo_root / "자료" / "Denvia_카카오_알림톡_메시지_안내서_v4.docx"
    build_document(output_path)
    print(f"[OK] Wrote {output_path}")


if __name__ == "__main__":
    main()
