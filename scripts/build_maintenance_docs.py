"""Denvia 유지보수 계약서 + 하자보수 범위·기간 안내서 빌더.

2026-06-05 — 개발 완료 인계 시점에 맞춰 두 개의 고객 전달용 문서를 동시에 생성한다.

산출물:
    자료/Denvia_유지보수_계약서.docx
    자료/Denvia_하자보수_범위_기간_안내서.docx

근거:
- 본 개발 계약서 제11조 / 제11조의2 / 제12조 를 이행 문서로 풀어낸 것.
- 무상 하자보수 30일, 그 이후의 유지보수는 별도 유상 계약 — 두 문서가 한 쌍.

CLAUDE.md "고객 전달용 문서 작성 규칙" 준수 — SDT(content control), 폼 보호,
문서 보호, ActiveX 등 일체 사용하지 않음. 어떤 워드 버전에서도 자유 편집 가능.

당사자·금액·기간 등 가변 값은 원 계약서와 동일하게 `[      ]` 표기로 두어
고객과 협의 후 직접 채워 넣도록 한다.

실행:
    py -3 scripts/build_maintenance_docs.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ──────────────────────────────────────────────────────────────────────────────
# 공통 상수
# ──────────────────────────────────────────────────────────────────────────────

KOREAN_FONT = "맑은 고딕"

PROJECT_NAME = "Denvia (치과 보험청구·데스크 행정 AI 어시스턴트)"
SERVICE_DOMAIN = "denvia.ai.kr"

CLIENT = {
    "label": "갑 (클라이언트)",
    "company": "치과챗봇",
    "rep": "이규성",
    "biz_no": "980602-1",
    "addr": "독산동 1004-12 더팰리스80",
    "tel": "010-2323-2753",
    "email": "dlrbtjd357@naver.com",
}

VENDOR = {
    "label": "을 (수행사)",
    "company": "브레이크 더 몰드 (Break The mould)",
    "rep": "정형우",
    "biz_no": "206-26-68137",
    "addr": "경기도 양주시 옥정서로 5길 60",
    "tel": "010-2484-0289",
    "email": "btmdesign@naver.com",
}


# ──────────────────────────────────────────────────────────────────────────────
# 스타일 헬퍼 (build_final_report.py 와 동일한 호환 안전 패턴)
# ──────────────────────────────────────────────────────────────────────────────

def _apply_kr_font(run, *, bold: bool = False, size_pt: float | None = None,
                   color: RGBColor | None = None) -> None:
    run.font.name = KOREAN_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement

        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), KOREAN_FONT)
    rfonts.set(qn("w:hAnsi"), KOREAN_FONT)
    rfonts.set(qn("w:ascii"), KOREAN_FONT)
    if bold:
        run.bold = True
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color


def _add_paragraph(doc, text: str, *, bold: bool = False, size_pt: float | None = None,
                   align: int | None = None, color: RGBColor | None = None,
                   space_after_pt: float | None = None,
                   left_indent_cm: float | None = None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_after_pt is not None:
        p.paragraph_format.space_after = Pt(space_after_pt)
    if left_indent_cm is not None:
        p.paragraph_format.left_indent = Cm(left_indent_cm)
    run = p.add_run(text)
    _apply_kr_font(run, bold=bold, size_pt=size_pt, color=color)


def _add_heading(doc, text: str, level: int = 1) -> None:
    sizes = {1: 16, 2: 13, 3: 11}
    spaces_before = {1: 14, 2: 10, 3: 6}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(spaces_before[level])
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _apply_kr_font(run, bold=True, size_pt=sizes[level])


def _shade_cell(cell, color_hex: str) -> None:
    from docx.oxml import OxmlElement

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, *, bold: bool = False, size_pt: float = 10,
                   color: RGBColor | None = None) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    for i, line in enumerate(text.split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        _apply_kr_font(r, bold=bold, size_pt=size_pt, color=color)


def _add_horizontal_rule(doc) -> None:
    from docx.oxml import OxmlElement

    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_margins(doc) -> None:
    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)


def _add_party_block(doc, party: dict) -> None:
    """제1조 (당사자) 형식의 정보 블록."""
    _add_paragraph(doc, party["label"], bold=True, size_pt=11)
    for k, v in [
        ("상호 또는 성명", party["company"]),
        ("대표자", party["rep"]),
        ("사업자등록번호(또는 주민등록번호)", party["biz_no"]),
        ("주소", party["addr"]),
        ("연락처", party["tel"]),
        ("이메일", party["email"]),
    ]:
        _add_paragraph(doc, f"  · {k}: {v}", size_pt=10.5)


def _add_signature_block(doc) -> None:
    """말미 서명란."""
    _add_paragraph(doc, "", size_pt=8)
    _add_paragraph(
        doc,
        "본 계약의 성립을 증명하기 위하여 계약서 2부를 작성하고, "
        "갑과 을이 각각 서명 또는 기명 날인 후 1부씩 보관한다.",
        size_pt=11, align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _add_paragraph(doc, "", size_pt=6)
    _add_paragraph(doc, "계약 체결일:       [        ]년    [    ]월    [    ]일",
                   size_pt=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "", size_pt=10)

    tbl = doc.add_table(rows=2, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Cm(8.25)
    tbl.columns[1].width = Cm(8.25)

    # 좌: 갑 / 우: 을
    sig_block = (
        "{label}\n"
        "  · 상호 또는 성명: {company}\n"
        "  · 대   표   자: {rep}\n"
        "  · 서명 / 날인:\n"
        "\n"
    )
    _set_cell_text(tbl.rows[0].cells[0],
                   sig_block.format(**CLIENT), size_pt=10.5)
    _set_cell_text(tbl.rows[0].cells[1],
                   sig_block.format(**VENDOR), size_pt=10.5)
    # 빈 줄 한 줄 추가
    _set_cell_text(tbl.rows[1].cells[0], " ", size_pt=10)
    _set_cell_text(tbl.rows[1].cells[1], " ", size_pt=10)


# ──────────────────────────────────────────────────────────────────────────────
# 문서 1 — 유지보수 계약서
# ──────────────────────────────────────────────────────────────────────────────

def build_maintenance_contract(output_path: Path) -> None:
    doc = Document()
    _set_margins(doc)

    # 표지
    _add_paragraph(doc, "Denvia", bold=True, size_pt=22,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "치과 보험청구·데스크 행정 AI 어시스턴트", size_pt=11,
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   color=RGBColor(0x55, 0x55, 0x55))
    _add_paragraph(doc, "", size_pt=10)
    _add_paragraph(doc, "유 지 보 수 계 약 서", bold=True, size_pt=22,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "— 개발 완료 후 운영 단계 유지보수에 관한 약정 —",
                   size_pt=11, align=WD_ALIGN_PARAGRAPH.CENTER,
                   color=RGBColor(0x55, 0x55, 0x55), space_after_pt=24)

    # 전문
    _add_paragraph(
        doc,
        "본 계약은 아래 당사자 간에 2026년 4월 17일 체결된 「개발용역 계약서」(이하 “본 개발계약”)의 "
        "제12조에 따라, 본 개발계약의 결과물(이하 “본 서비스”)에 대한 운영 단계 유지보수의 범위, "
        "대응 방식, 비용 및 기간에 관한 사항을 정함을 목적으로 하며, 당사자는 신의성실의 원칙에 따라 "
        "본 계약을 이행한다.",
        size_pt=11,
    )

    # 제1조 당사자
    _add_heading(doc, "제1조 (당사자)", level=2)
    _add_party_block(doc, CLIENT)
    _add_paragraph(doc, "", size_pt=4)
    _add_party_block(doc, VENDOR)

    # 제2조 정의
    _add_heading(doc, "제2조 (정의)", level=2)
    for n, t in [
        ("본 서비스", f"{PROJECT_NAME}. 운영 도메인 {SERVICE_DOMAIN} 에서 제공되는 웹·관리자 사이트, "
                    "백엔드 API, 데이터베이스, 작업 스케줄러, 외부 연동(결제·알림톡·소셜 로그인·환율 등)을 포함한다."),
        ("운영 환경", "AWS Lightsail(서울 리전) 위에 도커 컴포즈로 구동되는 7개 서비스(웹·API·DB·Redis·워커·스케줄러·nginx). "
                    "도메인·SSL 인증서·자동 백업 설정을 포함한다."),
        ("무상 하자보수 기간", "본 개발계약 제11조에 따라 최종 검수 완료일로부터 30일 동안 적용되는 무상 보수 기간을 말한다. "
                              "이 기간 동안의 보수는 본 계약과 별개로 무상 처리되며, 별도 청구가 발생하지 않는다."),
        ("유지보수", "무상 하자보수 기간 종료 후 본 서비스의 안정적 운영을 위해 을이 갑에게 제공하는 운영 지원·장애 대응·"
                  "정기 점검·소규모 보완 등 본 계약 제4조에 정한 업무를 말한다."),
    ]:
        _add_paragraph(doc, f"  · “{n}”: {t}", size_pt=10.5)

    # 제3조 계약의 목적
    _add_heading(doc, "제3조 (계약의 목적)", level=2)
    _add_paragraph(
        doc,
        "을은 본 서비스가 운영 환경에서 안정적으로 동작하도록 본 계약에서 정한 범위 내의 유지보수 "
        "업무를 수행하고, 갑은 그 대가를 본 계약에서 정한 바에 따라 지급한다.",
        size_pt=11,
    )

    # 제4조 유지보수 범위
    _add_heading(doc, "제4조 (유지보수 범위)", level=2)
    _add_paragraph(doc, "1. 본 계약에 따른 유지보수 업무에는 다음 각 호가 포함된다.",
                   size_pt=11)
    for t in [
        "1) 본 서비스의 장애·오작동에 대한 원인 분석 및 복구 대응",
        "2) 운영 환경(AWS Lightsail, 도커 컨테이너, nginx, PostgreSQL, Redis)의 상태 점검 및 재기동",
        "3) SSL 인증서 갱신 모니터링 및 도메인 설정 유지",
        "4) 자동 백업 실행 여부 점검 및 백업 데이터 복원 지원",
        "5) 외부 연동 서비스(토스페이먼츠, 알리고 알림톡, 카카오·네이버·구글 OAuth, OpenAI API, 한국수출입은행 환율 API)의 "
        "정책 변경·키 만료에 따른 설정 대응",
        "6) 의존 라이브러리·운영체제·런타임의 보안 패치 적용",
        "7) 운영자 문의에 대한 기술 자문(관리자 사이트 사용법, 데이터 조회·정정 방법 등)",
        "8) 운영 데이터 정정 요청 처리(예: 결제 상태 보정, 잘못 등록된 RAG 자료 삭제 등)",
        "9) 갑이 요청하는 분기별 점검 보고서 작성(요청 시)",
    ]:
        _add_paragraph(doc, f"  {t}", size_pt=10.5, left_indent_cm=0.4)

    _add_paragraph(doc, "2. 다음 각 호는 본 계약의 유지보수 범위에서 제외되며, 별도 견적·별도 계약에 따른다.",
                   size_pt=11)
    for t in [
        "1) 신규 기능의 추가 개발 및 기존 기능의 사양 변경",
        "2) 디자인 개편, 신규 화면 제작, UI/UX 리뉴얼",
        "3) 신규 외부 서비스 연동(예: 추가 결제 수단, 신규 소셜 로그인, 신규 알림 채널 등)",
        "4) 대규모 데이터 마이그레이션 또는 신규 RAG 모델·신규 LLM 모델로의 교체",
        "5) 갑 또는 제3자가 임의로 운영 환경·소스코드를 수정하여 발생한 장애의 복구",
        "6) 갑의 별도 계정·서버·도메인으로의 이관 작업",
        "7) 외부 API·PG사·소셜 로그인 제공사·OS·브라우저의 정책 변경에 따른 대규모 구조 변경",
        "8) 본 서비스 외 다른 시스템과의 연동 또는 데이터 동기화",
    ]:
        _add_paragraph(doc, f"  {t}", size_pt=10.5, left_indent_cm=0.4)

    _add_paragraph(
        doc,
        "3. 본 조 제2항에 해당하는 업무가 필요한 경우, 갑은 을에게 사전에 요청하고, 을은 "
        "5영업일 이내에 작업 범위·소요 기간·비용에 대한 견적을 회신한다. 양 당사자가 견적을 "
        "서면(전자 메일·메신저 등 전자적 기록을 포함한다)으로 합의한 후 작업에 착수한다.",
        size_pt=11,
    )

    # 제5조 대응 시간 (SLA)
    _add_heading(doc, "제5조 (대응 시간 — SLA)", level=2)
    _add_paragraph(doc, "장애·문의 유형에 따른 대응 시간은 다음과 같다.", size_pt=11)

    sla_tbl = doc.add_table(rows=1, cols=4)
    sla_tbl.autofit = False
    sla_tbl.columns[0].width = Cm(2.5)
    sla_tbl.columns[1].width = Cm(5.5)
    sla_tbl.columns[2].width = Cm(3.0)
    sla_tbl.columns[3].width = Cm(5.5)
    hdr = sla_tbl.rows[0].cells
    for i, h in enumerate(["등급", "정의", "1차 회신", "비고"]):
        _set_cell_text(hdr[i], h, bold=True, size_pt=10)
        _shade_cell(hdr[i], "E0E0E0")
    for row in [
        ("긴급",
         "운영 서버 다운, 로그인 불가, 결제 시스템 마비, 데이터 손실 등 서비스 전면 중단",
         "접수 후 24시간 이내",
         "주말·공휴일 포함. 채널 무관 접수 즉시 대응."),
        ("중요",
         "일부 기능 장애, 특정 사용자군의 사용 불가, 알림톡 발송 지연 등 부분 장애",
         "1~2 영업일 이내",
         "원인 파악 후 복구 일정 별도 안내."),
        ("일반",
         "사용법 문의, 데이터 조회·정정 요청, 소규모 설정 변경 등",
         "1~3 영업일 이내",
         "내용에 따라 추가 자료 요청 가능."),
    ]:
        cells = sla_tbl.add_row().cells
        _set_cell_text(cells[0], row[0], bold=True, size_pt=10)
        _set_cell_text(cells[1], row[1], size_pt=10)
        _set_cell_text(cells[2], row[2], size_pt=10)
        _set_cell_text(cells[3], row[3], size_pt=10,
                       color=RGBColor(0x55, 0x55, 0x55))
    _add_paragraph(doc, "", size_pt=4)
    _add_paragraph(
        doc,
        "※ 위 1차 회신 시간은 접수 확인 및 대응 계획 안내까지의 시간이며, 복구 완료 시간을 "
        "보장하는 것은 아니다. 다만 을은 가능한 최단 시간 내 복구를 위해 성실히 협조한다.",
        size_pt=10, color=RGBColor(0x55, 0x55, 0x55),
    )

    # 제6조 접수 채널
    _add_heading(doc, "제6조 (접수 채널)", level=2)
    _add_paragraph(doc, "1. 갑은 다음 채널 중 하나로 유지보수 업무를 접수한다.",
                   size_pt=11)
    for t in [
        f"1) 이메일: {VENDOR['email']}",
        f"2) 휴대전화(문자·통화): {VENDOR['tel']}",
        "3) 양 당사자가 합의한 메신저(카카오톡 등)",
    ]:
        _add_paragraph(doc, f"  {t}", size_pt=10.5, left_indent_cm=0.4)
    _add_paragraph(
        doc,
        "2. 긴급 장애(제5조 제1행)는 가급적 문자 또는 통화를 우선 사용하고, 이메일은 보조로 한다.",
        size_pt=11,
    )
    _add_paragraph(
        doc,
        "3. 을은 접수 즉시 접수 사실을 회신하며, 회신 시각을 1차 회신 SLA의 기산점으로 본다.",
        size_pt=11,
    )

    # 제7조 비용
    _add_heading(doc, "제7조 (유지보수 비용)", level=2)
    _add_paragraph(doc, "1. 본 계약의 유지보수 비용은 다음과 같이 정한다.",
                   size_pt=11)
    for t in [
        "1) 월 정기 비용: 월 [               ]원 (부가가치세 별도). "
        "매월 [       ]일에 갑이 을이 지정한 계좌로 선지급한다.",
        "2) 추가 작업 비용: 본 계약 제4조 제2항(범위 외 작업)에 대해서는 "
        "을의 견적에 따라 별도 청구한다. 통상 시간당 [           ]원 (부가가치세 별도)을 "
        "기준 단가로 한다.",
    ]:
        _add_paragraph(doc, f"  {t}", size_pt=10.5, left_indent_cm=0.4)
    _add_paragraph(
        doc,
        "2. 월 정기 비용에는 본 계약 제4조 제1항의 업무가 포함되며, 추가 청구되지 않는다. "
        "다만, 1회 4시간을 초과하는 장애 복구 작업이나 갑의 명시적 요청에 의한 운영자 교육·"
        "보고서 작성 등은 추가 작업 비용 항목으로 별도 청구할 수 있다.",
        size_pt=11,
    )
    _add_paragraph(
        doc,
        "3. AWS Lightsail 사용료, 도메인 갱신료, 토스페이먼츠 결제 수수료, 알리고 알림톡 발송 비용, "
        "OpenAI API 사용료 등 본 서비스 운영에 따른 외부 비용은 갑이 직접 부담한다.",
        size_pt=11,
    )

    # 제8조 계약 기간 및 갱신
    _add_heading(doc, "제8조 (계약 기간 및 갱신)", level=2)
    _add_paragraph(
        doc,
        "1. 본 계약의 유효 기간은 [       ]년 [   ]월 [   ]일부터 [       ]년 [   ]월 [   ]일까지로 한다. "
        "본 기간은 본 개발계약 제11조에 따른 무상 하자보수 기간(최종 검수 완료일로부터 30일) 종료일 "
        "다음 날부터 개시함을 원칙으로 한다.",
        size_pt=11,
    )
    _add_paragraph(
        doc,
        "2. 갑 또는 을이 계약 만료일 30일 전까지 서면으로 갱신 거절 또는 조건 변경 의사를 "
        "통지하지 아니한 경우, 본 계약은 동일한 조건으로 1년 단위로 자동 갱신된다.",
        size_pt=11,
    )
    _add_paragraph(
        doc,
        "3. 본 계약의 갱신 또는 종료 여부는 만료일 30일 전 양 당사자가 협의한 결과를 우선한다.",
        size_pt=11,
    )

    # 제9조 갑의 의무
    _add_heading(doc, "제9조 (갑의 의무)", level=2)
    for t in [
        "1. 갑은 본 서비스 운영에 필요한 외부 계정(AWS, 토스페이먼츠, 알리고, OAuth, OpenAI 등) 및 "
        "결제 수단을 정상 상태로 유지한다.",
        "2. 갑은 외부 서비스로부터 정책 변경·이용 약관 변경 통지를 받은 경우 지체 없이 을에게 공유한다.",
        "3. 갑은 본 서비스의 소스코드·데이터·운영 환경을 임의로 수정하지 아니하며, 수정이 필요한 "
        "경우 사전에 을과 협의한다.",
        "4. 갑은 본 계약 제7조에 정한 비용을 약정 기일에 지급한다.",
    ]:
        _add_paragraph(doc, t, size_pt=11)

    # 제10조 을의 의무
    _add_heading(doc, "제10조 (을의 의무)", level=2)
    for t in [
        "1. 을은 선량한 관리자의 주의의무로써 본 계약의 유지보수 업무를 수행한다.",
        "2. 을은 본 서비스의 운영 상태를 주기적으로 점검하고, 중대한 위험(저장 공간 부족, 인증서 만료 임박, "
        "외부 키 만료 임박 등)을 사전에 감지하여 갑에게 통지한다.",
        "3. 을은 본 서비스의 보안에 영향을 미치는 중대 취약점이 공개된 경우, 영향 분석 후 "
        "필요한 패치를 본 계약 범위 내에서 적용한다.",
        "4. 을은 본 계약 수행 중 알게 된 갑의 영업·기술 비밀 및 본 서비스 사용자의 개인정보를 "
        "본 계약 목적 외로 사용하거나 제3자에게 누설하지 아니한다.",
    ]:
        _add_paragraph(doc, t, size_pt=11)

    # 제11조 비밀유지 및 개인정보 보호
    _add_heading(doc, "제11조 (비밀유지 및 개인정보 보호)", level=2)
    for t in [
        "1. 갑과 을은 본 계약의 이행 과정에서 알게 된 상대방의 영업상·기술상 비밀 및 본 서비스 "
        "사용자의 개인정보를 본 계약의 목적 외로 사용하거나 제3자에게 누설하지 아니한다.",
        "2. 을은 본 서비스 사용자의 개인정보에 접근할 필요가 있는 경우 갑의 사전 동의를 받은 "
        "범위 내에서만 접근하며, 접근 사실 및 사유를 기록·보관한다.",
        "3. 본 조의 의무는 본 계약 종료 후 3년간 존속한다.",
    ]:
        _add_paragraph(doc, t, size_pt=11)

    # 제12조 계약의 해지
    _add_heading(doc, "제12조 (계약의 해지)", level=2)
    for t in [
        "1. 갑 또는 을은 상대방이 본 계약상 중대한 의무를 위반하고, 서면으로 상당한 기간을 "
        "정하여 시정을 최고하였음에도 이를 이행하지 아니하는 경우 본 계약을 해지할 수 있다.",
        "2. 갑 또는 을은 30일 전 서면 통지로써 본 계약을 임의로 해지할 수 있다. 이 경우 이미 "
        "지급된 정기 비용 중 잔여 기간에 해당하는 금액은 일할 계산하여 정산한다.",
        "3. 본 계약이 종료되더라도 본 개발계약 제9조에 따라 갑에게 이전된 결과물의 권리 귀속에는 "
        "영향이 없다.",
    ]:
        _add_paragraph(doc, t, size_pt=11)

    # 제13조 책임의 제한
    _add_heading(doc, "제13조 (책임의 제한)", level=2)
    for t in [
        "1. 을은 본 계약의 유지보수 업무 수행 중 고의 또는 중대한 과실로 갑에게 손해를 입힌 "
        "경우 그 손해를 배상할 책임을 진다.",
        "2. 다음 각 호의 사유로 인한 장애 또는 서비스 중단에 대해서는 을은 책임을 부담하지 아니한다.",
        "   1) 천재지변, 통신장애, 정전, 화재 등 불가항력적 사유",
        "   2) AWS·토스페이먼츠·알리고·카카오·네이버·구글·OpenAI 등 외부 서비스 제공사의 장애 "
        "또는 정책 변경",
        "   3) 갑 또는 제3자의 임의 수정·설정 변경으로 발생한 장애",
        "   4) 외부의 해킹·서비스 거부 공격 등에 의한 침해 사고(다만 을의 부주의에 의한 경우는 제외)",
        "3. 을의 본 계약상 손해배상 책임은 직전 12개월간 갑이 을에게 실제로 지급한 정기 비용의 "
        "총액을 한도로 한다.",
    ]:
        _add_paragraph(doc, t, size_pt=11)

    # 제14조 준거법 및 관할
    _add_heading(doc, "제14조 (준거법 및 관할)", level=2)
    for t in [
        "1. 본 계약의 해석 및 적용에 관하여는 대한민국 법령을 준거법으로 한다.",
        "2. 본 계약과 관련하여 분쟁이 발생하는 경우, 갑의 주소지를 관할하는 법원 또는 "
        "당사자가 합의한 법원을 제1심 전속적 관할 법원으로 한다.",
    ]:
        _add_paragraph(doc, t, size_pt=11)

    # 제15조 기타
    _add_heading(doc, "제15조 (기타)", level=2)
    for t in [
        "1. 본 계약에 정하지 아니한 사항 또는 본 계약의 해석에 다툼이 있는 경우, 양 당사자는 "
        "신의성실의 원칙에 따라 협의하여 정한다.",
        "2. 본 계약은 서면 서명 외에 전자문서 및 전자서명의 방식으로 체결할 수 있으며, "
        "이 경우에도 서면 계약과 동일한 법적 효력을 가진다.",
        "3. 본 계약과 본 개발계약의 내용이 상충하는 경우, 본 계약의 운영·유지보수에 관한 사항은 "
        "본 계약을 우선하여 적용한다.",
    ]:
        _add_paragraph(doc, t, size_pt=11)

    _add_signature_block(doc)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


# ──────────────────────────────────────────────────────────────────────────────
# 문서 2 — 하자보수 범위·기간 안내서
# ──────────────────────────────────────────────────────────────────────────────

def build_defect_warranty_notice(output_path: Path) -> None:
    doc = Document()
    _set_margins(doc)

    # 표지
    _add_paragraph(doc, "Denvia", bold=True, size_pt=22,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "치과 보험청구·데스크 행정 AI 어시스턴트", size_pt=11,
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   color=RGBColor(0x55, 0x55, 0x55))
    _add_paragraph(doc, "", size_pt=10)
    _add_paragraph(doc, "하자보수 범위·기간 안내서", bold=True, size_pt=22,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "— 무상 하자보수 30일 운영 기준 —", size_pt=11,
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   color=RGBColor(0x55, 0x55, 0x55))
    _add_paragraph(doc, "작성일: 2026년 6월 5일", size_pt=11,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=20)

    # 1. 문서 목적
    _add_heading(doc, "1. 문서 목적", level=1)
    _add_paragraph(
        doc,
        "본 안내서는 2026년 4월 17일 체결된 「개발용역 계약서」(이하 “본 개발계약”) 제11조 및 "
        "제11조의2에 정한 무상 하자보수의 범위·기간·신고 방법·대응 절차를 고객(갑)이 쉽게 "
        "파악할 수 있도록 정리한 문서입니다. 본 안내서의 내용이 본 개발계약과 상충하는 경우 "
        "본 개발계약의 내용이 우선합니다.",
        size_pt=11,
    )

    # 2. 핵심 요약
    _add_heading(doc, "2. 핵심 요약 (한눈에 보기)", level=1)

    summary_tbl = doc.add_table(rows=0, cols=2)
    summary_tbl.autofit = False
    summary_tbl.columns[0].width = Cm(4.5)
    summary_tbl.columns[1].width = Cm(12.0)
    for label, value in [
        ("대상 서비스",
         f"{PROJECT_NAME}\n운영 도메인: {SERVICE_DOMAIN}"),
        ("무상 보수 기간",
         "최종 검수 완료일로부터 30일\n(예: 2026년 6월 5일 검수 완료 시 → 2026년 7월 5일까지)"),
        ("무상 보수 범위",
         "본 개발계약에서 합의된 기능명세 내에서 발견된 오작동·오류·버그의 수정"),
        ("무상 보수 제외",
         "기능 추가·디자인 변경·외부 정책 변경 대응 등 (본 안내서 5항 참조)"),
        ("신고 채널",
         f"이메일: {VENDOR['email']}\n휴대전화: {VENDOR['tel']}"),
        ("긴급 장애 대응",
         "접수 후 24시간 이내 1차 회신 (주말·공휴일 포함)"),
        ("기간 종료 이후",
         "별도 「유지보수 계약」으로 전환 (월 정기 비용 + 추가 작업 별도 견적)"),
    ]:
        row = summary_tbl.add_row().cells
        _set_cell_text(row[0], label, bold=True, size_pt=10)
        _shade_cell(row[0], "F2F2F2")
        _set_cell_text(row[1], value, size_pt=10)
    _add_paragraph(doc, "", size_pt=6)

    # 3. 무상 하자보수란
    _add_heading(doc, "3. 무상 하자보수란", level=1)
    _add_paragraph(
        doc,
        "납품 후 일정 기간 동안, 본 개발계약에서 합의된 기능명세를 기준으로 “정상이라고 약속한 "
        "동작이 일어나지 않는 경우”에 한해 추가 비용 없이 수정하는 것을 말합니다. 본 안내서에서 "
        "“하자”라 함은 다음을 의미합니다.",
        size_pt=11,
    )
    for t in [
        "  · 기능명세상 동작해야 하는 화면·버튼·기능이 동작하지 않는 경우",
        "  · 기능명세상 명시된 결과와 다른 결과가 출력되는 경우",
        "  · 통상 사용 환경(최신 버전 크롬·사파리·엣지·삼성 인터넷 / 일반 모바일·태블릿·PC)에서 "
        "재현되는 화면 깨짐 또는 오작동",
        "  · 본 서비스 자체의 결함으로 인한 데이터 손실 또는 결제 오류",
    ]:
        _add_paragraph(doc, t, size_pt=11)

    # 4. 무상 보수 기간
    _add_heading(doc, "4. 무상 보수 기간", level=1)
    _add_paragraph(
        doc,
        "무상 하자보수 기간은 본 개발계약 제11조 제1항에 따라 최종 검수 완료일로부터 30일 "
        "(달력일 기준)입니다. 검수 완료일의 기준은 본 개발계약 제8조 제2항·제3항에 따릅니다.",
        size_pt=11,
    )

    period_tbl = doc.add_table(rows=0, cols=2)
    period_tbl.autofit = False
    period_tbl.columns[0].width = Cm(5.0)
    period_tbl.columns[1].width = Cm(11.5)
    for label, value in [
        ("검수 요청일",
         "을이 갑에게 검수 요청 통지를 발송한 날"),
        ("검수 완료일",
         "갑이 ‘검수 합격’을 통지한 날. 또는 검수 요청일로부터 5영업일 내 별도 통지가 없으면 "
         "본 개발계약 제8조 제3항에 따라 검수에 합격한 것으로 본다 (이 경우 합격 간주일을 "
         "검수 완료일로 본다)."),
        ("무상 보수 시작일",
         "검수 완료일 다음 날"),
        ("무상 보수 종료일",
         "시작일로부터 30일째 되는 날의 자정"),
        ("종료일 예시 (참고)",
         "2026년 6월 5일 검수 완료 → 2026년 6월 6일 ~ 2026년 7월 5일까지 무상"),
    ]:
        row = period_tbl.add_row().cells
        _set_cell_text(row[0], label, bold=True, size_pt=10)
        _shade_cell(row[0], "F2F2F2")
        _set_cell_text(row[1], value, size_pt=10)
    _add_paragraph(doc, "", size_pt=6)

    # 5. 무상 보수 대상 (포함)
    _add_heading(doc, "5. 무상 보수 대상 — 포함되는 항목", level=1)
    _add_paragraph(
        doc,
        "다음과 같은 사항은 무상 하자보수 기간 내에 신고하시면 추가 비용 없이 수정합니다.",
        size_pt=11,
    )

    include_tbl = doc.add_table(rows=1, cols=2)
    include_tbl.autofit = False
    include_tbl.columns[0].width = Cm(5.5)
    include_tbl.columns[1].width = Cm(11.0)
    hdr = include_tbl.rows[0].cells
    for i, h in enumerate(["분류", "예시"]):
        _set_cell_text(hdr[i], h, bold=True, size_pt=10)
        _shade_cell(hdr[i], "E0E0E0")
    for cat, ex in [
        ("사용자 화면 (Q&A 채팅·구독·마이페이지·받은 쪽지함·1:1 문의 등)",
         "버튼 클릭 시 오류 메시지가 뜨고 동작하지 않는 경우 / 로그인 후 자동 전환되지 않는 경우 / "
         "결제 카드 등록이 완료되었으나 Pro 권한이 적용되지 않는 경우"),
        ("관리자 화면 (대시보드·회원·통계·매출·RAG·CS·예산통제·관리자 관리 등)",
         "검색 결과가 표시되지 않는 경우 / 권한 변경 저장이 적용되지 않는 경우 / "
         "통계 그래프의 수치가 명백히 잘못 계산되는 경우"),
        ("결제·환불",
         "정기 결제가 약정일에 실행되지 않는 경우 / 환불 처리는 완료되었으나 사용자 알림톡이 "
         "발송되지 않는 경우 / 결제 실패 후 자동 재시도가 누락되는 경우"),
        ("알림톡 발송",
         "발송 트리거 조건이 충족되었으나 알림톡이 발송되지 않는 경우 / 발송된 알림톡의 "
         "치환 변수가 누락되거나 잘못 표기되는 경우"),
        ("AI 응답 엔진",
         "정상 질문에 대해 답변이 생성되지 않고 오류로 종료되는 경우 / RAG 자료 재빌드 후 "
         "관리자 화면에서 반영이 확인되지 않는 경우"),
        ("운영 환경",
         "서비스 자체의 결함으로 발생한 운영 서버 다운 또는 접속 불가 / 정상 동작하던 "
         "외부 연동이 본 서비스 측 결함으로 갑자기 끊기는 경우"),
    ]:
        cells = include_tbl.add_row().cells
        _set_cell_text(cells[0], cat, bold=True, size_pt=10)
        _set_cell_text(cells[1], ex, size_pt=10)
    _add_paragraph(doc, "", size_pt=6)

    # 6. 무상 보수 제외 (별도 유상)
    _add_heading(doc, "6. 무상 보수 제외 — 별도 유상 작업", level=1)
    _add_paragraph(
        doc,
        "다음 사항은 본 개발계약 제11조 제3항에 따라 무상 보수 대상에서 제외되며, "
        "별도 견적·유상 작업으로 진행됩니다.",
        size_pt=11,
    )

    exclude_tbl = doc.add_table(rows=1, cols=2)
    exclude_tbl.autofit = False
    exclude_tbl.columns[0].width = Cm(5.5)
    exclude_tbl.columns[1].width = Cm(11.0)
    hdr = exclude_tbl.rows[0].cells
    for i, h in enumerate(["분류", "예시"]):
        _set_cell_text(hdr[i], h, bold=True, size_pt=10)
        _shade_cell(hdr[i], "E0E0E0")
    for cat, ex in [
        ("기능의 추가 또는 사양 변경",
         "신규 결제 수단 추가 / 신규 메뉴 추가 / 기존 통계 항목의 계산 방식 변경 / "
         "RAG 자료 양식 변경 / 답변 길이·톤 정책 변경 등"),
        ("디자인 수정",
         "색상·폰트·레이아웃 변경 / 로고·이미지 교체 / 화면 리뉴얼 / 신규 화면 디자인"),
        ("정책 변경에 따른 수정",
         "구독 가격·플랜 구조 변경 / 환불 정책 변경 / 무료 한도 변경 / 이상탐지 기준 변경 / "
         "운영 시간 변경 등"),
        ("외부 서비스 정책 변경 대응",
         "토스페이먼츠·알리고·카카오·네이버·구글·OpenAI·한국수출입은행 환율 API 등 외부 "
         "서비스 제공사의 약관·API 변경에 따른 구조 수정"),
        ("브라우저·OS 정책 변경 대응",
         "iOS·Android·Windows 등의 보안 정책 변경, 신규 브라우저 출시 등에 따른 호환성 작업"),
        ("갑 또는 제3자의 임의 수정",
         "본 서비스 소스코드·데이터베이스·운영 환경을 을과 협의 없이 수정하여 발생한 장애의 복구"),
        ("운영 데이터 입력·정정·이관",
         "RAG 자료 신규 대량 등록 / 회원 정보 수기 정정 / 외부 시스템으로의 데이터 이관"),
        ("운영자 교육·매뉴얼 제작",
         "별도 교육 세션 진행 / 신규 매뉴얼 작성 / 영상 자료 제작"),
    ]:
        cells = exclude_tbl.add_row().cells
        _set_cell_text(cells[0], cat, bold=True, size_pt=10)
        _set_cell_text(cells[1], ex, size_pt=10)
    _add_paragraph(doc, "", size_pt=6)

    # 7. 신고 방법
    _add_heading(doc, "7. 신고 방법", level=1)
    _add_paragraph(
        doc,
        "하자가 발견된 경우 다음 채널 중 하나로 신고해 주십시오. 가능한 한 빨리 대응할 수 있도록 "
        "재현 정보를 함께 보내 주시면 도움이 됩니다.",
        size_pt=11,
    )

    contact_tbl = doc.add_table(rows=0, cols=2)
    contact_tbl.autofit = False
    contact_tbl.columns[0].width = Cm(4.0)
    contact_tbl.columns[1].width = Cm(12.5)
    for label, value in [
        ("이메일", VENDOR["email"]),
        ("휴대전화 (문자·통화)", VENDOR["tel"]),
        ("메신저", "양 당사자가 합의한 카카오톡 등"),
    ]:
        row = contact_tbl.add_row().cells
        _set_cell_text(row[0], label, bold=True, size_pt=10)
        _shade_cell(row[0], "F2F2F2")
        _set_cell_text(row[1], value, size_pt=10)
    _add_paragraph(doc, "", size_pt=6)

    _add_paragraph(doc, "신고 시 함께 보내 주시면 좋은 정보", bold=True, size_pt=11)
    for t in [
        "  · 발생 시각 (대략 몇 시쯤)",
        "  · 어떤 화면에서 어떤 동작을 했는지 (예: ‘마이페이지에서 구독 해지 버튼을 눌렀더니’)",
        "  · 어떤 결과가 나왔는지 (예: ‘오류 팝업이 떴다 / 화면이 흰색으로 변했다’)",
        "  · 사용한 기기 (모바일·태블릿·PC) 와 브라우저 (크롬·사파리 등)",
        "  · 영향을 받은 회원 정보 (이름·이메일 등) — 결제 관련 문의의 경우",
        "  · 가능하다면 화면 캡쳐 또는 영상",
    ]:
        _add_paragraph(doc, t, size_pt=10.5)

    # 8. 대응 절차 및 시간
    _add_heading(doc, "8. 대응 절차 및 시간", level=1)
    _add_paragraph(
        doc,
        "신고 접수 후 다음 절차로 진행됩니다. 1차 회신 시간은 접수 확인 및 대응 계획 안내까지의 "
        "시간이며, 복구 완료 시간을 보장하는 것은 아닙니다.",
        size_pt=11,
    )

    step_tbl = doc.add_table(rows=1, cols=3)
    step_tbl.autofit = False
    step_tbl.columns[0].width = Cm(2.0)
    step_tbl.columns[1].width = Cm(4.5)
    step_tbl.columns[2].width = Cm(10.0)
    hdr = step_tbl.rows[0].cells
    for i, h in enumerate(["단계", "구분", "내용"]):
        _set_cell_text(hdr[i], h, bold=True, size_pt=10)
        _shade_cell(hdr[i], "E0E0E0")
    for n, k, v in [
        ("1", "접수 회신",
         "신고 접수 후 가능한 빠르게 접수 사실을 회신 (긴급 24시간 이내 / 일반 1~3 영업일 이내)"),
        ("2", "원인 분석",
         "재현 시도, 로그·DB·외부 응답 확인. 무상 보수 대상인지 여부를 판별하여 갑에게 통지"),
        ("3", "대응",
         "무상 대상: 즉시 수정·배포\n무상 대상이 아닌 경우: 견적 안내 후 갑의 승인 시 작업 진행"),
        ("4", "결과 안내",
         "수정 내용·반영 시각·재현 방법을 갑에게 통지하고 정상 동작 여부를 확인 요청"),
    ]:
        cells = step_tbl.add_row().cells
        _set_cell_text(cells[0], n, bold=True, size_pt=10,
                       color=RGBColor(0x10, 0x80, 0x30))
        _set_cell_text(cells[1], k, bold=True, size_pt=10)
        _set_cell_text(cells[2], v, size_pt=10)
    _add_paragraph(doc, "", size_pt=6)

    # 9. 기간 종료 후 (제11조의2)
    _add_heading(doc, "9. 무상 보수 기간 종료 후의 대응 (본 개발계약 제11조의2)", level=1)
    _add_paragraph(
        doc,
        "무상 하자보수 기간이 종료된 이후에 발생한 장애 중, 본 서비스 자체의 결함이 명백한 "
        "경우에는 을이 적극적으로 보완에 협조합니다. 다만 다음 각 호에 해당하는 사유로 발생한 "
        "장애는 무상 보완 의무가 없으며, 별도 유지보수 계약에 따른 유상 작업으로 진행됩니다.",
        size_pt=11,
    )
    for t in [
        "  · 갑 또는 제3자에 의한 소스코드·운영 환경 수정으로 발생한 장애",
        "  · 서버·호스팅·클라우드·데이터베이스·운영체제·브라우저·외부 API·PG사 등 운영 환경의 "
        "변경 또는 외부 서비스 장애로 발생한 장애",
        "  · 서비스 정책·운영 방식·기능 요구사항 변경에 따른 수정",
        "  · 통신장애·천재지변 등 불가항력적 사유",
    ]:
        _add_paragraph(doc, t, size_pt=11)

    _add_paragraph(
        doc,
        "기간 종료 이후의 안정적 운영을 위해 별도 「유지보수 계약」 체결을 권장드리며, "
        "유지보수 계약은 본 안내서와 함께 전달된 「Denvia 유지보수 계약서」를 참고해 주십시오.",
        size_pt=11,
    )

    # 10. 문의처
    _add_heading(doc, "10. 문의처", level=1)

    contact_summary = doc.add_table(rows=0, cols=2)
    contact_summary.autofit = False
    contact_summary.columns[0].width = Cm(4.0)
    contact_summary.columns[1].width = Cm(12.5)
    for label, value in [
        ("수행사", VENDOR["company"]),
        ("담당자", VENDOR["rep"]),
        ("이메일", VENDOR["email"]),
        ("휴대전화", VENDOR["tel"]),
    ]:
        row = contact_summary.add_row().cells
        _set_cell_text(row[0], label, bold=True, size_pt=10)
        _shade_cell(row[0], "F2F2F2")
        _set_cell_text(row[1], value, size_pt=10)
    _add_paragraph(doc, "", size_pt=8)

    _add_horizontal_rule(doc)
    _add_paragraph(
        doc,
        "본 안내서는 본 개발계약의 이행을 보조하는 문서이며, 법적 효력은 본 개발계약에 따릅니다.",
        size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER,
        color=RGBColor(0x55, 0x55, 0x55),
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


# ──────────────────────────────────────────────────────────────────────────────
# 실행
# ──────────────────────────────────────────────────────────────────────────────

def main() -> None:
    project_root = Path(__file__).resolve().parent.parent
    out_dir = project_root / "자료"

    contract_path = out_dir / "Denvia_유지보수_계약서.docx"
    notice_path = out_dir / "Denvia_하자보수_범위_기간_안내서.docx"

    build_maintenance_contract(contract_path)
    print(f"[OK] {contract_path}")

    build_defect_warranty_notice(notice_path)
    print(f"[OK] {notice_path}")


if __name__ == "__main__":
    main()
