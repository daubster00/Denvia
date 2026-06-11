"""VSCode로 운영 서버(denvia-prod-01)에 SSH 접속하는 방법 안내서 빌더.

2026-06-05 — 비개발자도 따라할 수 있도록 비유 위주, 단계별 그림 포함.

산출물:
    자료/Denvia_VSCode_SSH_접속_안내서.docx

CLAUDE.md "고객 전달용 문서 작성 규칙" 준수 — SDT/문서 보호/잠금 일체 금지.
어떤 워드 버전(MS Word, 한컴, Google Docs, LibreOffice)에서도 자유 편집 가능.

실행:
    py -3 scripts/build_vscode_ssh_guide.py
"""
from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parent.parent
SHOT_CLEAN = ROOT / "outputs" / "vscode_ssh" / "clean"
ILLUS_DIR = ROOT / "outputs" / "vscode_ssh" / "illus"
ILLUS_DIR.mkdir(parents=True, exist_ok=True)

DOC_TITLE = "VSCode 원격 접속 안내서"
DOC_SUBTITLE = "Denvia 운영 서버 (denvia-prod-01) 접속용 가이드"
DOC_DATE = "작성일: 2026년 6월 5일"

KOREAN_FONT = "맑은 고딕"
MONO_FONT = "Consolas"


# ──────────────────────────────────────────────────────────────────────────────
# docx 헬퍼
# ──────────────────────────────────────────────────────────────────────────────

def _apply_font(run, *, font: str = KOREAN_FONT, bold: bool = False,
                size_pt: float | None = None, color: RGBColor | None = None) -> None:
    run.font.name = font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:ascii"), font)
    if bold:
        run.bold = True
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color


def _para(doc, text, *, bold=False, size_pt=11, align=None,
          color=None, font=KOREAN_FONT, space_after_pt=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_after_pt is not None:
        p.paragraph_format.space_after = Pt(space_after_pt)
    r = p.add_run(text)
    _apply_font(r, font=font, bold=bold, size_pt=size_pt, color=color)
    return p


def _heading(doc, text, level=1):
    sizes = {1: 20, 2: 15, 3: 12}
    spaces = {1: 22, 2: 16, 3: 10}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(spaces[level])
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    _apply_font(r, bold=True, size_pt=sizes[level])
    # 제목 1 아래 가는 밑줄
    if level == 1:
        _hr(doc)


def _hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "BBBBBB")
    pBdr.append(bot)
    pPr.append(pBdr)


def _shade(cell, hex_):
    pr = cell._tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), hex_)
    pr.append(s)


def _callout(doc, lines: list[str], *, fill="FFF8DC",
             title: str | None = None, title_color="996600"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Cm(16)
    c = tbl.cell(0, 0)
    c.width = Cm(16)
    _shade(c, fill)
    first_done = False
    if title:
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        _apply_font(r, bold=True, size_pt=11, color=RGBColor.from_string(title_color))
        first_done = True
    for line in lines:
        p = c.paragraphs[0] if not first_done else c.add_paragraph()
        first_done = True
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        _apply_font(r, size_pt=10.5)
    doc.add_paragraph()


def _code_block(doc, code: str, *, fill="F2F2F2"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Cm(16)
    c = tbl.cell(0, 0)
    c.width = Cm(16)
    _shade(c, fill)
    first = True
    for line in code.split("\n"):
        p = c.paragraphs[0] if first else c.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line if line else " ")
        _apply_font(r, font=MONO_FONT, size_pt=10)
    doc.add_paragraph()


def _bullet(doc, text, *, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.0 * level)
    r = p.add_run(text)
    _apply_font(r, size_pt=11)


def _kv_table(doc, rows: list[tuple[str, str]], *,
              label_w=4.0, value_w=12.0, label_fill="F2F2F2"):
    tbl = doc.add_table(rows=0, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Cm(label_w)
    tbl.columns[1].width = Cm(value_w)
    for label, value in rows:
        row = tbl.add_row().cells
        row[0].text = ""; row[1].text = ""
        row[0].width = Cm(label_w); row[1].width = Cm(value_w)
        for c in row:
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _shade(row[0], label_fill)
        p0 = row[0].paragraphs[0]
        r0 = p0.add_run(label)
        _apply_font(r0, bold=True, size_pt=10.5)
        first = True
        for line in value.split("\n"):
            p = row[1].paragraphs[0] if first else row[1].add_paragraph()
            first = False
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(line)
            _apply_font(r, size_pt=10.5)


def _figure(doc, path: Path, *, caption: str, width_cm: float = 15.0):
    """그림 + 캡션."""
    if not path.exists():
        _para(doc, f"[그림 누락: {path.name}]", color=RGBColor(0xCC, 0x00, 0x00))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    cr = cap.add_run(caption)
    _apply_font(cr, size_pt=10, color=RGBColor(0x66, 0x66, 0x66))


# ──────────────────────────────────────────────────────────────────────────────
# PIL 합성 일러스트
# ──────────────────────────────────────────────────────────────────────────────

VS_BG = (30, 30, 30)
VS_SIDE = (37, 37, 38)
VS_PANEL = (45, 45, 48)
VS_TEXT = (220, 220, 220)
VS_DIM = (140, 140, 140)
VS_ACCENT = (0, 122, 204)
VS_GREEN = (35, 134, 54)
HL = (255, 90, 90)


def _font(size: int, *, bold: bool = False) -> ImageFont.ImageFont:
    name = "malgunbd.ttf" if bold else "malgun.ttf"
    p = f"C:/Windows/Fonts/{name}"
    if Path(p).exists():
        return ImageFont.truetype(p, size)
    return ImageFont.load_default()


def _circle_arrow(d: ImageDraw.ImageDraw, x: int, y: int, label: str,
                  *, r: int = 16):
    d.ellipse((x - r, y - r, x + r, y + r), outline=HL, width=3)
    f = _font(15, bold=True)
    bbox = d.textbbox((0, 0), label, font=f)
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    pad = 6
    lx, ly = x + r + 10, y - th // 2 - pad
    d.rectangle((lx - pad, ly - pad, lx + tw + pad, ly + th + pad),
                fill=(50, 50, 50), outline=HL, width=2)
    d.text((lx, ly), label, font=f, fill=(255, 255, 255))


def make_extensions_illus() -> Path:
    """Extensions 패널에서 'Remote - SSH' 검색 결과 일러스트."""
    W, H = 900, 460
    img = Image.new("RGB", (W, H), VS_BG)
    d = ImageDraw.Draw(img)
    # 좌측 활동바
    d.rectangle((0, 0, 48, H), fill=VS_SIDE)
    # Extensions 아이콘 자리 (네번째)
    cx, cy = 24, 150
    d.rectangle((cx - 10, cy - 10, cx + 10, cy + 10), outline=VS_TEXT, width=2)
    d.rectangle((cx - 5, cy - 5, cx + 5, cy + 5), outline=VS_TEXT, width=1)
    # Extensions 사이드바 패널
    d.rectangle((48, 0, 330, H), fill=VS_PANEL)
    f_lbl = _font(13, bold=True)
    d.text((64, 14), "EXTENSIONS: MARKETPLACE", font=f_lbl, fill=VS_DIM)
    # 검색창
    d.rectangle((64, 42, 312, 70), fill=(60, 60, 60), outline=(80, 80, 80))
    f_in = _font(14)
    d.text((74, 48), "Remote - SSH", font=f_in, fill=VS_TEXT)
    # 결과 카드 1: Remote - SSH (Microsoft)
    d.rectangle((64, 84, 312, 184), fill=(50, 50, 52))
    d.rectangle((76, 96, 116, 136), fill=VS_ACCENT)
    d.text((86, 104), ">_", font=_font(20, bold=True), fill=(255, 255, 255))
    d.text((128, 96), "Remote - SSH", font=_font(15, bold=True), fill=VS_TEXT)
    d.text((128, 118), "Open any folder on a remote", font=_font(12), fill=VS_DIM)
    d.text((128, 134), "machine using SSH.", font=_font(12), fill=VS_DIM)
    d.text((128, 154), "Microsoft", font=_font(11), fill=VS_DIM)
    # Install 버튼
    d.rectangle((220, 152, 296, 174), fill=VS_ACCENT)
    d.text((240, 156), "Install", font=_font(12, bold=True), fill=(255, 255, 255))
    # 강조: Install 버튼
    _circle_arrow(d, x=258, y=163, label="여기 클릭")

    # 본문 영역
    d.text((364, 220), "이렇게 'Remote - SSH'를 찾아 Install을 누르면",
           font=_font(15), fill=VS_TEXT)
    d.text((364, 248), "한 번에 셋업이 끝납니다 (필요 시 자동 재시작).",
           font=_font(15), fill=VS_TEXT)

    out = ILLUS_DIR / "step_extensions_install.png"
    img.save(out, "PNG")
    return out


def make_connected_status_illus() -> Path:
    """접속 성공 후 좌측 하단 상태바 일러스트."""
    W, H = 900, 220
    img = Image.new("RGB", (W, H), VS_BG)
    d = ImageDraw.Draw(img)
    # VSCode 좌측 상태바 한 줄
    bar_y = 140
    d.rectangle((0, bar_y, W, bar_y + 28), fill=VS_GREEN)
    # SSH 배지
    d.rectangle((8, bar_y + 3, 220, bar_y + 25), fill=(20, 90, 40))
    d.text((20, bar_y + 6), ">< SSH: denvia-prod-01", font=_font(13, bold=True),
           fill=(255, 255, 255))
    # 나머지는 어두운 톤
    d.text((240, bar_y + 6), "main",
           font=_font(13), fill=(220, 220, 220))
    d.text((310, bar_y + 6), "0 ⚠ 0",
           font=_font(13), fill=(220, 220, 220))

    # 강조
    _circle_arrow(d, x=120, y=bar_y + 14, label="여기 표시되면 접속 성공")
    # 안내 문구
    d.text((40, 50), "화면 가장 아래 좌측 모서리에 위 모양의 초록 배지가 뜨면",
           font=_font(16, bold=True), fill=VS_TEXT)
    d.text((40, 78), "Denvia 운영 서버 안에 들어와 있는 상태입니다.",
           font=_font(16), fill=VS_TEXT)
    out = ILLUS_DIR / "step_connected_status.png"
    img.save(out, "PNG")
    return out


def make_remote_folder_picker_illus() -> Path:
    """원격 폴더 선택 — 상단 입력창에 /opt/denvia 입력."""
    W, H = 900, 360
    img = Image.new("RGB", (W, H), VS_BG)
    d = ImageDraw.Draw(img)
    # 상단 입력 패널 (Quick Pick 스타일)
    d.rectangle((100, 30, 800, 220), fill=VS_PANEL, outline=(60, 60, 60))
    # 입력창
    d.rectangle((116, 46, 784, 76), fill=(60, 60, 60))
    d.text((128, 52), "/opt/denvia", font=_font(15), fill=VS_TEXT)
    # 안내 텍스트
    d.text((116, 88), "서버 안의 폴더 경로를 입력하세요",
           font=_font(13), fill=VS_DIM)
    # OK 버튼
    d.rectangle((664, 178, 784, 206), fill=VS_ACCENT)
    d.text((704, 184), "OK", font=_font(14, bold=True), fill=(255, 255, 255))

    _circle_arrow(d, x=200, y=62, label="/opt/denvia 입력 후 Enter")
    # 하단 설명
    d.text((40, 260), "Denvia 서버의 소스 코드는 항상 /opt/denvia 폴더에 있습니다.",
           font=_font(15), fill=VS_TEXT)
    d.text((40, 288), "이 경로를 입력하면 서버 안 파일을 바로 열 수 있습니다.",
           font=_font(15), fill=VS_TEXT)
    out = ILLUS_DIR / "step_folder_picker.png"
    img.save(out, "PNG")
    return out


def make_concept_diagram() -> Path:
    """1장: 원격 접속 개념도 (내 노트북 ↔ SSH ↔ 서버)."""
    W, H = 900, 340
    img = Image.new("RGB", (W, H), (250, 250, 250))
    d = ImageDraw.Draw(img)
    # 좌측 박스 — 내 노트북
    d.rectangle((40, 80, 280, 240), fill=(255, 255, 255), outline=(120, 120, 120), width=2)
    d.text((90, 100), "내 노트북", font=_font(20, bold=True), fill=(40, 40, 40))
    d.text((70, 140), "VSCode 실행", font=_font(14), fill=(80, 80, 80))
    d.text((70, 162), "+ 개인키(.pem)", font=_font(14), fill=(80, 80, 80))
    d.text((70, 184), "+ ssh config", font=_font(14), fill=(80, 80, 80))
    # 우측 박스 — 서버
    d.rectangle((620, 80, 860, 240), fill=(255, 255, 255), outline=(120, 120, 120), width=2)
    d.text((680, 100), "운영 서버", font=_font(20, bold=True), fill=(40, 40, 40))
    d.text((650, 140), "denvia-prod-01", font=_font(14, bold=True), fill=VS_ACCENT)
    d.text((650, 162), "(AWS 안에 있음)", font=_font(13), fill=(80, 80, 80))
    d.text((650, 184), "/opt/denvia/", font=_font(13), fill=(80, 80, 80))
    # 중간 화살표 — SSH 터널
    d.line((280, 160, 620, 160), fill=VS_ACCENT, width=4)
    d.polygon([(610, 152), (625, 160), (610, 168)], fill=VS_ACCENT)
    d.polygon([(285, 152), (270, 160), (285, 168)], fill=VS_ACCENT)
    d.text((370, 120), "SSH 암호화 연결", font=_font(16, bold=True), fill=VS_ACCENT)
    d.text((370, 175), "(자물쇠+열쇠 방식 — 키가 맞으면 통과)",
           font=_font(12), fill=(120, 120, 120))
    # 상단 안내
    d.text((40, 30), "VSCode 원격 접속은 이 그림 한 장으로 끝납니다.",
           font=_font(17, bold=True), fill=(40, 40, 40))
    # 하단 안내
    d.text((40, 280), "왼쪽 내 노트북에서 VSCode를 열면 마치 서버 안 파일을 "
                       "직접 편집하는 것처럼 보이지만,", font=_font(13), fill=(80, 80, 80))
    d.text((40, 302), "실제 작업은 모두 오른쪽 서버에서 일어나고 결과만 화면에 비춰주는 구조입니다.",
           font=_font(13), fill=(80, 80, 80))
    out = ILLUS_DIR / "concept_diagram.png"
    img.save(out, "PNG")
    return out


# ──────────────────────────────────────────────────────────────────────────────
# GitHub 챕터 일러스트
# ──────────────────────────────────────────────────────────────────────────────

def make_github_three_places_diagram() -> Path:
    """3장: 내 PC / GitHub / 운영서버 — 자동·수동 관계 표시."""
    W, H = 980, 380
    img = Image.new("RGB", (W, H), (250, 250, 250))
    d = ImageDraw.Draw(img)
    # 박스 3개
    boxes = [
        (40, 100, 290, 260, "내 PC", "VSCode + Git",
         "코드를 직접 수정하는 곳"),
        (370, 100, 620, 260, "GitHub", "(내 저장소)",
         "인터넷에 있는 코드 창고"),
        (700, 100, 950, 260, "운영 서버", "denvia-prod-01",
         "고객이 실제 접속하는 사이트"),
    ]
    for x1, y1, x2, y2, title, sub, desc in boxes:
        d.rectangle((x1, y1, x2, y2), fill=(255, 255, 255),
                    outline=(120, 120, 120), width=2)
        d.text((x1 + 20, y1 + 14), title,
               font=_font(19, bold=True), fill=(40, 40, 40))
        d.text((x1 + 20, y1 + 50), sub,
               font=_font(13, bold=True), fill=VS_ACCENT)
        d.text((x1 + 20, y1 + 90), desc,
               font=_font(12), fill=(80, 80, 80))
    # 화살표 1: 내PC → GitHub (자동·간편, 파란색)
    d.line((290, 180, 370, 180), fill=VS_ACCENT, width=4)
    d.polygon([(360, 172), (375, 180), (360, 188)], fill=VS_ACCENT)
    d.text((295, 145), "VSCode 버튼 한 번",
           font=_font(12, bold=True), fill=VS_ACCENT)
    d.text((295, 195), "(자동·쉬움)",
           font=_font(11), fill=(80, 80, 80))
    # 화살표 2: GitHub → 서버 (수동, 빨간색)
    d.line((620, 180, 700, 180), fill=HL, width=4)
    d.polygon([(690, 172), (705, 180), (690, 188)], fill=HL)
    d.text((625, 145), "사람이 직접",
           font=_font(12, bold=True), fill=HL)
    d.text((625, 195), "(SSH 접속 + 명령어)",
           font=_font(11), fill=(80, 80, 80))
    # 상단 안내
    d.text((40, 30), "코드는 3곳에 따로 존재합니다 — 자동으로 연결되는 곳은 한 곳뿐",
           font=_font(17, bold=True), fill=(40, 40, 40))
    # 하단 안내
    d.text((40, 290), "왼쪽 두 칸은 VSCode 클릭 한 번이면 동기화되지만,",
           font=_font(13), fill=(80, 80, 80))
    d.text((40, 314), "오른쪽 운영 서버는 별개입니다 — 5장(원격 접속) 절차로 사람이 직접 들어가서 갱신해야 합니다.",
           font=_font(13), fill=(80, 80, 80))
    out = ILLUS_DIR / "github_three_places.png"
    img.save(out, "PNG")
    return out


def make_github_signup_illus() -> Path:
    """github.com 회원가입 화면 모형."""
    W, H = 900, 480
    img = Image.new("RGB", (W, H), (250, 250, 250))
    d = ImageDraw.Draw(img)
    # 브라우저 프레임
    d.rectangle((20, 20, W - 20, H - 20), fill=(255, 255, 255),
                outline=(180, 180, 180), width=2)
    d.rectangle((20, 20, W - 20, 60), fill=(240, 240, 240))
    d.ellipse((35, 32, 51, 48), fill=(255, 95, 86))
    d.ellipse((55, 32, 71, 48), fill=(255, 189, 46))
    d.ellipse((75, 32, 91, 48), fill=(39, 201, 63))
    d.rectangle((110, 28, W - 40, 52), fill=(255, 255, 255),
                outline=(200, 200, 200))
    d.text((120, 32), "https://github.com/signup",
           font=_font(13), fill=(80, 80, 80))
    # GitHub 로고
    d.ellipse((40, 80, 78, 118), fill=(36, 41, 47))
    d.text((90, 86), "GitHub", font=_font(22, bold=True), fill=(36, 41, 47))
    # 폼
    d.text((50, 150), "이메일 (Email)",
           font=_font(13, bold=True), fill=(40, 40, 40))
    d.rectangle((50, 174, 850, 210), fill=(255, 255, 255),
                outline=(180, 180, 180))
    d.text((60, 180), "example@gmail.com",
           font=_font(13), fill=(160, 160, 160))

    d.text((50, 230), "비밀번호 (Password) — 영문·숫자·기호 섞어 15자 이상",
           font=_font(13, bold=True), fill=(40, 40, 40))
    d.rectangle((50, 254, 850, 290), fill=(255, 255, 255),
                outline=(180, 180, 180))
    d.text((60, 260), "••••••••••••••••",
           font=_font(13), fill=(160, 160, 160))

    d.text((50, 310), "사용자명 (Username) — 영어 + 숫자, 한글 안 됨",
           font=_font(13, bold=True), fill=(40, 40, 40))
    d.rectangle((50, 334, 850, 370), fill=(255, 255, 255),
                outline=HL, width=2)
    d.text((60, 340), "my-clinic-2026",
           font=_font(13, bold=True), fill=(40, 40, 40))

    # 가입 버튼
    d.rectangle((50, 400, 850, 444), fill=VS_GREEN)
    d.text((380, 412), "Create account",
           font=_font(15, bold=True), fill=(255, 255, 255))

    # 강조
    _circle_arrow(d, x=820, y=352, label="이 이름을 개발팀에게 알려주셔야 합니다")
    out = ILLUS_DIR / "github_signup.png"
    img.save(out, "PNG")
    return out


def make_git_install_illus() -> Path:
    """git-scm.com 다운로드 페이지 모형."""
    W, H = 900, 360
    img = Image.new("RGB", (W, H), (250, 250, 250))
    d = ImageDraw.Draw(img)
    # 브라우저 프레임
    d.rectangle((20, 20, W - 20, H - 20), fill=(255, 255, 255),
                outline=(180, 180, 180), width=2)
    d.rectangle((20, 20, W - 20, 60), fill=(240, 240, 240))
    d.rectangle((110, 28, W - 40, 52), fill=(255, 255, 255),
                outline=(200, 200, 200))
    d.text((120, 32), "https://git-scm.com/download/win",
           font=_font(13), fill=(80, 80, 80))
    # 본문
    d.text((50, 90), "Download Git for Windows",
           font=_font(22, bold=True), fill=(40, 40, 40))
    d.text((50, 130), "Click here to download the latest version (64-bit Setup)",
           font=_font(14), fill=(80, 80, 80))
    # 큰 다운로드 버튼
    d.rectangle((50, 170, 380, 224), fill=(240, 80, 50))
    d.text((85, 184), "Click here to download",
           font=_font(15, bold=True), fill=(255, 255, 255))
    # 강조
    _circle_arrow(d, x=215, y=197, label="이 버튼 클릭")
    # 안내
    d.text((50, 250), "내려받은 설치파일을 실행하시고, 모든 화면에서 [ Next ] 만 눌러도 안전합니다.",
           font=_font(14), fill=(40, 40, 40))
    d.text((50, 278), "설치가 끝나면 PC를 한 번 재시작해 주세요.",
           font=_font(14), fill=(40, 40, 40))
    out = ILLUS_DIR / "git_install.png"
    img.save(out, "PNG")
    return out


def make_vscode_signin_github_illus() -> Path:
    """VSCode 좌측 하단 사람 아이콘 → Sign in to GitHub."""
    W, H = 900, 420
    img = Image.new("RGB", (W, H), VS_BG)
    d = ImageDraw.Draw(img)
    # 좌측 활동바
    d.rectangle((0, 0, 48, H), fill=VS_SIDE)
    # 사람 아이콘 위치 (좌측 하단)
    px, py = 24, H - 60
    d.ellipse((px - 9, py - 14, px + 9, py + 4), outline=VS_TEXT, width=2)
    d.arc((px - 14, py - 2, px + 14, py + 26), 200, 340, fill=VS_TEXT, width=2)
    # 팝업 메뉴
    d.rectangle((60, H - 220, 380, H - 40),
                fill=VS_PANEL, outline=(80, 80, 80))
    d.text((78, H - 210), "Accounts",
           font=_font(13, bold=True), fill=VS_DIM)
    d.line((78, H - 188, 360, H - 188), fill=(80, 80, 80), width=1)
    d.text((78, H - 178), "Sign in with GitHub",
           font=_font(14, bold=True), fill=VS_TEXT)
    d.text((78, H - 152), "Sign in with GitHub for Settings Sync",
           font=_font(12), fill=VS_DIM)
    d.text((78, H - 128), "Sign in with Microsoft",
           font=_font(12), fill=VS_DIM)

    _circle_arrow(d, x=px, y=py, label="사람 모양 아이콘 클릭")
    _circle_arrow(d, x=300, y=H - 173, label="이 항목 선택")

    # 상단 안내
    d.text((40, 40), "VSCode 좌측 활동바 가장 아래의 사람 아이콘을 클릭하면",
           font=_font(15), fill=VS_TEXT)
    d.text((40, 70), "위와 같은 작은 메뉴가 뜹니다. [ Sign in with GitHub ] 을 선택하세요.",
           font=_font(15), fill=VS_TEXT)
    d.text((40, 110), "브라우저가 자동으로 열리며 GitHub 로그인 페이지가 나옵니다 —",
           font=_font(14), fill=VS_DIM)
    d.text((40, 138), "방금 8-1 단계에서 만든 계정으로 로그인하시고 [ Authorize ] 버튼을 누르세요.",
           font=_font(14), fill=VS_DIM)
    out = ILLUS_DIR / "vscode_signin_github.png"
    img.save(out, "PNG")
    return out


def make_vscode_publish_to_github_illus() -> Path:
    """Source Control 패널의 'Publish to GitHub' 버튼."""
    W, H = 900, 460
    img = Image.new("RGB", (W, H), VS_BG)
    d = ImageDraw.Draw(img)
    # 좌측 활동바
    d.rectangle((0, 0, 48, H), fill=VS_SIDE)
    # Source Control 아이콘 (세번째) — 나뭇가지
    sx, sy = 24, 110
    d.ellipse((sx - 4, sy - 14, sx + 4, sy - 6), fill=VS_TEXT)
    d.ellipse((sx - 4, sy + 6, sx + 4, sy + 14), fill=VS_TEXT)
    d.ellipse((sx + 10, sy - 4, sx + 18, sy + 4), fill=VS_TEXT)
    d.line((sx, sy - 10, sx, sy + 10), fill=VS_TEXT, width=2)
    d.line((sx, sy, sx + 12, sy), fill=VS_TEXT, width=2)
    # Source Control 패널
    d.rectangle((48, 0, 360, H), fill=VS_PANEL)
    d.text((64, 14), "SOURCE CONTROL",
           font=_font(13, bold=True), fill=VS_DIM)
    # 큰 Publish 버튼
    d.rectangle((64, 60, 344, 110), fill=VS_ACCENT)
    d.text((110, 76), "Publish to GitHub",
           font=_font(16, bold=True), fill=(255, 255, 255))
    # 설명
    d.text((64, 130), "또는 'Publish Branch' 라고",
           font=_font(12), fill=VS_DIM)
    d.text((64, 152), "표시될 수도 있습니다 — 같은 기능입니다.",
           font=_font(12), fill=VS_DIM)

    _circle_arrow(d, x=sx, y=sy, label="나뭇가지 모양 아이콘")
    _circle_arrow(d, x=200, y=85, label="이 버튼 한 번 클릭")

    # 우측 안내
    d.text((400, 40), "버튼을 누르면 위쪽 검색창이 뜹니다.",
           font=_font(15), fill=VS_TEXT)
    d.text((400, 72), "[ Publish to GitHub private repository ] 선택",
           font=_font(14, bold=True), fill=(255, 220, 100))
    d.text((400, 110), "(나만 볼 수 있는 비공개 창고로 만들기)",
           font=_font(12), fill=VS_DIM)
    d.text((400, 150), "이어서 저장소 이름을 물어봅니다 —",
           font=_font(14), fill=VS_TEXT)
    d.text((400, 180), "예) denvia-chatbot 처럼 영문으로 입력 후 Enter",
           font=_font(13), fill=VS_DIM)
    d.text((400, 230), "끝나면 자동으로 본인 GitHub에 저장소가",
           font=_font(14), fill=VS_TEXT)
    d.text((400, 260), "만들어지고 코드도 한꺼번에 올라갑니다.",
           font=_font(14), fill=VS_TEXT)
    d.text((400, 310), "주소는 다음 형태로 생성됩니다 —",
           font=_font(13), fill=VS_DIM)
    d.text((400, 340), "https://github.com/[본인ID]/denvia-chatbot",
           font=_font(13, bold=True), fill=(255, 220, 100))
    d.text((400, 378), "이 주소를 9장에서 개발팀에 전달하시면 됩니다.",
           font=_font(13), fill=VS_DIM)
    out = ILLUS_DIR / "vscode_publish_to_github.png"
    img.save(out, "PNG")
    return out


def make_vscode_commit_sync_illus() -> Path:
    """일상 사용 — 변경, 커밋 메시지, 동기화."""
    W, H = 900, 460
    img = Image.new("RGB", (W, H), VS_BG)
    d = ImageDraw.Draw(img)
    # Source Control 패널 (좌측)
    d.rectangle((0, 0, 360, H), fill=VS_PANEL)
    d.text((16, 14), "SOURCE CONTROL",
           font=_font(13, bold=True), fill=VS_DIM)
    # 커밋 메시지 입력창
    d.rectangle((16, 40, 344, 78), fill=(60, 60, 60), outline=(80, 80, 80))
    d.text((26, 50), "메시지 입력 (예: 운영자 안내 문구 수정)",
           font=_font(12), fill=VS_TEXT)
    # 커밋 버튼
    d.rectangle((16, 90, 344, 124), fill=VS_ACCENT)
    d.text((130, 100), "✓ Commit",
           font=_font(14, bold=True), fill=(255, 255, 255))
    # 변경 파일 목록
    d.text((16, 144), "Changes (2)",
           font=_font(12, bold=True), fill=VS_DIM)
    for i, name in enumerate(["src/page.tsx", "README.md"]):
        y = 168 + i * 26
        d.text((28, y), f"M  {name}",
               font=_font(12), fill=(255, 200, 100))

    # 하단 상태바
    d.rectangle((0, H - 32, W, H), fill=VS_GREEN)
    d.text((20, H - 26), "main",
           font=_font(13, bold=True), fill=(255, 255, 255))
    # 동기화 아이콘 자리
    d.text((80, H - 26), "↑1 ↓0",
           font=_font(13, bold=True), fill=(255, 255, 255))
    _circle_arrow(d, x=95, y=H - 18, label="여기 클릭 = GitHub로 업로드")
    _circle_arrow(d, x=180, y=107, label="② 커밋")
    _circle_arrow(d, x=180, y=58, label="① 메시지 입력")

    # 우측 흐름 설명
    d.text((400, 30), "일상 사용 흐름 (세 단계)",
           font=_font(17, bold=True), fill=VS_TEXT)
    steps = [
        ("①", "코드 또는 파일을 수정합니다.",
         "VSCode 왼쪽 Source Control 패널에 자동으로 변경 파일이 나타납니다."),
        ("②", "변경 사항을 '커밋'합니다.",
         "위 입력창에 한 줄 메모(예: '상담 안내 문구 수정')를 적고 Commit."),
        ("③", "GitHub로 '푸시(업로드)'합니다.",
         "화면 가장 아래 'main ↑1' 부분을 클릭하면 끝. 잠시 후 GitHub에 반영됩니다."),
    ]
    y = 70
    for num, title, desc in steps:
        d.text((400, y), num, font=_font(22, bold=True), fill=(255, 220, 100))
        d.text((430, y + 2), title, font=_font(14, bold=True), fill=VS_TEXT)
        d.text((430, y + 30), desc, font=_font(12), fill=VS_DIM)
        y += 76

    out = ILLUS_DIR / "vscode_commit_sync.png"
    img.save(out, "PNG")
    return out


def make_handover_form_illus() -> Path:
    """개발팀 전달 정보 — 빈 양식 모형."""
    W, H = 900, 440
    img = Image.new("RGB", (W, H), (250, 250, 250))
    d = ImageDraw.Draw(img)
    # 양식 프레임
    d.rectangle((30, 30, W - 30, H - 30), fill=(255, 255, 255),
                outline=(180, 180, 180), width=2)
    d.text((50, 50), "개발팀 전달 양식",
           font=_font(20, bold=True), fill=(40, 40, 40))
    d.line((50, 90, W - 50, 90), fill=(200, 200, 200), width=1)

    fields = [
        ("GitHub 사용자명 (Username)",
         "예) my-clinic-2026   ← 8-1단계에서 정한 이름"),
        ("GitHub 가입 이메일",
         "예) owner@my-clinic.kr"),
        ("새로 만드신 저장소 주소 (URL)",
         "예) https://github.com/my-clinic-2026/denvia-chatbot"),
        ("개발팀을 '협업자(Collaborator)'로 추가하실까요?",
         "[      ] 네, 추가합니다.        [      ] 아니요, 추가 안 합니다."),
    ]
    y = 110
    for label, hint in fields:
        d.text((50, y), label,
               font=_font(13, bold=True), fill=(40, 40, 40))
        d.rectangle((50, y + 24, W - 50, y + 56),
                    fill=(248, 248, 248), outline=(200, 200, 200))
        d.text((60, y + 32), hint,
               font=_font(12), fill=(140, 140, 140))
        y += 72

    out = ILLUS_DIR / "handover_form.png"
    img.save(out, "PNG")
    return out


# ──────────────────────────────────────────────────────────────────────────────
# 문서 본문
# ──────────────────────────────────────────────────────────────────────────────

def build(output_path: Path) -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = KOREAN_FONT
    style.font.size = Pt(11)

    for sec in doc.sections:
        sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)
        sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)

    # ─── 표지
    _para(doc, DOC_TITLE, bold=True, size_pt=24,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, DOC_SUBTITLE, size_pt=13,
          align=WD_ALIGN_PARAGRAPH.CENTER,
          color=RGBColor(0x55, 0x55, 0x55))
    _para(doc, DOC_DATE, size_pt=11,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=18)

    # ─── 0. 이 문서를 읽는 분께
    _heading(doc, "이 문서를 읽는 분께", level=2)
    _para(doc,
          "Denvia 서비스가 운영되고 있는 서버에 직접 접속해 파일을 열어보거나 "
          "코드를 손봐야 할 때가 있습니다. 이 안내서는 그 작업을 "
          "Visual Studio Code(이하 VSCode)라는 무료 편집기로 할 수 있도록 단계별로 "
          "그림과 함께 정리했습니다.")
    _para(doc,
          "IT 비전공자분도 따라하실 수 있도록 어려운 용어는 비유로 풀어서 설명했습니다. "
          "한 번만 셋업해두면 다음부터는 클릭 두 번이면 서버에 들어갈 수 있습니다.")

    _callout(doc,
        ["• 평소 사이트 점검·관리자 페이지 사용만 하실 거면 이 문서를 굳이 보실 필요는 없습니다.",
         "• 서버에 직접 들어가야 하는 상황(긴급 점검·로그 확인·파일 직접 수정 등)에 한해 사용하세요.",
         "• 어렵게 느껴지시면 언제든 개발팀에 전화 주세요. 사고를 막는 것이 가장 중요합니다."],
        title="이런 분들이 보시면 됩니다",
        fill="FFF8DC",
    )

    # ─── 1. 원격 접속이 뭔가요? (개념)
    _heading(doc, "1. 원격 접속이 뭔가요?", level=1)
    _para(doc,
          "'원격 접속'은 말 그대로 멀리 있는 컴퓨터(서버)에 마치 그 컴퓨터 앞에 앉아 있는 것처럼 "
          "내 노트북에서 들어가 작업하는 일입니다. 사무실에서 집에 있는 PC를 조종하는 '팀뷰어'와 "
          "비슷한 개념이지만, 화면을 통째로 보여주는 대신 파일과 명령어만 안전하게 주고받습니다.")

    _figure(doc, make_concept_diagram(),
            caption="그림 1. 원격 접속 개념 — 내 노트북에서 SSH를 통해 운영 서버 안 파일을 열고 닫습니다.")

    _kv_table(doc, [
        ("내 노트북", "지금 이 문서를 보고 계신 컴퓨터. VSCode가 설치되어 있어야 합니다."),
        ("서버", "Denvia 사이트가 실제로 돌아가는 컴퓨터. AWS라는 회사의 클라우드 안에 있고 "
               "주소 이름은 denvia-prod-01입니다."),
        ("SSH", "두 컴퓨터 사이를 안전하게 잇는 길. 자물쇠(서버)와 열쇠(내 노트북의 .pem 파일)가 "
               "정확히 맞아야만 열립니다."),
        ("개인키 (.pem)", "서버 자물쇠를 열 수 있는 유일한 열쇠 파일. 이 파일이 유출되면 누구든 "
                       "서버에 들어올 수 있으므로 절대 카톡·이메일·USB로 공유하시면 안 됩니다."),
    ])

    # ─── 2. 준비물
    _heading(doc, "2. 준비물 체크리스트", level=1)
    _para(doc, "본격적으로 따라하시기 전에 아래 세 가지가 준비되어 있는지 확인해 주세요.")
    _bullet(doc, "VSCode가 설치되어 있다 — code.visualstudio.com 에서 무료 다운로드")
    _bullet(doc, "개인키 파일(.pem)을 받아두었다 — 인계 시 USB·암호화 압축으로 전달받으심")
    _bullet(doc, "그 .pem 파일을 [ 사용자 폴더 ] / .ssh / 폴더에 넣어두었다")

    _callout(doc,
        ["Windows 기준 [ 사용자 폴더 ] 의 위치는 다음과 같이 찾으실 수 있습니다.",
         "    파일 탐색기 주소창에 %USERPROFILE% 입력 후 Enter",
         "그 안에 .ssh 라는 폴더가 없으면 직접 만드신 뒤 .pem 파일을 넣어 두세요."],
        title="개인키(.pem) 파일을 어디에 두나요?",
        fill="EAF4FF", title_color="003366",
    )

    # ─── 3. 한 번만 하는 셋업
    _heading(doc, "3. 한 번만 하는 셋업 (15분)", level=1)
    _para(doc, "이 단계는 새 노트북을 받으셨을 때 등 처음 한 번만 하면 됩니다. "
               "두 번째부터는 4번 챕터부터 보시면 됩니다.")

    _heading(doc, "3-1. VSCode 설치 확인", level=2)
    _para(doc, "이미 VSCode를 사용 중이시라면 건너뛰셔도 됩니다. 미설치 시:")
    _bullet(doc, "https://code.visualstudio.com 접속")
    _bullet(doc, "큰 [ Download for Windows ] 버튼 클릭")
    _bullet(doc, "다운로드된 설치파일 실행 → [ 다음 ] 만 반복해서 설치 완료")

    _heading(doc, "3-2. Remote-SSH 확장 설치", level=2)
    _para(doc,
          "VSCode 자체에는 SSH 기능이 들어 있지 않습니다. 'Remote - SSH'라는 무료 확장(애드온)을 "
          "한 번 깔아두면 SSH 접속 기능이 추가됩니다.")
    _bullet(doc, "VSCode 왼쪽의 네모 4개(블록) 아이콘 클릭 — Extensions 패널이 열립니다")
    _bullet(doc, "상단 검색창에 Remote - SSH 입력")
    _bullet(doc, "결과 첫 번째 'Remote - SSH' (제작사: Microsoft)의 [ Install ] 버튼 클릭")

    _figure(doc, make_extensions_illus(),
            caption="그림 2. Extensions 패널에서 'Remote - SSH' 검색 후 Install 클릭")

    _callout(doc,
        ["설치가 끝나면 VSCode 왼쪽 가장 아래에 모니터 모양의 작은 아이콘이 새로 생깁니다. "
         "그 아이콘이 보이면 확장 설치는 성공입니다."],
        title="설치 성공 확인",
    )

    _heading(doc, "3-3. SSH 설정 파일(config)에 서버 정보 적어두기", level=2)
    _para(doc,
          "매번 접속할 때마다 서버 주소·사용자 이름·열쇠 위치를 입력하기 번거롭습니다. "
          "한 번만 'config'라는 작은 텍스트 파일에 적어두면, 다음부터는 'denvia-prod-01'이라는 "
          "별명으로 한 번에 들어갈 수 있습니다.")
    _para(doc, "다음 내용을 메모장으로 만들어 [ 사용자 폴더 ] / .ssh / config 라는 이름으로 저장하세요. "
               "(확장자 .txt 안 붙게 주의 — 저장 시 파일 형식을 '모든 파일'로 변경)")

    _code_block(doc,
"""Host denvia-prod-01
    HostName 15.164.114.152
    User ubuntu
    IdentityFile C:\\Users\\본인_윈도우_이름\\.ssh\\denvia-prod-01.pem
    IdentitiesOnly yes""")

    _callout(doc,
        ["• HostName 의 IP 주소(15.164.114.152)는 AWS Lightsail이 부여한 고정 IP 입니다. "
         "혹시 AWS 콘솔에서 IP를 변경하셨다면 이 줄도 함께 바꿔주셔야 합니다.",
         "• IdentityFile 의 경로는 본인 PC에 맞게 수정하세요. 예: 윈도우 사용자 이름이 'denvia' 라면 "
         "C:\\Users\\denvia\\.ssh\\denvia-prod-01.pem 이 됩니다.",
         "• 백슬래시(\\)는 두 번씩 적어야 합니다."],
        title="config 파일을 채우실 때 주의점",
        fill="FFF5F5", title_color="990000",
    )

    # ─── 4. 매번 접속하는 방법
    _heading(doc, "4. 매번 접속하는 방법 (1분)", level=1)
    _para(doc, "셋업이 끝났다면 이제 매 접속은 키보드 단축키 두 번이면 끝납니다.")

    _heading(doc, "4-1. 명령 팔레트 열기", level=2)
    _para(doc, "VSCode를 켠 뒤 키보드에서 ", size_pt=11)
    _para(doc, "        Ctrl + Shift + P", bold=True, font=MONO_FONT, size_pt=12)
    _para(doc, "를 동시에 누르면 화면 위쪽에 검색창(명령 팔레트)이 뜹니다. 거기에 다음을 입력하세요.")
    _para(doc, "        Remote-SSH: Connect to Host",
          bold=True, font=MONO_FONT, size_pt=12)
    _para(doc, "오타가 없도록 그대로 복사해 붙여넣어도 됩니다.")

    _figure(doc, SHOT_CLEAN / "03_command_palette.png",
            caption="그림 3. 명령 팔레트에 'Remote-SSH: Connect to Host' 입력 후 첫 번째 항목 선택")

    _heading(doc, "4-2. 서버 선택", level=2)
    _para(doc, "Enter를 누르면 등록된 서버 목록이 보입니다. 'denvia-prod-01' 을 선택하세요.")

    _figure(doc, SHOT_CLEAN / "04_host_picker.png",
            caption="그림 4. 서버 목록에서 denvia-prod-01 선택")

    _callout(doc,
        ["처음 접속할 때 한 번 '서버 지문(fingerprint)이 맞는지 확인하시겠습니까?' 라는 영문 메시지가 "
         "뜰 수 있습니다. 그때만 [ Continue ] 를 눌러주시면 됩니다. 이후로는 다시 묻지 않습니다."],
        title="처음 한 번만 묻는 메시지",
        fill="EAF4FF", title_color="003366",
    )

    _heading(doc, "4-3. 접속 완료 확인", level=2)
    _para(doc,
          "잠시(약 10~20초) 후 새 VSCode 창이 열리며 접속이 완료됩니다. "
          "화면 왼쪽 가장 아래에 초록색으로 'SSH: denvia-prod-01' 이라는 글자가 표시되면 성공입니다.")

    _figure(doc, make_connected_status_illus(),
            caption="그림 5. 좌측 하단 상태바 — 초록색 'SSH: denvia-prod-01' 표시")

    # ─── 5. 서버 폴더 열기
    _heading(doc, "5. 서버 안의 폴더 열어보기", level=1)
    _para(doc,
          "접속만으로는 아무 파일도 보이지 않습니다. 'Open Folder'로 어느 폴더를 들여다볼지 정해야 "
          "합니다. Denvia 서버의 소스 코드는 항상 다음 경로에 있습니다.")
    _para(doc, "        /opt/denvia",
          bold=True, font=MONO_FONT, size_pt=12)
    _para(doc, "방법:")
    _bullet(doc, "상단 메뉴 [ File ] → [ Open Folder ] 클릭")
    _bullet(doc, "또는 단축키 Ctrl + K → Ctrl + O 차례로 누르기")
    _bullet(doc, "나오는 입력창에 /opt/denvia 입력 후 [ OK ] 클릭")

    _figure(doc, make_remote_folder_picker_illus(),
            caption="그림 6. 폴더 경로 입력창에 /opt/denvia 입력")

    _para(doc,
          "처음 폴더를 여실 때 'Do you trust the authors?' 라는 영문 안전 확인 메시지가 뜨면 "
          "[ Yes, I trust the authors ] 를 선택해 주세요. 이후 왼쪽에 서버 안 파일·폴더 목록이 "
          "주르륵 나타납니다. 이때부터는 마치 내 노트북의 파일을 다루듯이 열고 닫고 편집할 수 있습니다.")

    # ─── 6. 자주 만나는 문제
    _heading(doc, "6. 잘 안 될 때 — 자주 만나는 문제", level=1)

    _kv_table(doc, [
        ("Permission denied (publickey)",
         "열쇠(.pem) 파일을 못 찾았거나 권한이 너무 열려 있는 경우입니다.\n"
         "→ config 의 IdentityFile 경로가 실제 파일 위치와 일치하는지 확인해 주세요.\n"
         "→ .pem 파일을 마우스 오른쪽 클릭 → [ 속성 ] → [ 보안 ] 에서 본인 외 사용자 권한을 모두 제거."),

        ("Connection timed out / refused",
         "서버 IP 주소가 바뀌었거나 본인 인터넷에 문제가 있을 수 있습니다.\n"
         "→ AWS Lightsail 콘솔에서 denvia-prod-01의 IP 주소를 다시 확인하고 config 의 HostName 수정.\n"
         "→ 그래도 안 되면 개발팀에 연락 (서버가 꺼져 있을 가능성)."),

        ("Could not establish connection",
         "회사·집 공유기에서 SSH 트래픽을 막고 있을 가능성이 있습니다.\n"
         "→ 다른 네트워크(스마트폰 핫스팟 등)로 바꿔서 시도해 보시거나 IT 담당자에게 문의."),

        ("열쇠 파일을 잃어버렸어요",
         "AWS Lightsail 콘솔에서 새 키 페어를 발급받아야 합니다.\n"
         "기존 키를 가진 사람은 더 이상 들어올 수 없게 됩니다.\n"
         "→ 분실 의심 시 즉시 개발팀에 알려 키 교체를 요청해 주세요. 보안 사고와 직결됩니다."),
    ], label_w=5.5, value_w=10.5)

    # ─── 7. 보안 수칙
    _heading(doc, "7. 꼭 지켜주실 보안 수칙", level=1)
    _callout(doc,
        ["1. 개인키(.pem) 파일은 절대 카톡·이메일·구글드라이브로 공유하지 않습니다. "
         "전달이 필요하면 USB 또는 암호 걸린 압축파일로만.",
         "2. 작업이 끝났으면 VSCode 좌하단의 'SSH: denvia-prod-01' 부분을 클릭 → "
         "[ Close Remote Connection ] 으로 연결을 닫아주세요.",
         "3. 서버 안에서 무엇을 지우거나 수정하실 일이 있으시면 반드시 사전에 개발팀과 협의. "
         "'rm' 같은 명령은 휴지통이 없어 즉시 영구 삭제됩니다.",
         "4. 본 안내서·.pem 파일·config 내용은 회사 내부 자료입니다. 외부에 공유하지 않습니다."],
        title="이 네 가지만 지켜주시면 안전합니다",
        fill="FFF5F5", title_color="990000",
    )

    # ─── 8. GitHub 처음 설정하기
    _heading(doc, "8. GitHub 처음 설정하기 (Git이 처음이신 분 전용)", level=1)
    _para(doc,
          "이 챕터는 Git이나 GitHub를 한 번도 써 보신 적이 없으신 분을 위해 "
          "맨 처음부터 차근차근 안내합니다. 이미 GitHub 계정과 VSCode 연동이 끝나 있으시면 "
          "이 챕터는 건너뛰셔도 됩니다.")

    _heading(doc, "8-0. GitHub가 도대체 뭔가요?", level=2)
    _para(doc,
          "GitHub는 '인터넷에 있는 코드 창고'입니다. 우리가 작업한 코드는 세 곳에 따로 존재합니다.")
    _figure(doc, make_github_three_places_diagram(),
            caption="그림 7. 코드가 존재하는 3곳 — 왼쪽 두 곳만 자동으로 연결되고, 운영 서버는 별개")

    _callout(doc,
        ["• 내 PC ↔ GitHub : VSCode 버튼 한 번이면 동기화됩니다.",
         "• GitHub ↔ 운영 서버 : 자동이 아닙니다. 사람이 서버에 직접 들어가 갱신해야 합니다. "
         "(5장 원격 접속 절차)",
         "• GitHub는 '백업 + 변경 이력 + 협업' 세 가지를 한 번에 해결해 줍니다. "
         "어느 시점으로든 되돌릴 수 있는 안전망이라고 보시면 됩니다."],
        title="가장 중요한 사실",
        fill="FFF8DC",
    )

    _heading(doc, "8-1. GitHub 계정 만들기 (5분)", level=2)
    _para(doc, "다음 순서대로 따라하시면 됩니다.")
    _bullet(doc, "브라우저로 https://github.com/signup 접속")
    _bullet(doc, "이메일 → 비밀번호 → 사용자명(Username) 순서로 입력")
    _bullet(doc, "사용자명은 영어와 숫자만, 한글 안 됩니다. 한 번 정하면 바꾸기 번거로우니 신중하게.")
    _bullet(doc, "이메일 인증 코드 입력 → 가입 완료")

    _figure(doc, make_github_signup_illus(),
            caption="그림 8. GitHub 회원가입 — 사용자명은 한 번 정하면 사실상 영구적이라 신중하게")

    _callout(doc,
        ["가입 후 정하신 사용자명(Username)은 9장에서 개발팀에 알려주셔야 하는 필수 정보입니다. "
         "쪽지로 적어두시거나 사진을 찍어두세요."],
        title="꼭 기록해두실 정보",
        fill="EAF4FF", title_color="003366",
    )

    _heading(doc, "8-2. Git 프로그램 설치 (3분)", level=2)
    _para(doc,
          "GitHub와 통신하려면 'Git' 이라는 작은 프로그램이 내 PC에 깔려 있어야 합니다. "
          "택배회사로 비유하자면 GitHub가 본사 창고이고, Git은 우리 집 앞까지 와주는 택배 기사님입니다.")
    _bullet(doc, "https://git-scm.com/download/win 접속")
    _bullet(doc, "큰 다운로드 버튼 클릭 → 내려받은 파일 실행")
    _bullet(doc, "설치 화면에서 [ Next ] 만 끝까지 눌러도 안전합니다 (기본값이 가장 안전한 설정)")
    _bullet(doc, "설치가 끝나면 PC를 한 번 재시작")

    _figure(doc, make_git_install_illus(),
            caption="그림 9. Git for Windows 다운로드 페이지")

    _callout(doc,
        ["설치 확인 방법 — VSCode를 열고 메뉴 [ Terminal ] → [ New Terminal ] 클릭. "
         "검은 창에 다음을 입력하고 Enter:",
         "    git --version",
         "버전 번호 (예: git version 2.45.0.windows.1) 가 나오면 정상 설치된 것입니다."],
        title="설치 성공 확인",
    )

    _heading(doc, "8-3. VSCode에서 GitHub 계정 연결 (1분)", level=2)
    _para(doc,
          "VSCode에 본인의 GitHub 계정 정보를 한 번만 로그인해두면, 이후로는 "
          "매번 아이디/비밀번호를 입력하지 않아도 됩니다.")
    _bullet(doc, "VSCode 왼쪽 활동바 가장 아래의 [ 사람 모양 아이콘 ] 클릭")
    _bullet(doc, "팝업 메뉴에서 [ Sign in with GitHub ] 선택")
    _bullet(doc, "기본 브라우저가 자동으로 열리며 GitHub 로그인 페이지가 뜹니다 — 로그인")
    _bullet(doc, "이어서 나오는 [ Authorize Visual-Studio-Code ] 버튼 클릭")
    _bullet(doc, "'VSCode를 열겠습니까?' 라는 메시지가 뜨면 [ 열기 ] 선택")

    _figure(doc, make_vscode_signin_github_illus(),
            caption="그림 10. VSCode 좌측 하단 사람 아이콘 → Sign in with GitHub")

    _callout(doc,
        ["연결이 끝나면 좌측 하단 사람 아이콘 옆에 본인 사용자명이 작게 표시됩니다. "
         "그 표시가 보이면 성공입니다."],
        title="연결 성공 확인",
    )

    _heading(doc, "8-4. 프로젝트 폴더를 GitHub에 올리기 (2분)", level=2)
    _para(doc,
          "개발팀에서 압축 파일로 전달받으신 프로젝트 폴더를 본인의 GitHub 창고로 올릴 차례입니다. "
          "이 작업도 버튼 한 번이면 끝납니다.")
    _bullet(doc, "VSCode에서 메뉴 [ File ] → [ Open Folder ] 로 프로젝트 폴더(예: D:\\Denvia) 열기")
    _bullet(doc, "왼쪽 활동바에서 [ 나뭇가지 모양 아이콘 ] (Source Control) 클릭")
    _bullet(doc, "큰 파란색 [ Publish to GitHub ] 버튼 클릭")
    _bullet(doc, "위쪽 검색창에 [ Publish to GitHub private repository ] 선택 (비공개 권장)")
    _bullet(doc, "저장소 이름(예: denvia-chatbot) 입력 후 Enter — 끝")

    _figure(doc, make_vscode_publish_to_github_illus(),
            caption="그림 11. Source Control 패널의 'Publish to GitHub' 버튼 한 번이면 끝")

    _callout(doc,
        ["• 비공개(private) vs 공개(public) — 반드시 'private' 선택. "
         "공개로 올리면 누구나 코드를 볼 수 있게 됩니다.",
         "• 저장소 이름은 영문 + 하이픈만 (예: denvia-chatbot). 한글·공백·특수문자 금지.",
         "• 완료되면 브라우저로 https://github.com/[본인사용자명] 에 접속해 보세요. "
         "방금 만든 저장소가 보이면 성공입니다.",
         "• 이때 만들어진 주소 https://github.com/[본인사용자명]/[저장소이름] 를 "
         "9장에서 개발팀에 전달하시면 됩니다."],
        title="이 단계 핵심 주의사항",
        fill="FFF5F5", title_color="990000",
    )

    _heading(doc, "8-5. 일상 사용 — 변경, 저장, 업로드", level=2)
    _para(doc,
          "8-4까지 끝났다면 이제부터의 작업 흐름은 매우 단순합니다. "
          "파일을 고칠 때마다 세 단계만 반복하시면 됩니다.")

    _figure(doc, make_vscode_commit_sync_illus(),
            caption="그림 12. 일상 사용 흐름 — 메시지 입력 → 커밋 → 동기화")

    _kv_table(doc, [
        ("① 수정",
         "원하는 파일을 열어 평소처럼 편집합니다. 저장(Ctrl+S)하면 "
         "왼쪽 Source Control 패널에 자동으로 변경 파일이 표시됩니다."),
        ("② 커밋 (Commit)",
         "변경한 내용에 짧은 메모를 다는 단계입니다. "
         "예) '상담 안내 문구 수정', '운영팀 연락처 변경'. "
         "메시지 입력창에 적고 그 아래 [ ✓ Commit ] 버튼 클릭."),
        ("③ 푸시 (Sync)",
         "본인 PC의 변경 내역을 GitHub 창고로 올리는 단계입니다. "
         "화면 가장 아래 [ main ↑1 ] 부분 클릭. 잠시 후 GitHub에 반영됩니다."),
    ], label_w=4.0, value_w=12.0)

    _callout(doc,
        ["② 커밋 메시지는 본인이 나중에 보고 알아볼 수 있게 한국어로 적으셔도 됩니다. "
         "'어떤 화면을 어떻게 바꿨는지' 한 줄로 적는 것이 가장 좋습니다.",
         "③ 푸시까지 끝나야 GitHub에 실제로 반영됩니다. 커밋만 하고 푸시를 안 하면 본인 PC에만 남아 있습니다.",
         "운영 서버에는 여전히 자동으로 안 올라갑니다 — 8-0의 그림 7을 다시 확인하세요. "
         "운영 서버 반영은 5장 절차로 사람이 직접 들어가 갱신해야 합니다."],
        title="일상 사용 시 자주 묻는 질문",
        fill="EAF4FF", title_color="003366",
    )

    # ─── 9. 개발팀 전달 정보
    _heading(doc, "9. 인계 시 개발팀에게 전달하실 정보", level=1)
    _para(doc,
          "위 8장의 셋업이 모두 끝나면, 개발팀이 '운영 서버가 어디 GitHub를 바라봐야 하는지'를 "
          "다시 설정해야 합니다. 매번이 아니라 단 한 번만 하면 끝나는 작업이지만, "
          "이를 위해 다음 정보가 필요합니다. 아래 양식을 채워 이메일이나 카톡으로 보내주세요.")

    _figure(doc, make_handover_form_illus(),
            caption="그림 13. 개발팀 전달 양식 — 네 가지만 채워주시면 됩니다")

    _kv_table(doc, [
        ("① GitHub 사용자명",
         "8-1에서 정하신 영문 사용자명.\n"
         "예) my-clinic-2026"),
        ("② GitHub 가입 이메일",
         "GitHub에서 발송되는 알림(보안 경고·협업 초대 등)이 이쪽 메일함으로 갑니다.\n"
         "예) owner@my-clinic.kr"),
        ("③ 새 저장소(repository) 주소",
         "8-4에서 'Publish to GitHub' 후 만들어진 URL 전체.\n"
         "예) https://github.com/my-clinic-2026/denvia-chatbot"),
        ("④ 개발팀을 협업자로 추가 여부",
         "추가하시면 개발팀이 GitHub에서 직접 코드를 확인·수정할 수 있어 유지보수가 빨라집니다.\n"
         "추가 안 하셔도 동작에는 문제 없습니다. 본인 결정사항입니다."),
    ], label_w=5.0, value_w=11.0)

    _callout(doc,
        ["• 협업자 추가 방법(선택) — GitHub 저장소 페이지 → [ Settings ] → [ Collaborators ] → "
         "[ Add people ] → 개발팀이 알려드린 GitHub 사용자명 입력 → [ Add ]",
         "• 위 정보를 전달해 주시면 개발팀은 운영 서버의 GitHub 연결 주소를 "
         "딱 한 번 본인 저장소로 바꿔드립니다(30초 작업). 이후로는 본인이 푸시한 변경이 "
         "정상적으로 운영 서버 갱신 절차에 쓰일 수 있게 됩니다.",
         "• 비밀번호·개인키(.pem) 같은 민감 정보는 함께 보내지 마세요. 양식의 네 가지로 충분합니다."],
        title="추가 안내",
        fill="EAF4FF", title_color="003366",
    )

    # ─── 부록
    _heading(doc, "부록. 서버 안에서 가끔 쓰는 명령어", level=1)
    _para(doc,
          "VSCode 안에서 메뉴 [ Terminal ] → [ New Terminal ] 을 누르면 서버 안 터미널 창이 "
          "하단에 열립니다. 거기서 아래 명령을 입력하시면 서버 상태를 확인하실 수 있습니다. "
          "단, '명령을 직접 실행하는 일'은 가능한 한 개발팀에 맡겨주세요.")

    _kv_table(doc, [
        ("pwd",   "현재 내 위치(폴더 경로)를 알려줍니다."),
        ("ls",    "현재 폴더에 어떤 파일이 있는지 목록을 보여줍니다."),
        ("cd /opt/denvia",
                  "Denvia 소스 코드 폴더로 이동합니다."),
        ("docker compose -f infra/docker-compose.prod.yml ps",
                  "서비스(웹·API·데이터베이스 등)가 정상 동작 중인지 확인합니다."),
        ("exit",  "터미널 창을 닫고 서버에서 빠져나옵니다."),
    ], label_w=7.0, value_w=9.0)

    _para(doc, " ")
    _para(doc,
          "문서 끝. 더 궁금하신 점은 개발팀(연락처 별도 안내)으로 문의해 주세요.",
          size_pt=10, color=RGBColor(0x66, 0x66, 0x66),
          align=WD_ALIGN_PARAGRAPH.CENTER)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"저장 완료 -> {output_path}")


def main() -> None:
    out = ROOT / "자료" / "Denvia_VSCode_SSH_접속_안내서.docx"
    build(out)


if __name__ == "__main__":
    main()
