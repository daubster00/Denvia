"""고객 인계용 서버 운영 안내서 빌더.

2026-06-05 — 운영 중인 Denvia 서버(AWS Lightsail / denvia.ai.kr)를 클라이언트에게
인계할 시점에 전달하는 안내서입니다. 비개발자가 읽어도 이해할 수 있도록 비유 위주로
작성합니다.

산출물:
    자료/Denvia_서버_운영_인계서.docx

CLAUDE.md "고객 전달용 문서 작성 규칙" 준수 — SDT(content control), 폼 보호,
문서 보호, ActiveX 등 일체 사용하지 않음. 어떤 워드 버전에서도 자유 편집 가능.

실행:
    py -3 scripts/build_server_handover_guide.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


KOREAN_FONT = "맑은 고딕"

DOC_TITLE = "Denvia 서버 운영 인계서"
DOC_SUBTITLE = "AWS Lightsail (서울 리전) · https://denvia.ai.kr"
DOC_DATE = "작성일: 2026년 6월 5일"


# ──────────────────────────────────────────────────────────────────────────────
# 스타일 헬퍼 (build_alimtalk_guide_v4.py와 동일 패턴)
# ──────────────────────────────────────────────────────────────────────────────

def _apply_kr_font(run, *, bold: bool = False, size_pt: float | None = None,
                   color: RGBColor | None = None) -> None:
    run.font.name = KOREAN_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), KOREAN_FONT)
    rfonts.set(qn("w:hAnsi"), KOREAN_FONT)
    rfonts.set(qn("w:ascii"), KOREAN_FONT)
    if bold:
        run.bold = True
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color


def _add_paragraph(doc, text, *, bold=False, size_pt=None, align=None,
                   color=None, space_after_pt=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_after_pt is not None:
        p.paragraph_format.space_after = Pt(space_after_pt)
    run = p.add_run(text)
    _apply_kr_font(run, bold=bold, size_pt=size_pt, color=color)


def _add_heading(doc, text, level=1) -> None:
    sizes = {1: 18, 2: 14, 3: 12}
    spaces_before = {1: 18, 2: 14, 3: 10}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(spaces_before[level])
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _apply_kr_font(run, bold=True, size_pt=sizes[level])


def _shade_cell(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _add_horizontal_rule(doc) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_callout_box(doc, body: str, *, fill: str = "FFF8DC", size_pt: float = 11) -> None:
    """안내 박스(연한 배경의 단일 셀 표)."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Cm(16)
    cell = tbl.cell(0, 0)
    cell.width = Cm(16)
    _shade_cell(cell, fill)
    for i, line in enumerate(body.split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        _apply_kr_font(r, size_pt=size_pt)
    doc.add_paragraph()


def _add_table_header(table, headers: list[str], *, fill: str = "E0E0E0") -> None:
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        _apply_kr_font(r, bold=True, size_pt=10)
        _shade_cell(hdr_cells[i], fill)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _add_table_row(table, values: list[str], *, bold_first: bool = False) -> None:
    row = table.add_row().cells
    for i, v in enumerate(values):
        row[i].text = ""
        row[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        for j, line in enumerate(v.split("\n")):
            p = row[i].paragraphs[0] if j == 0 else row[i].add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(line)
            _apply_kr_font(r, bold=(bold_first and i == 0), size_pt=10)


def _add_kv_table(doc, rows: list[tuple[str, str]], *,
                  label_width_cm: float = 4.0,
                  value_width_cm: float = 12.0,
                  label_fill: str = "F2F2F2") -> None:
    """라벨-값 두 칸 표."""
    tbl = doc.add_table(rows=0, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Cm(label_width_cm)
    tbl.columns[1].width = Cm(value_width_cm)
    for label, value in rows:
        row = tbl.add_row().cells
        row[0].text = ""
        row[1].text = ""
        row[0].width = Cm(label_width_cm)
        row[1].width = Cm(value_width_cm)
        for cell in row:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _shade_cell(row[0], label_fill)
        # 라벨
        p0 = row[0].paragraphs[0]
        r0 = p0.add_run(label)
        _apply_kr_font(r0, bold=True, size_pt=10)
        # 값
        for i, line in enumerate(value.split("\n")):
            p = row[1].paragraphs[0] if i == 0 else row[1].add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(line)
            _apply_kr_font(r, size_pt=10)


# ──────────────────────────────────────────────────────────────────────────────
# 문서 빌드
# ──────────────────────────────────────────────────────────────────────────────

def build_document(output_path: Path) -> None:
    doc = Document()

    # 본문 기본 폰트
    style = doc.styles["Normal"]
    style.font.name = KOREAN_FONT
    style.font.size = Pt(11)

    # 페이지 여백
    for section in doc.sections:
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # ─── 표지
    _add_paragraph(doc, DOC_TITLE, bold=True, size_pt=22,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, DOC_SUBTITLE, size_pt=12,
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   color=RGBColor(0x55, 0x55, 0x55))
    _add_paragraph(doc, DOC_DATE, size_pt=11,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=24)

    # ─── 0. 이 문서를 읽는 분께
    _add_heading(doc, "이 문서를 읽는 분께", level=2)
    _add_paragraph(
        doc,
        "본 문서는 Denvia 서비스가 운영되고 있는 서버를 클라이언트께서 인계받으실 때 "
        "필요한 모든 정보를 정리한 안내서입니다. 어려운 기술 용어는 가능한 한 풀어서 "
        "설명했으니 IT 전공이 아니시더라도 편하게 읽으실 수 있습니다.",
        size_pt=11,
    )
    _add_paragraph(
        doc,
        "역할 분담은 다음과 같습니다.",
        size_pt=11,
    )
    _add_kv_table(doc, [
        ("클라이언트", "AWS·도메인·관리자 계정의 주인. 평소 사이트가 잘 살아 있는지 모니터링하시고, "
                  "회원 응대·공지 발송·결제 현황 확인 등을 관리자 페이지에서 수행하십니다."),
        ("개발자", "서버 점검·코드 배포·장애 복구·데이터베이스 작업 등 기술 작업을 담당합니다. "
                "AWS 콘솔이나 서버에 직접 들어가야 하는 작업은 개발자에게 요청해 주시면 됩니다."),
    ], label_width_cm=3.0, value_width_cm=13.0)

    # ─── 1. 한눈에 보기
    _add_heading(doc, "1. 한눈에 보기", level=1)
    _add_kv_table(doc, [
        ("사이트 주소", "https://denvia.ai.kr"),
        ("관리자 페이지", "https://denvia.ai.kr/admin"),
        ("서버 위치", "AWS Lightsail (서울 리전)"),
        ("서버 사양", "2 vCPU · 8 GB RAM · 160 GB SSD"),
        ("운영체제", "Ubuntu 22.04 LTS"),
        ("월 비용", "약 6.1만원 (Lightsail $40 + 자동 백업 $4 ≒ $44)"),
        ("자동 백업", "매일 1회 스냅샷 · 7일 보관"),
    ])

    # ─── 2. 사이트 접속 안내
    _add_heading(doc, "2. 사이트 접속 안내", level=1)
    _add_paragraph(
        doc,
        "Denvia는 사용자용 사이트와 관리자 페이지가 같은 도메인 아래에 있습니다. "
        "관리자 페이지는 운영관리자 계정으로 로그인됩니다.",
        size_pt=11,
    )
    tbl = doc.add_table(rows=1, cols=3)
    tbl.autofit = False
    tbl.columns[0].width = Cm(4.5)
    tbl.columns[1].width = Cm(6.0)
    tbl.columns[2].width = Cm(5.5)
    _add_table_header(tbl, ["구분", "주소", "용도"])
    _add_table_row(tbl, [
        "일반 사용자 페이지",
        "https://denvia.ai.kr",
        "회원가입·로그인·구독·1:1 문의",
    ])
    _add_table_row(tbl, [
        "관리자 페이지",
        "https://denvia.ai.kr/admin",
        "회원 관리·결제 현황·문의 응답·공지 발송",
    ])
    _add_table_row(tbl, [
        "운영관리자 계정",
        "admin@denvia.ai.kr",
        "초기 비밀번호는 인계 시 별도 안내",
    ])
    doc.add_paragraph()

    _add_callout_box(
        doc,
        "※ 인계 시 전달드린 초기 비밀번호는 첫 로그인 직후 클라이언트께서 직접 변경해 주세요. "
        "비밀번호 변경 방법은 관리자 페이지 우측 상단 [내 정보 → 비밀번호 변경]입니다.\n"
        "※ 운영관리자(operator) 계정은 회원 관리·결제 현황 확인·1:1 문의 응답·공지 발송 등 "
        "일상 운영에 필요한 모든 권한을 가집니다. 같은 권한으로 운영자를 추가하실 수도 있습니다.",
        fill="FFF8DC",
    )

    # ─── 3. 서버 안에는 어떤 것들이 들어 있나요?
    _add_heading(doc, "3. 서버 안에는 어떤 것들이 들어 있나요?", level=1)
    _add_paragraph(
        doc,
        "서버 한 대는 “여러 개의 방이 있는 한 건물”이라고 생각하시면 쉽습니다. "
        "Denvia 서버 건물 안에는 7개의 방이 들어 있고, 각 방은 정해진 역할을 맡고 있습니다.",
        size_pt=11,
    )
    tbl = doc.add_table(rows=1, cols=3)
    tbl.autofit = False
    tbl.columns[0].width = Cm(4.0)
    tbl.columns[1].width = Cm(7.5)
    tbl.columns[2].width = Cm(4.5)
    _add_table_header(tbl, ["방 이름", "역할 (비유로)", "기술 이름"])
    _add_table_row(tbl, [
        "안내 데스크",
        "건물 1층 로비. 모든 손님(사용자)이 처음 들어오는 입구이며, "
        "안전한 통신(HTTPS)을 처리합니다.",
        "nginx",
    ])
    _add_table_row(tbl, [
        "매장 화면",
        "손님이 직접 보는 매장 내부. 회원가입·로그인·구독 화면 등 "
        "사용자가 눈으로 보는 모든 화면을 보여줍니다.",
        "web (Next.js)",
    ])
    _add_table_row(tbl, [
        "본사 사무실",
        "매장 뒤편의 두뇌. 회원 정보 확인·결제 처리·문의 접수 등 "
        "보이지 않는 모든 처리를 담당합니다.",
        "api (FastAPI)",
    ])
    _add_table_row(tbl, [
        "야간 근무 직원",
        "정해진 시각에 자동으로 일을 합니다. 예: 매월 정기 결제, "
        "알림톡 발송, 만료된 차단 해제 등.",
        "worker · beat (Celery)",
    ])
    _add_table_row(tbl, [
        "큰 창고",
        "모든 회원·결제·문의·공지 자료가 영구 보관되는 창고입니다. "
        "데이터의 본진(本陣).",
        "postgres (PostgreSQL)",
    ])
    _add_table_row(tbl, [
        "책상 위 메모지",
        "자주 보는 정보를 빠르게 꺼내쓰기 위한 임시 보관소. "
        "전원이 꺼지면 사라지지만 속도가 매우 빠릅니다.",
        "redis",
    ])
    doc.add_paragraph()
    _add_callout_box(
        doc,
        "이 7개 방은 자동으로 함께 동작합니다.\n"
        "손님이 사이트에 접속 → 안내 데스크가 받음 → 매장 화면을 보여줌 → "
        "매장이 본사 사무실에 물어봄 → 사무실이 창고/메모지에서 정보를 꺼내 답함.\n"
        "어느 한 방이 멈추면 사이트가 비어 보이거나 일부 기능이 동작하지 않을 수 있습니다.",
        fill="EAF3FF",
    )

    # ─── 4. 데이터 보관과 백업
    _add_heading(doc, "4. 데이터 보관과 백업", level=1)
    _add_paragraph(
        doc,
        "회원 정보·결제 내역·1:1 문의·공지·알림톡 발송 이력 등 "
        "Denvia의 모든 데이터는 서버 안의 데이터베이스(PostgreSQL)에 저장됩니다.",
        size_pt=11,
    )
    tbl = doc.add_table(rows=1, cols=4)
    tbl.autofit = False
    tbl.columns[0].width = Cm(4.0)
    tbl.columns[1].width = Cm(3.0)
    tbl.columns[2].width = Cm(3.0)
    tbl.columns[3].width = Cm(6.0)
    _add_table_header(tbl, ["백업 종류", "주기", "보관 기간", "설명"])
    _add_table_row(tbl, [
        "AWS 자동 스냅샷",
        "매일 1회",
        "7일",
        "Lightsail이 자동으로 서버 전체를 통째로 백업합니다.",
    ])
    _add_table_row(tbl, [
        "DB 단위 백업",
        "배포 직전",
        "수동 보관",
        "새 코드를 배포할 때 개발자가 안전을 위해 데이터베이스만 별도로 백업합니다.",
    ])
    doc.add_paragraph()
    _add_callout_box(
        doc,
        "※ 7일 이상의 장기 보관(예: 30일·1년 단위 분기별 백업)이 필요하시면 별도 협의를 통해 "
        "정책을 추가할 수 있습니다. 추가 비용이 소액 발생합니다.",
        fill="FFF8DC",
    )

    # ─── 5. 평소 운영 — 살아있는지 확인하는 법
    _add_heading(doc, "5. 평소 운영 — 사이트가 살아있는지 확인하는 법", level=1)
    _add_paragraph(
        doc,
        "특별한 기술 없이도 클라이언트께서 직접 다음 3가지를 주기적으로 확인하시면 됩니다.",
        size_pt=11,
    )

    _add_heading(doc, "5-1. 사이트가 정상으로 보이는지", level=2)
    _add_paragraph(
        doc,
        "https://denvia.ai.kr 에 접속해 메인 화면이 정상적으로 보이면 됩니다. "
        "화면이 비어 있거나 “이 사이트에 연결할 수 없음” 메시지가 뜨면 장애 가능성이 있습니다.",
        size_pt=11,
    )

    _add_heading(doc, "5-2. 빠른 진단 URL — 헬스체크", level=2)
    _add_paragraph(
        doc,
        "다음 주소를 브라우저 주소창에 입력하시면 서버 자체의 건강 상태를 한 줄로 알려줍니다.",
        size_pt=11,
    )
    _add_callout_box(
        doc,
        "https://denvia.ai.kr/healthz",
        fill="F2F2F2",
        size_pt=11,
    )
    _add_paragraph(
        doc,
        "정상이면 다음과 같은 문장이 보입니다.",
        size_pt=11,
    )
    _add_callout_box(
        doc,
        '{"status":"ok","db":"up","redis":"up"}',
        fill="EAF7EA",
        size_pt=11,
    )
    _add_paragraph(
        doc,
        "“ok”가 아니거나 “db” / “redis” 중 하나라도 “down”으로 표시되면 개발자에게 캡처와 함께 즉시 알려주세요.",
        size_pt=11,
    )

    _add_heading(doc, "5-3. 관리자 페이지에서 일상 점검", level=2)
    _add_paragraph(
        doc,
        "관리자 페이지에 로그인하셔서 아래 2가지가 정상인지 가볍게 확인하시면 됩니다.",
        size_pt=11,
    )
    _add_kv_table(doc, [
        ("회원 현황", "신규 가입자 수가 평소와 큰 차이 없이 늘고 있는지"),
        ("결제 현황", "최근 결제가 비어 있지 않은지, 결제 실패가 비정상적으로 늘지 않았는지"),
    ], label_width_cm=3.5, value_width_cm=12.5)
    _add_callout_box(
        doc,
        "※ 알림톡 발송 성공/실패 이력은 Denvia 관리자 페이지가 아니라 "
        "알리고(Aligo) 콘솔에서 확인합니다. 알리고에 로그인하신 후 "
        "[전송결과 조회 → 알림톡 발송 이력] 메뉴에서 일자별 성공/실패 건수와 "
        "실패 사유를 확인하실 수 있습니다.",
        fill="EAF3FF",
    )

    # ─── 6. 장애가 의심되면
    _add_heading(doc, "6. 장애가 의심되면 — 증상별 대응표", level=1)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.autofit = False
    tbl.columns[0].width = Cm(4.0)
    tbl.columns[1].width = Cm(6.0)
    tbl.columns[2].width = Cm(6.0)
    _add_table_header(tbl, ["증상", "클라이언트가 먼저 할 일", "그 다음 단계"])
    _add_table_row(tbl, [
        "사이트가 안 열린다",
        "1~2분 기다린 후 다시 접속해보기. "
        "그래도 안 열리면 모바일·다른 인터넷으로 접속 시도.",
        "여전히 안 열리면 개발자에게 연락 (장애 가능).",
    ])
    _add_table_row(tbl, [
        "화면은 뜨는데 텅 비어 있다",
        "https://denvia.ai.kr/healthz 결과 화면을 캡처.",
        "캡처와 함께 개발자에게 전달.",
    ])
    _add_table_row(tbl, [
        "결제가 실패한다",
        "토스 페이먼츠 가맹점 콘솔에서 해당 결제의 거절 사유 확인.",
        "콘솔 화면 캡처와 함께 개발자에게 전달.",
    ])
    _add_table_row(tbl, [
        "알림톡이 발송되지 않는다",
        "알리고(Aligo) 콘솔 > 전송결과 조회 > 알림톡 발송 이력에서 실패 사유 확인.",
        "실패 사유 캡처와 함께 개발자에게 전달.",
    ])
    _add_table_row(tbl, [
        "관리자 페이지 로그인이 안 된다",
        "비밀번호 재확인. 모바일에서 시도. 다른 관리자 계정으로 시도.",
        "그래도 안 되면 개발자에게 비밀번호 초기화 요청.",
    ])
    doc.add_paragraph()
    _add_callout_box(
        doc,
        "※ 서버 재시작·코드 배포·데이터베이스 복구는 모두 개발자만 수행합니다.\n"
        "클라이언트께서 AWS Lightsail 콘솔에서 임의로 인스턴스를 중지(Stop)하거나 삭제(Delete)하시면 "
        "데이터 손실 위험이 있으니 절대 하지 말아 주세요. "
        "사양 변경·재시작이 필요하시면 반드시 개발자에게 먼저 문의해 주세요.",
        fill="FFEDED",
    )

    # ─── 7. 월 비용 안내
    _add_heading(doc, "7. 월 비용 안내", level=1)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.autofit = False
    tbl.columns[0].width = Cm(7.0)
    tbl.columns[1].width = Cm(4.0)
    tbl.columns[2].width = Cm(5.0)
    _add_table_header(tbl, ["항목", "월 비용 (USD)", "월 비용 (KRW 환산)"])
    _add_table_row(tbl, ["Lightsail $40 플랜 (서버)", "$40", "약 56,000원"])
    _add_table_row(tbl, ["자동 스냅샷 (백업)", "약 $4", "약 5,600원"])
    _add_table_row(tbl, ["합계", "약 $44", "약 61,600원"])
    doc.add_paragraph()
    _add_paragraph(
        doc,
        "※ 환율 1 USD = 1,400원 기준. 환율에 따라 ±5% 변동.",
        size_pt=10, color=RGBColor(0x55, 0x55, 0x55),
    )
    _add_paragraph(
        doc,
        "※ 위 금액은 서버 비용만 포함합니다. 아래 항목은 별도로 청구됩니다.",
        size_pt=11,
    )
    _add_kv_table(doc, [
        ("알리고 (알림톡)", "발송 건수에 따라 청구. 알림톡 1건 약 8.5원, SMS 1건 약 19원."),
        ("토스 페이먼츠 (결제)", "결제 수수료. 결제 금액의 약 2.9% (카드사·결제 수단에 따라 다름)."),
        ("도메인 (denvia.ai.kr)", "연 1회 갱신 비용. 약 22,000원 / 년 (가비아 기준)."),
    ], label_width_cm=4.0, value_width_cm=12.0)

    # ─── 8. 인계 자산 체크리스트
    _add_heading(doc, "8. 인계 자산 체크리스트", level=1)
    _add_paragraph(
        doc,
        "인계 시점에 클라이언트께서 보유하셔야 할 자산 목록입니다. "
        "각 자산의 비밀번호는 인계 직후 클라이언트께서 직접 변경하시기를 권장합니다.",
        size_pt=11,
    )
    tbl = doc.add_table(rows=1, cols=3)
    tbl.autofit = False
    tbl.columns[0].width = Cm(5.5)
    tbl.columns[1].width = Cm(5.5)
    tbl.columns[2].width = Cm(5.0)
    _add_table_header(tbl, ["자산", "보관처", "인계 후 비밀번호 변경"])
    _add_table_row(tbl, [
        "AWS 콘솔 계정 (Lightsail 포함)",
        "클라이언트 단독 보유",
        "권장",
    ])
    _add_table_row(tbl, [
        "도메인 denvia.ai.kr (가비아 등)",
        "클라이언트 단독 보유",
        "권장",
    ])
    _add_table_row(tbl, [
        "알리고 (알림톡) 계정",
        "클라이언트 단독 보유",
        "권장",
    ])
    _add_table_row(tbl, [
        "토스 페이먼츠 가맹점 계정",
        "클라이언트 단독 보유",
        "권장",
    ])
    _add_table_row(tbl, [
        "관리자 master 계정 (btmdesign@naver.com)",
        "클라이언트 보유",
        "필수",
    ])
    _add_table_row(tbl, [
        "서버 SSH 접속 키",
        "개발자 보관 (긴급 점검·배포용)",
        "인계 종료 시점에 협의",
    ])
    doc.add_paragraph()

    _add_horizontal_rule(doc)
    _add_paragraph(doc, "감사합니다.", bold=True, size_pt=12,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    repo_root = Path(__file__).resolve().parent.parent
    output_path = repo_root / "자료" / "Denvia_서버_운영_인계서.docx"
    build_document(output_path)
    print(f"[OK] Wrote {output_path}")


if __name__ == "__main__":
    main()
