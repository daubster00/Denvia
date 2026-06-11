"""Denvia 최종 납품 확인서 빌더.

2026-06-05 — 클라이언트 서명용 공식 납품 확인서.

산출물:
    자료/Denvia_최종납품확인서_2026-06-05.docx

CLAUDE.md "고객 전달용 문서 작성 규칙" 준수 — SDT(content control), 폼 보호,
문서 보호, ActiveX 등 일체 사용하지 않음. 어떤 워드 버전에서도 자유 편집 가능.

용도:
- 클라이언트가 납품 내역을 확인하고 서명·인 처리하는 공식 문서.
- 1) 계약 범위 vs 실제 납품 대조표
  2) 운영 자산 인계 내역(서버·도메인·계정·외부 연동)
  3) 유지보수 범위 및 면책
  4) 수령 확인·서명란
  순서로 구성.

실행:
    py -3 scripts/build_delivery_acceptance.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ──────────────────────────────────────────────────────────────────────────────
# 콘텐츠 정의
# ──────────────────────────────────────────────────────────────────────────────

DOC_TITLE = "최종 납품 확인서"
DOC_SUBTITLE = "Denvia · 치과 보험청구·데스크 행정 AI 어시스턴트"
DOC_DATE = "2026년 6월 5일"

PROJECT_NAME = "Denvia (치과 보험청구·데스크 행정 AI 어시스턴트) 구축"
SERVICE_URL = "https://denvia.ai.kr"


# 계약 범위 vs 실제 납품 대조 (MVP 범위 요약서 + 계약서 부록 기준)
SCOPE_VS_DELIVERY: list[dict] = [
    {
        "scope": "회원가입·로그인 (이메일 + 카카오·네이버·구글 소셜)",
        "delivery": (
            "이메일·카카오·네이버·구글 4종 가입/로그인 완료. "
            "휴대폰 본인인증, 가입유형(치과의사/치과위생사/기타) 선택, "
            "이메일·비밀번호 찾기, 무차별 시도 차단 포함."
        ),
        "status": "완료",
    },
    {
        "scope": "AI Q&A 채팅 (치과 보험청구·데스크 행정)",
        "delivery": (
            "RAG(지식 검색) + GPT 응답 엔진 적용. "
            "무료 사용자(일 한도 + 지연 응답) / 유료 사용자(무제한 + 실시간 스트리밍) 분기. "
            "답변 피드백, 후속 질문 인식, 대화 초기화 포함."
        ),
        "status": "완료",
    },
    {
        "scope": "구독·결제 시스템",
        "delivery": (
            "토스페이먼츠 빌링키 연동. 카드 등록 → 즉시 결제 → Pro 권한 활성화. "
            "매월 자동 결제, 결제 실패 시 자동 재시도(1일·3일), "
            "구독 해지/해지 철회, 결제 내역 5년 보관."
        ),
        "status": "완료",
    },
    {
        "scope": "환불 처리 (청약철회 + 운영 환불)",
        "delivery": (
            "청약철회 7일 자동 환불(즉시 처리). "
            "관리자 수동환불(부분환불 가능, 일할 권장값 자동 계산, 2단계 확인). "
            "구독 취소 분기 별도 처리."
        ),
        "status": "완료",
    },
    {
        "scope": "마이페이지 (구독 상태·결제 내역·탈퇴)",
        "delivery": (
            "현재 요금제·다음 결제일·결제 내역 표 확인, "
            "구독 해지/철회, 청약철회 환불 신청, "
            "회원 탈퇴(휴대폰 본인인증 후 즉시 처리) 포함."
        ),
        "status": "완료",
    },
    {
        "scope": "받은 쪽지함 + 1:1 문의 게시판",
        "delivery": (
            "쪽지함(공지·환불·이상탐지 안내 수신). "
            "1:1 문의 게시판 6종 유형(결제/계정/기능/오류/환불/기타), "
            "이미지 최대 5장 첨부, 관리자 답변 시 알림톡 자동 발송."
        ),
        "status": "완료",
    },
    {
        "scope": "관리자 페이지 — 대시보드·회원·통계·매출",
        "delivery": (
            "대시보드(가입자·구독자·매출·API 비용 위젯). "
            "회원 통합 검색·상세·권한 편집(Pro 부여/한도/속도/차단). "
            "통계(가입자 추이·피드백·매출·가입유형 분포) + 엑셀 내보내기."
        ),
        "status": "완료",
    },
    {
        "scope": "관리자 페이지 — RAG 데이터·프롬프트 관리",
        "delivery": (
            "RAG 자료(.txt) 등록·수정·삭제, 무중단 재빌드 버튼, "
            "동의어 사전 관리. 시스템 프롬프트·차단 키워드 편집, "
            "모델 파라미터(창의성·길이) 조정, 변경 이력 자동 기록."
        ),
        "status": "완료",
    },
    {
        "scope": "관리자 페이지 — 이상탐지·예산 통제(Kill Switch)",
        "delivery": (
            "비밀번호 5회 오류·동일 질문 반복·다중 IP 로그인 자동 감지 + 관리자 알림톡. "
            "월 예산 게이지(80%/95%/100% 자동 차단) + 수동 전체 정지(사유 + 2단계 확인). "
            "정지 기간만큼 유료 구독자 다음 결제일 자동 연기."
        ),
        "status": "완료",
    },
    {
        "scope": "관리자 페이지 — CS 관리(문의 답변·휴지통)",
        "delivery": (
            "신규 문의 알림톡 발송, 답변 서식(글머리·굵게·링크) 지원, "
            "답변 시 사용자 알림톡 자동 발송. CS 휴지통 별도 관리."
        ),
        "status": "완료",
    },
    {
        "scope": "관리자 페이지 — 팝업·전역 설정·관리자 관리",
        "delivery": (
            "공지 팝업 등록·노출 기간 설정. 글로벌 토글(구독 ON/OFF·1일 한도·지연시간). "
            "관리자 4등급 권한 체계(마스터/운영자/부운영자/대기) + 활동 로그."
        ),
        "status": "완료",
    },
    {
        "scope": "알림톡 메시지 (사용자 + 관리자용 전체)",
        "delivery": (
            "사용자용 13종 + 관리자용 7종 카카오 비즈 검수 통과, "
            "알리고 콘솔 템플릿 등록 + 실제 발송 테스트 완료. "
            "메시지 안내서(별첨) 제공."
        ),
        "status": "완료",
    },
    {
        "scope": "운영 서버 구축·도메인·SSL·자동 백업",
        "delivery": (
            "AWS Lightsail 서울 리전, Ubuntu 22.04, 2 vCPU / 7.6 GB / 155 GB. "
            "denvia.ai.kr 도메인 연결, SSL 인증서, 7개 컨테이너 서비스(웹·API·DB·Redis·워커·스케줄러·nginx) 가동, "
            "DB·RAG 자동 백업."
        ),
        "status": "완료",
    },
    {
        "scope": "법적 문서 (이용약관·개인정보처리방침)",
        "delivery": (
            "이용약관·개인정보처리방침 최종본 사이트 게재. "
            "개인정보 보호책임자 확정(denvia@naver.com)."
        ),
        "status": "완료",
    },
]


# 운영 자산 인계 — 카테고리별 상세
ASSET_HANDOVER: list[dict] = [
    {
        "category": "도메인",
        "rows": [
            ("도메인", "denvia.ai.kr"),
            ("등록처", "카페24"),
            ("계정 정보", "별도 ‘자산·계정 통합 인계서’로 전달 (ID/PW 포함)"),
        ],
    },
    {
        "category": "운영 서버 (AWS Lightsail)",
        "rows": [
            ("인스턴스", "denvia-prod-01 (서울 리전 ap-northeast-2)"),
            ("사양", "Ubuntu 22.04 LTS · 2 vCPU · 7.6 GB RAM · 155 GB SSD"),
            ("공인 IP", "15.164.114.152"),
            ("배포 경로", "/opt/denvia (docker compose)"),
            ("AWS 콘솔 접속", "별도 인계서로 전달 (계정·MFA·SSH 키)"),
        ],
    },
    {
        "category": "결제 (토스페이먼츠)",
        "rows": [
            ("벤더", "토스페이먼츠 (Toss Payments)"),
            ("키 관리", "관리자 페이지 → 설정 → 결제 키 관리에서 직접 갱신 가능 (마스킹 처리)"),
            ("운영 모드", "관리자 페이지에서 LIVE/SANDBOX 토글 (Redis SSOT)"),
        ],
    },
    {
        "category": "알림톡 (알리고)",
        "rows": [
            ("벤더", "알리고 (smartsms.aligo.in)"),
            ("등록 템플릿", "사용자 13종 + 관리자 7종 (총 20종)"),
            ("템플릿 SSOT", "docs/ALIMTALK_TEMPLATES.md + 안내서 별첨"),
        ],
    },
    {
        "category": "소셜 로그인 (OAuth)",
        "rows": [
            ("카카오", "OAuth 클라이언트 발급·등록 완료"),
            ("네이버", "OAuth 클라이언트 발급·등록 완료"),
            ("구글", "OAuth 클라이언트 발급·등록 완료 (운영 도메인 콜백 등록)"),
        ],
    },
    {
        "category": "AI / 외부 API",
        "rows": [
            ("LLM", "OpenAI API 키 (운영 적용 완료, 관리자 콘솔 환경변수)"),
            ("환율", "한국수출입은행 API 키 (매출 KRW 환산용)"),
            ("RAG 데이터", "초기 등록본 적재 완료. 이후 관리자 화면에서 직접 추가/수정 가능"),
        ],
    },
    {
        "category": "관리자 계정",
        "rows": [
            ("초기 운영 계정", "admin@denvia.ai.kr — 등급 ‘운영자’"),
            ("초기 비밀번호", "인계서 별도 전달 (수령 후 즉시 변경 권장)"),
            ("권한 체계", "마스터(개발사 1인) / 운영자 / 부운영자 / 대기 — 4등급"),
        ],
    },
    {
        "category": "소스 코드",
        "rows": [
            ("전달 방식", "전체 소스 코드 ZIP 압축본으로 일괄 전달 (Git 저장소 별도 이관 없음)"),
            ("포함 범위",
             "프론트엔드(web) + 백엔드(api) + 운영 docker compose 설정 + "
             "DB 마이그레이션 + 운영용 .env 예시 파일"),
            ("운영 서버 코드", "AWS Lightsail /opt/denvia 경로에 배포된 동일 코드와 일치"),
            ("보관 책임", "수령 이후 소스 코드 백업·보관 책임은 클라이언트에 있음"),
        ],
    },
]


# 검수 체크리스트 — 클라이언트가 확인하고 표시
ACCEPTANCE_CHECKLIST: list[dict] = [
    {
        "title": "1) 사용자 화면 동작 확인",
        "items": [
            "회원가입 (이메일·카카오·네이버·구글) 정상 동작 확인",
            "로그인 / 비밀번호 찾기 / 아이디 찾기 정상 동작 확인",
            "Q&A 채팅 — 무료/유료 사용자 분기 정상 동작 확인",
            "Pro 구독 결제 (실 카드 등록 → 즉시 결제) 정상 동작 확인",
            "마이페이지 — 구독 상태·결제 내역·해지·청약철회 정상 동작 확인",
            "쪽지함 / 1:1 문의 게시판 작성·열람 정상 동작 확인",
        ],
    },
    {
        "title": "2) 관리자 화면 동작 확인",
        "items": [
            "관리자 로그인 / 권한 등급 부여 정상 동작 확인",
            "대시보드 위젯·예산 게이지 정상 표시 확인",
            "회원 검색·상세·권한 편집·차단/해제 정상 동작 확인",
            "통계 / 매출 / 가입유형 분포 정상 표시 확인",
            "RAG 자료 등록·재빌드·동의어 사전 정상 동작 확인",
            "이상탐지·예산 통제(Kill Switch)·CS 답변 정상 동작 확인",
            "팝업·전역 설정·관리자 관리·활동 로그 정상 동작 확인",
        ],
    },
    {
        "title": "3) 알림톡 발송 확인",
        "items": [
            "사용자 알림톡 13종 발송 확인 (회원가입·결제·환불·문의 답변 등)",
            "관리자 알림톡 7종 발송 확인 (이상탐지·예산 경고·신규 문의 등)",
        ],
    },
    {
        "title": "4) 운영 자산 인계 확인",
        "items": [
            "도메인(denvia.ai.kr) 정상 접속 확인",
            "운영 서버 SSL 인증서 정상 적용 확인",
            "관리자 초기 계정 수령 및 비밀번호 변경 확인",
            "토스페이먼츠 / 알리고 / OAuth / OpenAI / 환율 API 키 인계 확인",
            "Git 저장소 접근 권한 확인",
        ],
    },
    {
        "title": "5) 문서 인계 확인",
        "items": [
            "Denvia 개발완료보고서 수령 확인",
            "Denvia 서버 안내서 수령 확인",
            "카카오 알림톡 메시지 안내서 수령 확인",
            "이용약관 / 개인정보처리방침 최종본 수령 확인",
        ],
    },
]


# 유지보수 / 면책
MAINTENANCE_TABLE: list[tuple[str, str]] = [
    ("무상 하자보수 범위",
     "합의된 기능명세 내에서 발견된 오작동·오류·버그의 수정 (계약서 제11조 제2항)"),
    ("무상 하자보수 기간",
     "본 확인서 서명일로부터 30일 (계약서 제11조 제1항)"),
    ("무상 범위 제외",
     "기능 추가·디자인 변경·외부 API/벤더 정책 변경 대응·운영 자료 추가 입력은 별도 유료 작업"),
    ("긴급 장애 대응",
     "운영 서버 다운·결제 시스템 마비 등 서비스 중단 장애는 즉시 대응 (24시간 이내 1차 회신)"),
    ("일반 문의 대응",
     "평일 업무시간 기준 1~3 영업일 내 회신"),
    ("운영 책임 범위",
     "본 확인서 서명 이후 발생하는 운영비(서버·OpenAI 토큰·알림톡·PG 수수료 등)는 클라이언트 부담"),
    ("AI 응답 면책",
     "Denvia는 치과 행정·보험청구 업무 보조 도구이며, 의료 진단·법적 판단 자료가 아닙니다. "
     "최종 결정 및 책임은 사용자(치과 의료진)에게 있습니다."),
]


# ──────────────────────────────────────────────────────────────────────────────
# 스타일 헬퍼 (build_final_report.py와 동일)
# ──────────────────────────────────────────────────────────────────────────────

KOREAN_FONT = "맑은 고딕"


def _apply_kr_font(run, *, bold: bool = False, size_pt: float | None = None,
                   color: RGBColor | None = None) -> None:
    run.font.name = KOREAN_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement

        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), KOREAN_FONT)
    rfonts.set(qn("w:hAnsi"), KOREAN_FONT)
    rfonts.set(qn("w:ascii"), KOREAN_FONT)
    if bold:
        run.bold = True
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color


def _add_paragraph(doc, text: str, *, bold: bool = False, size_pt: float | None = None,
                   align: int | None = None, color: RGBColor | None = None,
                   space_after_pt: float | None = None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_after_pt is not None:
        p.paragraph_format.space_after = Pt(space_after_pt)
    run = p.add_run(text)
    _apply_kr_font(run, bold=bold, size_pt=size_pt, color=color)


def _add_heading(doc, text: str, level: int = 1) -> None:
    sizes = {1: 18, 2: 14, 3: 12}
    spaces = {1: 16, 2: 12, 3: 8}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(spaces[level])
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _apply_kr_font(run, bold=True, size_pt=sizes[level])


def _shade_cell(cell, color_hex: str) -> None:
    from docx.oxml import OxmlElement

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, *, bold: bool = False, size_pt: float = 10,
                   color: RGBColor | None = None) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    for i, line in enumerate(text.split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        _apply_kr_font(r, bold=bold, size_pt=size_pt, color=color)


def _add_horizontal_rule(doc) -> None:
    from docx.oxml import OxmlElement

    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ──────────────────────────────────────────────────────────────────────────────
# 문서 빌드
# ──────────────────────────────────────────────────────────────────────────────

def build_document(output_path: Path) -> None:
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # ── 표지 ──
    _add_paragraph(doc, "Denvia", bold=True, size_pt=24,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, "치과 보험청구·데스크 행정 AI 어시스턴트",
                   size_pt=11, align=WD_ALIGN_PARAGRAPH.CENTER,
                   color=RGBColor(0x55, 0x55, 0x55))
    _add_paragraph(doc, "", size_pt=8)
    _add_paragraph(doc, DOC_TITLE, bold=True, size_pt=22,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(doc, f"— {DOC_SUBTITLE} —", size_pt=12,
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   color=RGBColor(0x55, 0x55, 0x55))
    _add_paragraph(doc, f"작성일: {DOC_DATE}", size_pt=11,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=24)

    # ── 0. 문서 개요 ──
    _add_heading(doc, "1. 문서 개요", level=1)

    overview_tbl = doc.add_table(rows=0, cols=2)
    overview_tbl.autofit = False
    overview_tbl.columns[0].width = Cm(4.0)
    overview_tbl.columns[1].width = Cm(12.5)
    for label, value in [
        ("프로젝트명", PROJECT_NAME),
        ("서비스 URL", SERVICE_URL),
        ("작성일", DOC_DATE),
        ("문서 목적",
         "계약서 및 MVP 범위 요약서에 정한 개발 범위가 모두 완료되어, "
         "클라이언트가 납품 내역을 확인하고 수령을 공식 확인하기 위한 문서입니다."),
        ("효력",
         "본 확인서에 양 당사자가 서명 또는 기명 날인한 시점에 납품이 공식적으로 완료되며, "
         "계약서 제11조에 따른 무상 하자보수 기간이 개시됩니다."),
    ]:
        row = overview_tbl.add_row().cells
        _set_cell_text(row[0], label, bold=True, size_pt=10)
        _shade_cell(row[0], "F2F2F2")
        _set_cell_text(row[1], value, size_pt=10)
    _add_paragraph(doc, "", size_pt=6, space_after_pt=12)

    # ── 2. 계약 범위 vs 실제 납품 대조표 ──
    _add_heading(doc, "2. 계약 범위 vs 실제 납품 대조표", level=1)
    _add_paragraph(
        doc,
        "계약서 부록 및 MVP 범위 요약서에 명시된 모든 항목과, 추가 협의로 확정된 사항이 "
        "운영 서버(denvia.ai.kr)에 반영되어 정상 동작 중임을 확인합니다.",
        size_pt=11,
    )

    scope_tbl = doc.add_table(rows=1, cols=3)
    scope_tbl.autofit = False
    scope_tbl.columns[0].width = Cm(5.0)
    scope_tbl.columns[1].width = Cm(9.0)
    scope_tbl.columns[2].width = Cm(2.5)
    hdr = scope_tbl.rows[0].cells
    for i, h in enumerate(["계약 범위", "실제 납품 내역", "상태"]):
        _set_cell_text(hdr[i], h, bold=True, size_pt=10)
        _shade_cell(hdr[i], "E0E0E0")
    for item in SCOPE_VS_DELIVERY:
        cells = scope_tbl.add_row().cells
        _set_cell_text(cells[0], item["scope"], bold=True, size_pt=10)
        _set_cell_text(cells[1], item["delivery"], size_pt=10)
        _set_cell_text(cells[2], item["status"], size_pt=10,
                       color=RGBColor(0x10, 0x80, 0x30))
    _add_paragraph(doc, "", size_pt=6, space_after_pt=12)

    # ── 3. 운영 자산 인계 ──
    _add_heading(doc, "3. 운영 자산 인계 (서버·도메인·계정)", level=1)
    _add_paragraph(
        doc,
        "아래 자산의 소유권 및 접근 권한이 클라이언트에 이전됩니다. 민감한 ID/PW·API 키는 "
        "본 확인서와 별도로 ‘자산·계정 통합 인계서’를 통해 안전한 채널로 전달됩니다.",
        size_pt=11,
    )

    for asset in ASSET_HANDOVER:
        _add_heading(doc, asset["category"], level=3)
        asset_tbl = doc.add_table(rows=0, cols=2)
        asset_tbl.autofit = False
        asset_tbl.columns[0].width = Cm(4.5)
        asset_tbl.columns[1].width = Cm(12.0)
        for label, value in asset["rows"]:
            row = asset_tbl.add_row().cells
            _set_cell_text(row[0], label, bold=True, size_pt=10)
            _shade_cell(row[0], "F2F2F2")
            _set_cell_text(row[1], value, size_pt=10)
        _add_paragraph(doc, "", size_pt=4)

    # ── 4. 유지보수 범위 및 면책 ──
    _add_heading(doc, "4. 유지보수 범위 및 면책", level=1)
    _add_paragraph(
        doc,
        "본 확인서 서명일부터 계약서(제11조)에 정한 무상 하자보수가 개시되며, "
        "이후에는 별도 유지보수 계약을 통해 운영 지원이 진행됩니다.",
        size_pt=11,
    )

    maint_tbl = doc.add_table(rows=0, cols=2)
    maint_tbl.autofit = False
    maint_tbl.columns[0].width = Cm(4.5)
    maint_tbl.columns[1].width = Cm(12.0)
    for label, value in MAINTENANCE_TABLE:
        row = maint_tbl.add_row().cells
        _set_cell_text(row[0], label, bold=True, size_pt=10)
        _shade_cell(row[0], "F2F2F2")
        _set_cell_text(row[1], value, size_pt=10)
    _add_paragraph(doc, "", size_pt=6, space_after_pt=12)

    # ── 5. 검수 체크리스트 ──
    _add_heading(doc, "5. 검수 체크리스트", level=1)
    _add_paragraph(
        doc,
        "아래 항목을 직접 확인하신 뒤 ‘확인’ 칸에 V 표시 부탁드립니다. "
        "(워드·모바일·구글 문서 어디서 열어도 자유롭게 편집 가능합니다.)",
        size_pt=11,
    )

    for group in ACCEPTANCE_CHECKLIST:
        _add_heading(doc, group["title"], level=3)
        chk_tbl = doc.add_table(rows=1, cols=2)
        chk_tbl.autofit = False
        chk_tbl.columns[0].width = Cm(13.5)
        chk_tbl.columns[1].width = Cm(3.0)
        hdr = chk_tbl.rows[0].cells
        for i, h in enumerate(["확인 항목", "확인 (V)"]):
            _set_cell_text(hdr[i], h, bold=True, size_pt=10)
            _shade_cell(hdr[i], "E0E0E0")
        for it in group["items"]:
            cells = chk_tbl.add_row().cells
            _set_cell_text(cells[0], it, size_pt=10)
            # CLAUDE.md 규칙: SDT/체크박스 컨트롤 사용 금지. 평범한 대괄호로만 표시.
            _set_cell_text(cells[1], "[      ]", size_pt=11)
        _add_paragraph(doc, "", size_pt=4)

    # ── 6. 수령 확인·서명란 ──
    _add_horizontal_rule(doc)
    _add_heading(doc, "6. 수령 확인 및 서명", level=1)
    _add_paragraph(
        doc,
        "위 ‘2. 계약 범위 vs 실제 납품 대조표’ 및 ‘5. 검수 체크리스트’의 모든 항목을 "
        "확인하였으며, Denvia 프로젝트의 납품을 정상 수령하였음을 확인합니다.",
        size_pt=11,
    )
    _add_paragraph(doc, "", size_pt=6)

    # 계약서(자료/치과챗봇_계약서.docx) 제1조 당사자 정보 그대로 채움.
    # 서명/날인 칸만 공백 — 인쇄 후 실물 서명·인감용.
    SUPPLIER = {  # 을 = 개발사
        "회사명 / 상호": "Break The mould",
        "대표자": "정형우",
        "사업자번호 / 주민번호": "206-26-68137",
        "주소": "경기도 양주시 옥정서로 5길 60",
        "연락처": "010-2484-0289",
        "이메일": "btmdesign@naver.com",
        "납품(수령) 일자": DOC_DATE,
    }
    CLIENT = {  # 갑 = 클라이언트
        "회사명 / 상호": "치과챗봇 (Denvia)",
        "대표자": "이규성",
        "사업자번호 / 주민번호": "980602-1",
        "주소": "독산동 1004-12 더팰리스80",
        "연락처": "010-2323-2753",
        "이메일": "dlrbtjd357@naver.com",
        "납품(수령) 일자": DOC_DATE,
    }
    SIGN_LABELS = [
        "회사명 / 상호",
        "대표자",
        "사업자번호 / 주민번호",
        "주소",
        "연락처",
        "이메일",
        "납품(수령) 일자",
    ]

    sign_tbl = doc.add_table(rows=1, cols=3)
    sign_tbl.style = "Table Grid"
    sign_tbl.autofit = False
    sign_tbl.columns[0].width = Cm(4.0)
    sign_tbl.columns[1].width = Cm(6.25)
    sign_tbl.columns[2].width = Cm(6.25)

    # 헤더 행
    hdr_row = sign_tbl.rows[0].cells
    _set_cell_text(hdr_row[0], "구분", bold=True, size_pt=10)
    _shade_cell(hdr_row[0], "E0E0E0")
    _set_cell_text(hdr_row[1], "공급자 (을 / 개발사)", bold=True, size_pt=10)
    _shade_cell(hdr_row[1], "E0E0E0")
    _set_cell_text(hdr_row[2], "수령자 (갑 / 클라이언트)", bold=True, size_pt=10)
    _shade_cell(hdr_row[2], "E0E0E0")

    # 내용 행 — 계약서 값으로 미리 채움
    for label in SIGN_LABELS:
        row = sign_tbl.add_row().cells
        _set_cell_text(row[0], label, bold=True, size_pt=10)
        _shade_cell(row[0], "F2F2F2")
        _set_cell_text(row[1], SUPPLIER[label], size_pt=10)
        _set_cell_text(row[2], CLIENT[label], size_pt=10)

    # 서명·날인 행 — 실물 서명용 공백 (높이 확보)
    sign_row = sign_tbl.add_row().cells
    _set_cell_text(sign_row[0], "서명 또는\n기명 날인 (인)", bold=True, size_pt=10)
    _shade_cell(sign_row[0], "F2F2F2")
    _set_cell_text(sign_row[1], "\n\n\n", size_pt=10)
    _set_cell_text(sign_row[2], "\n\n\n", size_pt=10)

    _add_paragraph(doc, "", size_pt=8)
    _add_paragraph(
        doc,
        "※ 본 확인서는 2부 작성하여 공급자·수령자가 각 1부씩 보관합니다.",
        size_pt=10, color=RGBColor(0x55, 0x55, 0x55),
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    project_root = Path(__file__).resolve().parent.parent
    out = project_root / "자료" / "Denvia_최종납품확인서_2026-06-05.docx"
    build_document(out)
    print(f"[OK] {out}")


if __name__ == "__main__":
    main()
